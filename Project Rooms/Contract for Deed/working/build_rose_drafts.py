from datetime import datetime
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Inches, Pt


ROOT = Path(r"C:\Codex\Wiki Files\Project Rooms\Contract for Deed")
WORKBOOK = ROOT / "source" / "320 Rose project spreadsheet" / "28_Project Management - 320 Rose Pl.xlsm"
OUT = ROOT / "output"
SUFFIX = "DRAFT"


def money(value):
    return f"${float(value):,.2f}"


def date_text(value):
    if hasattr(value, "strftime"):
        return value.strftime("%B %-d, %Y") if False else value.strftime("%B %#d, %Y")
    return str(value)


def short_date(value):
    if hasattr(value, "strftime"):
        return value.strftime("%#m/%#d/%Y")
    return str(value)


def inclusive_months(start, end):
    if hasattr(start, "year") and hasattr(end, "year"):
        return (end.year - start.year) * 12 + (end.month - start.month) + 1
    return None


def get_docs_values():
    wb = openpyxl.load_workbook(str(WORKBOOK), data_only=True, read_only=True, keep_vba=False)
    ws = wb["Docs"]
    values = {}
    multi_value_rows = {"adverse conditions"}
    for col in range(1, ws.max_column + 1):
        key = ws.cell(1, col).value
        if key:
            values[str(key).strip()] = ws.cell(2, col).value
    for row in range(1, ws.max_row + 1):
        key = ws.cell(row, 1).value
        if key:
            clean_key = str(key).strip()
            if clean_key.lower() in multi_value_rows:
                cell_values = [
                    clean(ws.cell(row, col).value)
                    for col in range(2, ws.max_column + 1)
                ]
                row_value = "\n".join(str(value) for value in cell_values if value)
            else:
                row_value = ws.cell(row, 2).value
            if clean_key in values:
                i = 2
                while f"{clean_key}__{i}" in values:
                    i += 1
                values[f"{clean_key}__{i}"] = row_value
            else:
                values[clean_key] = row_value
    return values


def clean(value):
    if value in (None, ""):
        return ""
    if isinstance(value, (int, float)) and float(value) == 0:
        return ""
    text = str(value).strip()
    return "" if text == "0" else text


def normalize_values(v):
    selling_seller = clean(v.get("SellingSeller") or v.get("Selling -Seller"))
    trust = clean(v.get("Trust")) or selling_seller
    trustee = clean(v.get("Trustee")) or "Investment Services LLC"
    # CFD seller/trustee execution stays with Wes; spreadsheet manager-like
    # fields may refer to other project roles and should not replace the signer.
    manager = "Wes Browning"
    buyer1 = clean(v.get("SellingBuyer") or v.get("Selling-Buyer") or v.get("Selling-Buyer1"))
    buyer2 = clean(v.get("Selling-Buyer2") or v.get("SellingBuyer2"))
    buyer = " and ".join(part for part in [buyer1, buyer2] if part)
    sale_price = float(v.get("SellingPurchasePrice") or v.get("Selling Purchase Price:") or 0)
    down_payment = float(v.get("SellingDownPayment:") or v.get("Selling Down Payment:") or 0)
    earnest_money = float(
        v.get("SellingEarnestMoney:")
        or v.get("Selling Earnest Money:")
        or 0
    )
    remaining_down_payment = down_payment - earnest_money
    loan_amount = float(v.get("LoanAmount:") or v.get("Loan Amount:") or (sale_price - down_payment))
    monthly_pi = float(v.get("Monthly Payment1") or v.get("PrincipalInterst") or 0)
    insurance = float(v.get("PropertyInsurance") or v.get("Property Insurance:") or 0)
    tax = float(v.get("TaxEscrow") or v.get("Tax Escrow:") or 0)
    total_payment = monthly_pi + insurance + tax
    loan_start = v.get("Loan Start1") or v.get("LoanStart1")
    loan_end = v.get("Loan End1") or v.get("LoanEnd1")
    payment_count = inclusive_months(loan_start, loan_end)
    buyer_address_parts = [
        clean(v.get("Selling-Buyer Add1")),
        clean(v.get("Selling-Buyer Add2")),
        clean(v.get("Selling-Buyer Add1__2")),
        clean(v.get("Selling-Buyer Add2__2")),
        clean(v.get("SellingBuyerAddress")),
        clean(v.get("Selling Buyer Address")),
        clean(v.get("SellingBuyerAddress1")),
        clean(v.get("Selling Buyer Address1")),
        clean(v.get("PurchaserAddress")),
        clean(v.get("Purchaser Address")),
    ]
    buyer_address = ", ".join(str(part).strip() for part in buyer_address_parts if part not in (None, ""))
    seller = clean(trust) or selling_seller
    adverse = clean(
        v.get("Adverse Conditions")
        or v.get("AdverseConditions")
        or v.get("AdverseCondition")
        or v.get("Adverse Condition")
        or v.get("Pending Orders or Adverse Conditions")
        or v.get("Pending Orders")
    )
    if not adverse:
        adverse = "NOTE FOR ATTORNEY REVIEW: Confirm whether any pending orders, liens, adverse conditions, title matters, or required disclosures should be listed here before signing."

    return {
        "seller": str(seller),
        "selling_seller": str(selling_seller),
        "trust": str(clean(trust)),
        "trustee": str(trustee),
        "trustee_address": f"{clean(v.get('TrusteeAddress1') or v.get('Trustee Address1:'))}, {clean(v.get('TrusteeAddress2') or v.get('Trustee Address2:'))}".strip(", "),
        "manager": str(manager),
        "buyer": str(clean(buyer)),
        "buyer1": str(clean(buyer1)),
        "buyer2": str(clean(buyer2)),
        "buyer_address": buyer_address or "[BUYER ADDRESS TO BE INSERTED]",
        "property_address": str(v.get("PropertyAddress") or v.get("Address") or ""),
        "property_city_state": str(v.get("PropertyCityState") or v.get("City-State") or ""),
        "county": str(clean(v.get("PropertyCounty") or v.get("County"))),
        "parcel": str(clean(v.get("PropertyParcelID") or v.get("Property Parcel ID"))),
        "legal": str(v.get("LegalDescription") or v.get("Legal Description") or ""),
        "brief_legal": str(v.get("BriefLegalDescription") or v.get("Brief Legal Description") or ""),
        "sale_price": sale_price,
        "down_payment": down_payment,
        "earnest_money": earnest_money,
        "remaining_down_payment": remaining_down_payment,
        "loan_amount": loan_amount,
        "monthly_pi": monthly_pi,
        "insurance": insurance,
        "tax": tax,
        "total_payment": total_payment,
        "term_months": payment_count or int(float(v.get("TermMonths") or v.get("TermMonths1") or 360)),
        "term_years": int(float(v.get("Term1") or v.get("TermYears") or 30)),
        "interest": str(v.get("Interestrate1") or v.get("Interest rate1") or "7.500%"),
        "loan_start": loan_start,
        "loan_end": loan_end,
        "doc_month": int(float(v.get("DocMonth") or v.get("Doc Month:") or datetime.now().month)),
        "doc_year": int(float(v.get("DocYear") or v.get("Doc Year") or datetime.now().year)),
        "prepared_by": "Capital City Law, 116 N. Person Street, Raleigh, NC 27601",
        "additional_fees": 0,
        "adverse": adverse,
    }


def setup_doc(title):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(11.5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DRAFT - FOR REVIEW")
    run.bold = True
    run.font.size = Pt(10)
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(title)
    r.bold = True
    r.font.size = Pt(15)
    return doc


def add_page_numbers(doc):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)


def add_label(doc, label, value):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    p.add_run(str(value))


def add_label_red_value(doc, label, value):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    rv = p.add_run(str(value))
    rv.font.color.rgb = RGBColor(192, 0, 0)


def add_red_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.color.rgb = RGBColor(192, 0, 0)
    r.bold = bold
    return p


def add_sig_block(doc, party_label, name, manager=None):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(party_label).bold = True
    doc.add_paragraph(name)
    doc.add_paragraph("By: ______________________________________ (SEAL)")
    if manager:
        doc.add_paragraph(f"Printed Name: {manager}")
        doc.add_paragraph("Title: Manager")
    else:
        doc.add_paragraph("Printed Name: ______________________________________")
    doc.add_paragraph("Date: ______________________")


def add_seller_trustee_sig_block(doc, x):
    doc.add_paragraph()
    add_red_paragraph(doc, "SELLER:", bold=True)
    add_red_paragraph(doc, x["trust"] or x["seller"])
    add_red_paragraph(doc, f"By: {x['trustee']}, Trustee")
    add_red_paragraph(doc, "By: ______________________________________ (SEAL)")
    add_red_paragraph(doc, f"Printed Name: {x['manager']}, Manager")
    add_red_paragraph(doc, f"Title: Manager of {x['trustee']}, Trustee")
    add_red_paragraph(doc, "Date: ______________________")


def add_notary(doc, person_text, year):
    doc.add_paragraph("STATE OF NORTH CAROLINA")
    doc.add_paragraph("COUNTY OF ________________________")
    doc.add_paragraph(
        f"I, ______________________________, a Notary Public, certify that {person_text} personally appeared before me this day and acknowledged the due execution of the foregoing instrument."
    )
    doc.add_paragraph(f"Witness my hand and official seal, this ____ day of __________________, {year}.")
    doc.add_paragraph("Notary Public: ______________________________")
    doc.add_paragraph("My commission expires: ______________________")


def build_contract(x):
    doc = setup_doc("CONTRACT FOR DEED")
    doc.add_paragraph("1. The Parties:")
    add_red_paragraph(doc, f"{x['trust']}, by and through {x['trustee']}, Trustee, (Seller)")
    doc.add_paragraph("Whose principal address is:")
    doc.add_paragraph(x["trustee_address"] or "[SELLER ADDRESS TO BE INSERTED]")
    doc.add_paragraph("And the following Purchaser:")
    doc.add_paragraph(f"{x['buyer']}, (Purchaser)")
    doc.add_paragraph("Address:")
    doc.add_paragraph(x["buyer_address"])

    doc.add_paragraph("2. Date of Agreement: This Contract for Deed is made and effective as of the date of the last signature set forth below each party's signature.")

    doc.add_paragraph("3. Real Property: The Property which is subject to the Contract for Deed is described as follows (and fully described in the attached Exhibit A):")
    add_label(doc, "Property Address: ", "")
    doc.add_paragraph(x["property_address"])
    doc.add_paragraph(x["property_city_state"])
    doc.add_paragraph("Legal Description")
    doc.add_paragraph(x["brief_legal"] or x["legal"])
    add_label(doc, "County of: ", x["county"])

    doc.add_paragraph("4. Sales Price: Total Purchase Price to be paid by Purchaser is payable as follows:")
    add_label(doc, "a. Binder deposit which will remain as a binder until closing, unless sooner forfeited or returned, according to the provisions in this Agreement. ", money(x["down_payment"]))
    doc.add_paragraph("b. Additional binder deposit due within ______days after date of this agreement.")
    doc.add_paragraph("c. Balance due upon Seller's vacancy or Closing Date, whichever is last, (not including Purchaser's closing cost, prepaid items")
    doc.add_paragraph("d. or prorations) in U.S. cash or locally drawn certified or cashier's check approximately.")
    doc.add_paragraph("e. Proceeds of a new loan to be executed by Purchaser to any lender other than Seller.")
    add_label(doc, "f. Purchase money loan to Seller on terms set forth in Paragraph 2B. ", money(x["loan_amount"]))
    doc.add_paragraph("g. Other financing _____________________________________________________.")
    doc.add_paragraph("h. Existing mortgage balance encumbering the Property to be taken subject to by Purchaser (approximately).")
    add_label(doc, "i. Total Purchase Price - approximately______ exactly_______ ", money(x["sale_price"]))

    doc.add_paragraph("5. ADDITIONAL CHARGES AND FEES: Purchaser agrees to provide additional charges and fees associated with the recording of a Memorandum of (this) Contract for Deed and administrative costs of the Seller.")
    doc.add_paragraph("a. Purchaser Will Pay the following options if they want them:")
    doc.add_paragraph("[X] Appraisal     [X] Survey")
    doc.add_paragraph("[X] Wood Destroying Organism Report     [X] Real estate brokerage fee")
    doc.add_paragraph("[   ] Other______________________________________")
    doc.add_paragraph("b. Seller Will Pay:")
    doc.add_paragraph("[X] Closing Costs     [X] Transfer tax")
    doc.add_paragraph("[X] Title insurance policy     [X] Attorney's fee")
    doc.add_paragraph("[X] Recording fees     [X] Note stamps     [X] Intangible tax")
    doc.add_paragraph("[X] Satisfaction and recording fee")

    doc.add_paragraph(f"6. INTEREST RATE: Unpaid principal shall accrue interest at a rate of {x['interest']} (APR). The Amount of Each Installment Payment above includes any interest payable, property Insurance, and escrowed Property Tax. Late payments shall accrue interest at a rate of four percent (4%) on any delinquent monies owed.")

    doc.add_paragraph("7. INSTALLMENT PAYMENTS:")
    doc.add_paragraph(f"a. For the first {x['term_months']} instalments:")
    add_label(doc, "i. Due Date of the First Installment Payment: ", short_date(x["loan_start"]))
    add_label(doc, "ii. Due Date of the Last Installment Payment: ", short_date(x["loan_end"]))
    doc.add_paragraph("iii. Amount of Each Installment Payment includes")
    add_label(doc, "1. Principal and Interest: ", money(x["monthly_pi"]))
    add_label(doc, "2. Property Insurance: ", money(x["insurance"]))
    add_label(doc, "3. Property Tax: ", money(x["tax"]))
    add_label(doc, "4. Total Monthly Payment: ", money(x["total_payment"]))
    add_label(doc, "iv. Total Number of Installment Payments: ", x["term_months"])

    doc.add_paragraph("8. AMORTIZATION SCHEDULE: A Full amortization schedule, see Exhibit \"B\" will be provided at closing.")
    doc.add_paragraph("9. PENDING ORDERS OR ADVERSE CONDITIONS: If any pending order of a public agency or other matter of public record adversely affecting the property is known, said pending order or condition is set out below by the Seller:")
    doc.add_paragraph(x["adverse"])

    doc.add_paragraph("10. RIGHT TO CURE DEFAULT: A PURCHASER'S RIGHTS UNDER A CONTRACT FOR DEED SHALL NOT BE FORFEITED EXCEPT AS PROVIDED IN N.C. GEN. STAT. CHAPTER 47H. A CONTRACT FOR DEED CANNOT BE FORFEITED UNLESS A BREACH HAS OCCURRED IN ONE OR MORE OF THE PURCHASER'S EXPRESS OBLIGATIONS UNDER THE CONTRACT AND THE CONTRACT PROVIDES THAT AS A RESULT OF SUCH BREACH THE SELLER IS ENTITLED TO FORFEIT THE CONTRACT. EXAMPLES OF BREACH OF CONTRACT INCLUDE LACK OF PAYMENT EXTENDING BEYOND 15 DAYS OF DUE DATE, PROHIBITED USE OF PROPERTY, FRAUDULENT IDENTIFICATION, AND FRAUDULENT PAYMENTS. FURTHERMORE, THE PURCHASER'S RIGHTS SHALL NOT BE FORFEITED UNTIL THE PURCHASER HAS BEEN NOTIFIED OF THE INTENT TO FORFEIT IN ACCORDANCE WITH N.C. GEN. STAT. 47H-4 AND BEEN GIVEN A RIGHT TO CURE THE DEFAULT AND HAS FAILED TO DO SO WITHIN THE TIME PERIOD ALLOWED. A TIMELY TENDER OF CURE SHALL REINSTATE THE CONTRACT FOR DEED.")
    doc.add_paragraph("11. PROHIBITED USE OF PROPERTY: Any unauthorized or unpermitted construction, improvements, excavation (to include clearing, logging, and planting), or otherwise altering the property from its original condition at the time this contract is executed is prohibited. Authorization may be given by the Seller for certain improvements if Purchaser submits an application in writing to make changes and receives authorization in writing from Seller. Authorization will be withheld or granted solely at the discretion of the Seller. Additionally, any use not permitted by local, state, and federal laws and ordinances is prohibited.")
    doc.add_paragraph("12. RESPONSIBILITY FOR TAXES, INSURANCE, REPAIRS, DUES: Seller shall be responsible for taxes, and dues for the Property throughout the contractual term set forth herein.")
    doc.add_paragraph("13. NO PREPAYMENT PENALTY: Purchaser may prepay some or all of the remaining Balance Owed on this contract at any time.")
    doc.add_paragraph("14. RESIDENTIAL PROPERTY DISCLOSURE: Purchaser waives the right to obtain a Residential Property Disclosure pursuant to Chapter 47E of the N.C. General Statutes and has had an adequate opportunity to inspect the Property. The Property described herein is provided \"as is\" to the Purchaser.")
    doc.add_paragraph("15. PROHIBITED USE OF PROPERTY: Any unauthorized or unpermitted construction, improvements, excavation (to include clearing, logging, and planting), or otherwise altering the property from its original condition at the time this contract is executed is prohibited. Authorization may be given by the Seller for certain improvements if Purchaser submits an application in writing to make changes and receives authorization in writing from Seller. Authorization will be withheld or granted solely at the discretion of the Seller. Additionally, any use not permitted by local, state, and federal laws and ordinances is prohibited.")
    doc.add_paragraph("16. RESPONSIBILITY FOR TAXES, INSURANCE, REPAIRS, DUES: Seller shall be responsible for taxes, and dues for the Property throughout the contractual term set forth herein. The Purchaser shall be responsible for insurance, repairing any damage to the property, and any citations assigned to the property.")
    doc.add_paragraph("17. NO PREPAYMENT PENALTY: Purchaser may prepay some or all of the remaining Balance Owed on this contract at any time.")
    doc.add_paragraph("18. RESIDENTIAL PROPERTY DISCLOSURE: Purchaser waives the right to obtain a Residential Property Disclosure pursuant to Chapter 47E of the N.C. General Statutes and has had an adequate opportunity to inspect the Property. The Property described herein is provided \"as is\" to the Purchaser.")
    doc.add_paragraph("19. Seller agrees to deliver the Property in its PRESENT AS IS CONDITION except as otherwise specified herein. Seller does hereby certify and represent that Seller has legal authority and capacity to convey the property with all improvements. Seller further certifies and represents that Seller knows of no latent defects to the property and knows of no facts materially affecting the value of the property except the following: Description of problems: Purchaser will be responsible to dispose of all contents. Purchaser [X] has [  ] has not inspected the property and HAS NOT RELIED UPON ANY REPRESENTATIONS MADE BY ANY REAL ESTATE AGENT in describing the property, and Purchaser accepts the property in its PRESENT AS IS CONDITION, except as otherwise specified herein.")
    doc.add_paragraph("20. Title Examination and Time for Closing:")
    doc.add_paragraph("If title evidence and survey show Seller is vested with a marketable title, subject to the usual exceptions contained in title insurance commitments, the transaction will be closed and the deed and other closing papers delivered on or before [ ] __________________, [X] 15 days after the date of acceptance, or [ ] _____________days after date of satisfaction of all conditions unless extended by other conditions of this Agreement or this Agreement is cancelled by the Purchaser.")
    doc.add_paragraph("If title evidence or survey reveal any defects which render the title unmarketable, Purchaser will have 7 days from receipt of title commitment and survey to notify Seller of such title defects and Seller agrees to use reasonable diligence to cure such defects at Seller's expense and will have 30 days to do so, in which event this transaction will be closed within 10 days after delivery to Purchaser of evidence that such defects have been cured.")
    doc.add_paragraph("21. Additional Terms, Conditions or Addenda (lettered a,b,c,d, etc.)")
    doc.add_paragraph("Seller will provide an Amortization Schedule, annually, indicating the complete payments required, with its components.")
    doc.add_paragraph("Seller agrees to email copies of the bills validating these amounts.")
    doc.add_paragraph("This is explicitly a Contract for Deed. No transfer of deed will be done until the terms of the purchase are satisfied. See the Promissory Note for details.")
    doc.add_paragraph("If the Purchaser defaults on terms of this contract or the promissory note, the remedy for the Seller is eviction, not foreclosure.")
    doc.add_paragraph("Closing will be required to be with Capital City Law.")
    doc.add_paragraph("22. Personal Property: included in the purchase price are all fixed equipment including ceiling fans, drapery hardware, attached lighting fixtures, mailbox, fence, plants and shrubbery as now installed on the property.")
    doc.add_paragraph("And these additional items ____________________________________________________________.")
    doc.add_paragraph("Items specifically excluded from ____________________________________________________.")
    doc.add_paragraph("23. RIGHT TO CANCEL: Purchaser may exercise the right to cancel this contract for deed until midnight of the third business day following of this contract for deed or delivery of a copy of the contract with the required minimum contents, whichever occurs later. If Purchaser cancels this contract, the seller shall, not later than the tenth day after the date the seller receives the purchaser's notice of cancellation, return to the Purchaser any and all property exchanged or payments made by the purchaser under the contract minus an offset of an amount equal to the fair rental value of the use of the property during the duration of the tenant's possession of the property plus an amount necessary to compensate the seller for any damages caused to the property by the purchaser beyond normal wear and tear.")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("24. Signatures:")
    doc.add_paragraph("Signed sealed on the date herein stated")
    add_seller_trustee_sig_block(doc, x)
    doc.add_paragraph(f"[ ] Agent   [X] Seller, by the signature below, acknowledge receipt of {money(x['down_payment'])} [  ] Cash   [   ] Check, as binder deposit, which is the amount mentioned in paragraph 4A of this Agreement.")
    add_notary(doc, f"{x['manager']}, Manager of {x['trustee']}, Trustee of {x['trust']}", x["doc_year"])
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("Signed sealed on the date herein stated")
    add_sig_block(doc, "PURCHASER:", x["buyer"])
    add_notary(doc, x["buyer"], x["doc_year"])
    path = OUT / f"320 Rose - Contract for Deed Agreement - {SUFFIX}.docx"
    doc.save(path)
    return path


def build_memorandum(x):
    doc = setup_doc("MEMORANDUM OF A CONTRACT FOR DEED")
    add_page_numbers(doc)
    doc.add_paragraph(f"Prepared By: {x['prepared_by']}    Return To: Preparer")
    doc.add_paragraph("STATE OF NORTH CAROLINA    COUNTY OF ________________________")
    doc.add_paragraph("This Memorandum of a Contract for Deed is made and entered into this ____ day of __________________, 20____, to memorialize the Contract for Deed of even date hereto.")
    add_label_red_value(doc, "1. PARTIES. ", f"The Contract for Deed is by and between {x['trust']} (Seller), by and through {x['trustee']}, Trustee, and {x['buyer']} (Purchaser).")
    add_label(doc, "Seller address: ", x["trustee_address"] or "[SELLER ADDRESS TO BE INSERTED]")
    add_label(doc, "Purchaser address: ", x["buyer_address"])
    add_label(doc, "2. REAL PROPERTY. ", f"{x['property_address']}, {x['property_city_state']}")
    add_label(doc, "Legal Description: ", x["legal"])
    add_label(doc, "County / Parcel ID: ", f"{x['county']} / {x['parcel']}")
    add_label(doc, "3. TERM OF CONTRACT. ", "")
    add_label(doc, "Due Date of the First Installment Payment: ", short_date(x["loan_start"]))
    add_label(doc, "Due Date of the Last Installment Payment: ", short_date(x["loan_end"]))
    add_label(doc, "Total Number of Installment Payments: ", x["term_months"])
    doc.add_paragraph("4. RIGHT TO CURE DEFAULT. A PURCHASER'S RIGHTS UNDER A CONTRACT FOR DEED SHALL NOT BE FORFEITED EXCEPT AS PROVIDED IN N.C. GEN. STAT. CHAPTER 47H. A TIMELY TENDER OF CURE SHALL REINSTATE THE CONTRACT FOR DEED.")
    doc.add_paragraph("5. RIGHT TO CANCEL. Purchaser may exercise the right to cancel this contract for deed until midnight of the third business day following this contract for deed or delivery of a copy of the contract with the required minimum contents, whichever occurs later.")
    doc.add_paragraph("EXHIBIT A - LEGAL DESCRIPTION")
    doc.add_paragraph(x["legal"])
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_seller_trustee_sig_block(doc, x)
    add_notary(doc, f"{x['manager']}, Manager of {x['trustee']}, Trustee of {x['trust']}", x["doc_year"])
    add_sig_block(doc, "PURCHASER:", x["buyer"])
    add_notary(doc, x["buyer"], x["doc_year"])
    path = OUT / f"320 Rose - Memorandum of Contract for Deed - {SUFFIX}.docx"
    doc.save(path)
    return path


def build_note(x):
    doc = setup_doc("PROMISSORY NOTE")
    doc.add_paragraph("STATE OF NORTH CAROLINA    COUNTY OF ________________________")
    add_label(doc, "Date of Note: ", "[DATE OF SIGNING]")
    add_label(doc, "Maker: ", x["buyer"])
    add_label_red_value(doc, "Holder: ", f"{x['trust']}, by and through {x['trustee']}, Trustee")
    add_label(doc, "Holder Address: ", x["trustee_address"] or "[HOLDER ADDRESS TO BE INSERTED]")
    add_label(doc, "Principal Sum: ", money(x["loan_amount"]))
    doc.add_paragraph(f"FOR VALUE RECEIVED, Maker promises to pay Holder the principal sum of {money(x['loan_amount'])}, with interest from the date of this Note on the unpaid balance until paid or until default, in lawful money of the United States of America, according to the following terms.")
    doc.add_paragraph(f"Terms: {x['interest']} APR. Monthly principal and interest payment is {money(x['monthly_pi'])}. Payments begin {short_date(x['loan_start'])} and continue through {short_date(x['loan_end'])}, for {x['term_months']} scheduled installments, unless sooner paid or modified by written agreement of the parties.")
    doc.add_paragraph("This Note is made in connection with and secured by the Contract for Deed for the property described as 320 Rose Pl, Wendell, NC 27591, and by any security instrument executed with the Contract for Deed.")
    doc.add_paragraph("Maker may prepay all or any part of this Note at any time without penalty unless otherwise agreed in writing.")
    doc.add_paragraph("In the event of default in payment of any installment of principal or interest due under this Note, if such default is not cured within ten (10) days after written notice to Maker, Holder may declare the remaining principal sum, together with accrued interest and all other sums due, immediately due and payable.")
    doc.add_paragraph("Upon default, Holder may employ an attorney to enforce Holder's rights and remedies, and Maker agrees to pay reasonable attorney's fees not exceeding fifteen percent (15%) of the outstanding balance owing on this Note, plus all reasonable expenses incurred by Holder in exercising rights and remedies upon default. This Note shall be construed in accordance with the laws of the State of North Carolina.")
    doc.add_paragraph("IN TESTIMONY WHEREOF, Maker has executed this Promissory Note as of the date first written above.")
    add_sig_block(doc, "MAKER:", x["buyer"])
    add_notary(doc, x["buyer"], x["doc_year"])
    path = OUT / f"320 Rose - Promissory Note for Contract for Deed - {SUFFIX}.docx"
    doc.save(path)
    return path


def build_deed_of_trust(x):
    doc = setup_doc("NORTH CAROLINA DEED OF TRUST")
    doc.add_paragraph("After recording mail to: ______________________________")
    doc.add_paragraph("Trustee at: ______________________________")
    doc.add_paragraph(f"Prepared by: {x['prepared_by']}")
    add_label(doc, "Brief Description for Index: ", f"{x['property_address']}, {x['county']} County")
    doc.add_paragraph(f"This Deed of Trust is made this ____ day of __________________, {x['doc_year']}, by and between:")
    add_label(doc, "GRANTOR: ", x["buyer"])
    add_label(doc, "TRUSTEE: ", x["trustee"])
    add_label_red_value(doc, "BENEFICIARY: ", f"{x['trust']}, by and through {x['trustee']}, Trustee")
    doc.add_paragraph("The designations Grantor, Trustee, and Beneficiary shall include their heirs, successors, and assigns, and shall apply in singular, plural, masculine, feminine, or neuter as required by context.")
    doc.add_paragraph("WITNESSETH")
    doc.add_paragraph(f"Whereas, Grantor is indebted to Beneficiary in the principal sum of {money(x['loan_amount'])}, as evidenced by a Promissory Note of even date herewith, the terms of which are incorporated by reference. The final due date for payments under said Promissory Note, if not sooner paid, is {short_date(x['loan_end'])}.")
    doc.add_paragraph("Additionally, Grantor and Beneficiary have entered into a Contract for Deed, which outlines the terms, timeline, and conditions under which ownership of the property shall be transferred upon fulfillment of the contract's obligations.")
    doc.add_paragraph("Now, therefore, as security for said indebtedness, and any additional sums expended by Beneficiary pursuant to this Deed of Trust, Grantor conveys to Trustee, for the benefit of Beneficiary, a Deed of Trust on the property described as follows:")
    add_label(doc, "Property Address: ", f"{x['property_address']}, {x['property_city_state']}")
    add_label(doc, "Legal Description: ", x["legal"])
    add_label(doc, "County / Parcel ID: ", f"{x['county']} / {x['parcel']}")
    doc.add_paragraph("This Deed of Trust is intended to secure the Promissory Note and the Contract for Deed.")
    doc.add_paragraph("SECURITY AGREEMENT")
    doc.add_paragraph("If Grantor fully complies with all terms of the Promissory Note and Contract for Deed, this Deed of Trust shall be released at Beneficiary's expense. However, if Grantor defaults in any payment due under the Promissory Note, any obligations under the Contract for Deed, or any terms of this Deed of Trust, then Trustee, at Beneficiary's request, may initiate foreclosure proceedings and sell the property at public auction in accordance with North Carolina foreclosure laws.")
    doc.add_paragraph("COVENANTS AND AGREEMENTS")
    for clause in [
        "Grantor shall maintain insurance on the property against loss by fire, windstorm, and other casualties.",
        "Grantor shall timely pay all property taxes and assessments unless the parties' Contract for Deed states another payment arrangement.",
        "In the event of default, Beneficiary is entitled to collect all rents and profits from the property.",
        "Grantor shall maintain the property in good condition and prevent waste.",
        "If any portion of the property is taken through eminent domain, Beneficiary may apply any compensation received toward the secured debt or property improvements.",
        "Beneficiary has the right to replace Trustee at any time in writing.",
        "If Grantor defaults and fails to cure within fifteen (15) days of written notice, Beneficiary may direct Trustee to proceed under North Carolina law.",
    ]:
        doc.add_paragraph(clause, style=None)
    add_sig_block(doc, "GRANTOR:", x["buyer"])
    add_notary(doc, x["buyer"], x["doc_year"])
    path = OUT / f"320 Rose - Deed of Trust - Contract for Deed - {SUFFIX}.docx"
    doc.save(path)
    return path


def build_review_notes(x, paths):
    lines = [
        "# Rose Sale Draft Documents - Review Notes",
        "",
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## Drafts",
        "",
    ]
    for path in paths:
        lines.append(f"- `{path.name}`")
    lines.extend([
        "",
        "## Values Used From Docs Worksheet",
        "",
        f"- Seller shown in drafts: {x['trust']}, by and through {x['trustee']}, Trustee",
        f"- Prior selling-seller worksheet value: {x['selling_seller']}",
        f"- Trustee manager for printed name/signature block: {x['manager']}",
        f"- Buyer: {x['buyer']}",
        f"- Buyer address: {x['buyer_address']}",
        f"- Property: {x['property_address']}, {x['property_city_state']}",
        f"- County / Parcel: {x['county']} / {x['parcel']}",
        f"- Sale price: {money(x['sale_price'])}",
        f"- Down payment: {money(x['down_payment'])}",
        f"- Remaining balance / loan amount: {money(x['loan_amount'])}",
        f"- Principal and interest payment: {money(x['monthly_pi'])}",
        f"- Insurance escrow: {money(x['insurance'])}",
        f"- Tax escrow: {money(x['tax'])}",
        f"- Total monthly payment shown in contract draft: {money(x['total_payment'])}",
        f"- Interest: {x['interest']}",
        f"- Loan dates used: {short_date(x['loan_start'])} to {short_date(x['loan_end'])}",
        f"- Installments counted from LoanStart1 through LoanEnd1: {x['term_months']}",
        "",
        "## Needs Review Before Signing",
        "",
        "- Buyer mailing address is not in the Docs worksheet and is left as a bracketed placeholder.",
        "- Seller/trustee formulation has been changed in red font: the trust is seller, and the trustee signs through its manager.",
        "- Confirm the trustee manager name. The Docs worksheet currently does not provide it, so the drafts use [MANAGER NAME].",
        "- Confirm adverse-condition/title disclosure language. The Docs worksheet currently has no inserted adverse condition text.",
        "- Confirm whether the Deed of Trust should be part of this Rose package.",
        "- Confirm whether the 60-installment LoanStart1-to-LoanEnd1 term should include a balloon/refinance/remaining-balance clause.",
        "- Confirm whether total monthly payment should include tax and insurance escrow or whether only principal and interest should be shown.",
    ])
    path = ROOT / "working" / "draft-review-notes.md"
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    x = normalize_values(get_docs_values())
    paths = [
        build_contract(x),
        build_memorandum(x),
        build_note(x),
        build_deed_of_trust(x),
    ]
    notes = build_review_notes(x, paths)
    for path in paths:
        print(path)
    print(notes)


if __name__ == "__main__":
    main()
