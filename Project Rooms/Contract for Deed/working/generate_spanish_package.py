from __future__ import annotations

import io
import re
from contextlib import redirect_stdout
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import generate_spanish_contract
from package_delivery import PackageItem, deliver_package_item
from package_doc_footer import stamp_docx_footer


PROJECT = Path(r"C:\Codex\Wiki Files\Project Rooms\Contract for Deed")
OUTPUT = PROJECT / "output"
ENGLISH_TERM = OUTPUT / "320 Rose - Term Sheet - DRAFT.docx"
ENGLISH_ACK = OUTPUT / "320 Rose - Buyer Acknowledgment Addendum - DRAFT.docx"
SPANISH_TERM = OUTPUT / "320 Rose Pl - Term Sheet - SPANISH DRAFT.docx"
SPANISH_ACK = OUTPUT / "320 Rose Pl - Buyer Acknowledgment Addendum - SPANISH DRAFT.docx"
TEAMS_ROOT = Path(
    r"C:\Users\wesbr\Buy Your Home\Buy Your Home - Property\28-SYH-320 Rose Pl"
    r"\Selling\Ever Cardoza\Contract Package"
)
TEAMS_SPANISH = TEAMS_ROOT / "Spanish Package"
TEAMS_SPANISH_ARCHIVE = TEAMS_ROOT / "Archive" / "Spanish Package"

TERM_LABELS = {
    "Seller": "Vendedor",
    "Seller / trustee address": "Direccion del vendedor / trustee",
    "Purchaser": "Comprador",
    "Purchaser address": "Direccion del comprador",
    "Property": "Propiedad",
    "County": "Condado",
    "Parcel ID": "ID de parcela",
    "Legal description summary": "Resumen de descripcion legal",
    "Proposed purchase price": "Precio de compra propuesto",
    "Earnest money due at signing": "Dinero de arras pagadero al firmar",
    "Total down payment": "Pago inicial total",
    "Remaining down payment due at closing": "Pago inicial restante pagadero al cierre",
    "Amount financed": "Monto financiado",
    "Interest rate": "Tasa de interes",
    "Financing term": "Plazo de financiamiento",
    "Proposed first payment date": "Fecha propuesta del primer pago",
    "Proposed final payment date": "Fecha propuesta del pago final",
    "Principal and interest": "Principal e interes",
    "Property insurance escrow": "Reserva para seguro de propiedad",
    "Property tax escrow": "Reserva para impuestos de propiedad",
    "Estimated total monthly payment": "Pago mensual total estimado",
}

ACK_TRANSLATIONS = {
    "Buyer understands this is a Contract for Deed, not a deed transfer at signing or closing.": "El Comprador entiende que esta operacion es un Contrato de Compraventa con Escritura Diferida y no una transferencia de escritura al firmar ni al cierre.",
    "Buyer understands Seller retains legal title until Buyer satisfies the contract terms and title is transferred later as provided in the contract.": "El Comprador entiende que el Vendedor conserva el titulo legal hasta que el Comprador cumpla los terminos del contrato y el titulo se transfiera posteriormente segun lo dispuesto en el contrato.",
    "Buyer understands Buyer will have payment obligations under both the Contract for Deed and the Promissory Note.": "El Comprador entiende que tendra obligaciones de pago bajo el Contrato de Compraventa con Escritura Diferida y el Pagare.",
    "Buyer understands the monthly payment amount, due date, interest rate, loan term, and payment schedule.": "El Comprador entiende el monto del pago mensual, la fecha de vencimiento, la tasa de interes, el plazo del prestamo y el calendario de pagos.",
    "Buyer understands earnest money is paid at signing and credited toward the total down payment.": "El Comprador entiende que el dinero de arras se paga al firmar y se acredita al pago inicial total.",
    "Buyer understands the remaining down payment balance is due at closing.": "El Comprador entiende que el saldo restante del pago inicial vence en el cierre.",
    "Buyer understands payments will be made by ACH draft unless otherwise agreed in writing.": "El Comprador entiende que los pagos se haran mediante debito automatico ACH salvo acuerdo escrito en contrario.",
    "Buyer understands the property is being accepted in its stated condition, subject to the contract terms and any written addenda.": "El Comprador entiende que acepta la propiedad en la condicion indicada, sujeto a los terminos del contrato y cualquier anexo escrito.",
    "Buyer understands Seller has disclosed the listed pending orders, liens, mortgages, or adverse conditions known to Seller.": "El Comprador entiende que el Vendedor ha revelado las ordenes pendientes, gravamenes, hipotecas o condiciones adversas conocidas por el Vendedor.",
    "Buyer understands existing liens may remain during the contract term and that lienholders may have rights against the property if required payments are not made.": "El Comprador entiende que los gravamenes existentes pueden permanecer durante el plazo del contrato y que los acreedores garantizados pueden tener derechos contra la propiedad si los pagos requeridos no se hacen.",
    "Buyer understands the Contract or a Memorandum of Contract for Deed will be recorded with the county Register of Deeds.": "El Comprador entiende que el Contrato o un Memorandum del Contrato de Compraventa con Escritura Diferida se registrara en el Registro de Escrituras del condado.",
    "Buyer understands Buyer has a statutory right to cancel the Contract for Deed until midnight of the third business day after execution or delivery, whichever is later.": "El Comprador entiende que tiene un derecho legal de cancelar el Contrato de Compraventa con Escritura Diferida hasta la medianoche del tercer dia habil despues de la firma o entrega, lo que ocurra mas tarde.",
    "Buyer understands Buyer has a right to cure default before forfeiture as provided by North Carolina law and the contract.": "El Comprador entiende que tiene derecho a subsanar un incumplimiento antes de la perdida de derechos, segun la ley de Carolina del Norte y el contrato.",
    "Buyer understands Seller must provide periodic account statements during the contract term.": "El Comprador entiende que el Vendedor debe proporcionar estados de cuenta periodicos durante el plazo del contrato.",
    "Buyer understands Buyer may prepay the balance without penalty except as specifically allowed by the contract and applicable law.": "El Comprador entiende que puede pagar por adelantado el saldo sin penalidad, salvo que el contrato y la ley aplicable permitan algo distinto.",
    "Buyer understands responsibilities for taxes, insurance, dues, repairs, and other property charges are governed by the contract.": "El Comprador entiende que las responsabilidades por impuestos, seguro, cuotas, reparaciones y otros cargos de la propiedad se rigen por el contrato.",
    "Buyer understands Seller is not giving Buyer legal advice.": "El Comprador entiende que el Vendedor no le esta dando asesoria legal.",
    "Buyer understands Buyer may hire Buyer's own attorney, at Buyer's expense, before signing and before closing.": "El Comprador entiende que puede contratar su propio abogado, a su propio costo, antes de firmar y antes del cierre.",
    "Buyer understands Seller's attorney may require changes, and the parties may sign the contract again at closing or execute an amendment.": "El Comprador entiende que el abogado del Vendedor puede requerir cambios, y que las partes pueden firmar el contrato nuevamente en el cierre o ejecutar una enmienda.",
    "Buyer acknowledges that Buyer has read the Contract for Deed, Promissory Note, Memorandum, and all addenda before signing.": "El Comprador reconoce que ha leido el Contrato de Compraventa con Escritura Diferida, el Pagare, el Memorandum y todos los anexos antes de firmar.",
}


def plain(value: str) -> str:
    return " ".join(str(value).split())


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    paragraph.paragraph_format.space_after = Pt(0)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_styles(doc: Document, compact: bool = False) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75 if compact else 1)
    section.bottom_margin = Inches(0.75 if compact else 1)
    section.left_margin = Inches(0.7 if compact else 1)
    section.right_margin = Inches(0.7 if compact else 1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10 if compact else 11)
    normal.paragraph_format.space_after = Pt(5 if compact else 6)
    normal.paragraph_format.line_spacing = 1.05 if compact else 1.10

    for style_name, size, color in [
        ("Heading 1", 15 if compact else 16, "2E74B5"),
        ("Heading 2", 12 if compact else 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(6 if compact else 8)
        style.paragraph_format.space_after = Pt(3 if compact else 4)


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    set_cell_width(table.rows[0].cells[0], 3100)
    set_cell_width(table.rows[0].cells[1], 6260)
    set_cell_text(table.rows[0].cells[0], "Termino", bold=True)
    set_cell_text(table.rows[0].cells[1], "Monto / Detalle", bold=True)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "F2F4F7")
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_width(cells[0], 3100)
        set_cell_width(cells[1], 6260)
        set_cell_text(cells[0], TERM_LABELS.get(label, label), bold=True)
        set_cell_text(cells[1], value)
    doc.add_paragraph("")


def table_rows(table) -> list[tuple[str, str]]:
    rows = []
    for row in table.rows[1:]:
        rows.append((plain(row.cells[0].text), plain(row.cells[1].text)))
    return rows


def extract_term_adverse_items(doc: Document) -> list[str]:
    items = []
    collecting = False
    for paragraph in doc.paragraphs:
        text = plain(paragraph.text)
        if text == "Known Pending Orders or Adverse Conditions":
            collecting = True
            continue
        if text == "Next Steps":
            break
        if collecting and text:
            items.append(text)
    return items


def build_spanish_term_sheet(package_version: str) -> Path:
    source = Document(str(ENGLISH_TERM))
    subtitle = plain(source.paragraphs[1].text).replace(
        "Proposed Contract for Deed Terms for ", ""
    )
    adverse_items = extract_term_adverse_items(source)

    doc = Document()
    set_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Hoja de Terminos")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title.paragraph_format.space_after = Pt(3)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Terminos propuestos de contrato de compraventa con escritura diferida para {subtitle}").bold = True
    sub.paragraph_format.space_after = Pt(10)

    source_line = doc.add_paragraph()
    source_line.add_run("Fuente de los valores: ").bold = True
    source_line.add_run("Hoja Docs de la hoja de calculo del proyecto 320 Rose.")

    notice = doc.add_paragraph()
    notice.add_run("BORRADOR DE TRADUCCION AL ESPANOL: ").bold = True
    notice.add_run(
        "Este documento se proporciona solo como una traduccion preliminar de conveniencia. "
        "No es una traduccion legal certificada. Si existe algun conflicto entre este borrador en espanol y los documentos finales en ingles, los documentos finales en ingles controlan."
    )

    status = doc.add_paragraph()
    status.add_run("Estado: ").bold = True
    status.add_run(
        "Esta Hoja de Terminos resume los terminos comerciales propuestos para revision y discusion. "
        "No es el contrato final, no transfiere titulo, y esta sujeta a los documentos finales del contrato, las divulgaciones requeridas, la revision de calificacion del comprador y la revision de un abogado. "
        "Si esta Hoja de Terminos entra en conflicto con los documentos finales firmados, controlan los documentos finales firmados."
    )

    doc.add_heading("Partes y Propiedad", level=2)
    add_key_value_table(doc, table_rows(source.tables[0]))
    doc.add_heading("Terminos de Compra y Financiamiento", level=2)
    add_key_value_table(doc, table_rows(source.tables[1]))
    doc.add_heading("Pago Mensual Estimado", level=2)
    add_key_value_table(doc, table_rows(source.tables[2]))

    doc.add_heading("Ordenes Pendientes o Condiciones Adversas Conocidas", level=2)
    for item in adverse_items or ["None listed on the Docs worksheet."]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Proximos Pasos", level=2)
    for item in [
        "Confirmar la calificacion del comprador y los terminos comerciales finales.",
        "Preparar el Contrato de Compraventa con Escritura Diferida, el Memorandum del Contrato de Compraventa con Escritura Diferida, el Pagare, los reconocimientos requeridos y cualquier divulgacion requerida.",
        "Enviar los documentos para revision de abogado antes de la firma.",
        "Hacer que todos los documentos requeridos sean firmados, notarizados cuando corresponda, y registrados segun se requiera.",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("Acuse de Recibo", level=2)
    doc.add_paragraph("El Comprador acusa recibo de esta Hoja de Terminos para revision y discusion.")
    doc.add_paragraph("Comprador 1: ________________________________________ Fecha: ____________________")
    doc.add_paragraph("Comprador 2: ________________________________________ Fecha: ____________________")
    doc.add_paragraph("Vendedor / Trustee: ___________________________________ Fecha: ____________________")

    doc.save(str(SPANISH_TERM))
    stamp_docx_footer(SPANISH_TERM, package_version)
    return SPANISH_TERM


def add_label_value(paragraph, label: str, value: str) -> None:
    paragraph.add_run(label).bold = True
    paragraph.add_run(value)


def add_acknowledgment_table(doc: Document, source_table) -> list[str]:
    headers = [plain(cell.text) for cell in source_table.rows[0].cells]
    buyer1 = headers[1].replace(" Initials", "").strip()
    buyer2 = headers[2].replace(" Initials", "").strip()

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    widths = [520, 1150, 1150, 7080]
    spanish_headers = ["No.", f"{buyer1}\nIniciales", f"{buyer2}\nIniciales", "Reconocimiento"]
    for idx, (cell, header, width) in enumerate(zip(table.rows[0].cells, spanish_headers, widths)):
        set_cell_width(cell, width)
        set_cell_text(cell, header, bold=True, center=(idx < 3))
        set_cell_shading(cell, "F2F4F7")

    missing = []
    for row in source_table.rows[1:]:
        number = plain(row.cells[0].text)
        english = plain(row.cells[3].text)
        translated = ACK_TRANSLATIONS.get(english)
        if not translated:
            missing.append(english)
            translated = english
        cells = table.add_row().cells
        values = [number, "________", "________", translated]
        for idx, (cell, value, width) in enumerate(zip(cells, values, widths)):
            set_cell_width(cell, width)
            set_cell_text(cell, value, center=(idx < 3))
    return missing


def paragraph_value(doc: Document, label: str) -> str:
    for paragraph in doc.paragraphs:
        text = plain(paragraph.text)
        if text.startswith(label):
            return text[len(label) :].strip()
    return ""


def build_spanish_acknowledgment(package_version: str) -> Path:
    source = Document(str(ENGLISH_ACK))
    subtitle = plain(source.paragraphs[1].text).replace("Contract for Deed - ", "")
    buyer = paragraph_value(source, "Purchaser:")
    prop = paragraph_value(source, "Property:")
    seller = paragraph_value(source, "Seller:")
    price = paragraph_value(source, "Purchase price:")
    financed = paragraph_value(source, "Amount financed:")

    doc = Document()
    set_styles(doc, compact=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Anexo de Reconocimientos del Comprador")
    run.bold = True
    run.font.size = Pt(17)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Contrato de Compraventa con Escritura Diferida - {subtitle}").bold = True

    notice = doc.add_paragraph()
    notice.add_run("BORRADOR DE TRADUCCION AL ESPANOL: ").bold = True
    notice.add_run(
        "Este documento se proporciona solo como una traduccion preliminar de conveniencia. "
        "No es una traduccion legal certificada. Si existe algun conflicto entre este borrador en espanol y los documentos finales en ingles, los documentos finales en ingles controlan."
    )

    intro = doc.add_paragraph()
    intro.add_run("Proposito: ").bold = True
    intro.add_run(
        "Este Anexo documenta los reconocimientos del Comprador sobre la naturaleza de la operacion de Contrato de Compraventa con Escritura Diferida. "
        "Cada Comprador debe poner sus iniciales en cada reconocimiento y firmar abajo."
    )

    add_label_value(doc.add_paragraph(), "Comprador: ", buyer)
    add_label_value(doc.add_paragraph(), "Propiedad: ", prop)
    add_label_value(doc.add_paragraph(), "Vendedor: ", seller)
    add_label_value(doc.add_paragraph(), "Precio de compra: ", price)
    add_label_value(doc.add_paragraph(), "Monto financiado: ", financed)

    doc.add_heading("Reconocimientos con Iniciales del Comprador", level=2)
    missing = add_acknowledgment_table(doc, source.tables[0])
    if missing:
        raise RuntimeError(f"Missing Spanish acknowledgment translations: {missing}")

    doc.add_heading("Firmas de los Compradores", level=2)
    for paragraph in source.paragraphs:
        text = plain(paragraph.text)
        if text.endswith("Date: ____________________") and ":" in text:
            doc.add_paragraph(text.replace(" Date: ", " Fecha: "))

    notary_started = False
    for paragraph in source.paragraphs:
        text = plain(paragraph.text)
        if text == "Notary Acknowledgment":
            notary_started = True
            doc.add_heading(text, level=2)
            continue
        if notary_started and text:
            doc.add_paragraph(text)

    doc.save(str(SPANISH_ACK))
    stamp_docx_footer(SPANISH_ACK, package_version)
    return SPANISH_ACK


def preflight() -> None:
    lock_files = []
    for root in [OUTPUT, TEAMS_SPANISH]:
        if root.exists():
            lock_files.extend(str(path) for path in root.rglob("~$*.docx"))
    missing = [str(path) for path in [ENGLISH_TERM, ENGLISH_ACK] if not path.exists()]
    if lock_files or missing:
        raise RuntimeError({"lock_files": lock_files, "missing": missing})


def run_contract_generator() -> str:
    captured = io.StringIO()
    with redirect_stdout(captured):
        generate_spanish_contract.main()
    output = captured.getvalue()
    print(output, end="")
    match = re.search(r"Package footer version:\s*(.+)", output)
    if not match:
        raise RuntimeError("Spanish contract generator did not report a package footer version.")
    return match.group(1).strip()


def deliver_spanish_doc(label: str, source: Path, target_name: str) -> dict:
    item = PackageItem(
        label=label,
        source=source,
        target=TEAMS_SPANISH / target_name,
        archive_dir=TEAMS_SPANISH_ARCHIVE,
    )
    return deliver_package_item(item, TEAMS_ROOT)


def main() -> None:
    preflight()
    package_version = run_contract_generator()
    term = build_spanish_term_sheet(package_version)
    ack = build_spanish_acknowledgment(package_version)
    records = [
        deliver_spanish_doc("Spanish Term Sheet", term, "320 Rose Pl - Term Sheet - SPANISH DRAFT.docx"),
        deliver_spanish_doc(
            "Spanish Buyer Acknowledgment Addendum",
            ack,
            "320 Rose Pl - Buyer Acknowledgment Addendum - SPANISH DRAFT.docx",
        ),
    ]
    print(f"Spanish package footer version: {package_version}")
    for record in records:
        print(record)


if __name__ == "__main__":
    main()
