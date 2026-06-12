from __future__ import annotations

import re
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


VERSION_RE = re.compile(r"(?P<date>\d{2}-\d{2}-\d{2})\s+V(?P<version>\d+)", re.I)


def version_date(value: date | None = None) -> str:
    value = value or date.today()
    return value.strftime("%y-%m-%d")


def read_footer_versions(path: Path) -> list[tuple[str, int]]:
    if not path.exists() or path.suffix.lower() != ".docx":
        return []
    versions: list[tuple[str, int]] = []
    try:
        with zipfile.ZipFile(path) as docx:
            footer_names = [name for name in docx.namelist() if name.startswith("word/footer")]
            for name in footer_names:
                text = docx.read(name).decode("utf-8", errors="ignore")
                for match in VERSION_RE.finditer(text):
                    versions.append((match.group("date"), int(match.group("version"))))
    except zipfile.BadZipFile:
        return []
    return versions


def next_package_version(existing_paths: list[Path], run_date: date | None = None) -> str:
    date_text = version_date(run_date)
    max_seen = 0
    for path in existing_paths:
        for seen_date, version in read_footer_versions(path):
            if seen_date == date_text:
                max_seen = max(max_seen, version)
    return f"{date_text} V{max_seen + 1}"


def clear_paragraph(paragraph) -> None:
    paragraph.text = ""


def add_field(paragraph, instruction: str, fallback: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def stamp_docx_footer(path: Path, package_version: str) -> dict:
    doc = Document(str(path))
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        run = paragraph.add_run("Page ")
        run.font.size = Pt(9)
        add_field(paragraph, "PAGE", "1")
        run = paragraph.add_run(" of ")
        run.font.size = Pt(9)
        add_field(paragraph, "NUMPAGES", "1")
        run = paragraph.add_run(f" | {package_version}")
        run.font.size = Pt(9)
    doc.save(str(path))
    return {"path": str(path), "status": "stamped", "package_version": package_version}


def stamp_docx_files(paths: list[Path], package_version: str) -> list[dict]:
    records = []
    seen: set[Path] = set()
    for path in paths:
        resolved = path.resolve()
        if resolved in seen:
            continue
        seen.add(resolved)
        if not path.exists():
            records.append({"path": str(path), "status": "missing"})
            continue
        if path.suffix.lower() != ".docx":
            records.append({"path": str(path), "status": "skipped_non_docx"})
            continue
        records.append(stamp_docx_footer(path, package_version))
    return records
