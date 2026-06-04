from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROW_HEIGHT_TWIPS = 264  # 13.2 pt; matches 12 pt text at 1.1 line spacing.


def set_row_height(row, height_twips=ROW_HEIGHT_TWIPS):
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(height_twips))
    tr_height.set(qn("w:hRule"), "exact")


def adjust_payment_table(path):
    doc = Document(str(path))
    if not doc.tables:
        raise RuntimeError(f"No tables found in {path}")
    table = doc.tables[0]
    for row in table.rows:
        set_row_height(row)
    doc.save(str(path))


def main():
    root = Path(r"C:\Codex\Wiki Files\Project Rooms\Contract for Deed")
    paths = [
        root / "reference" / "Rose contract prototype" / "320 Rose - Contract for Deed Agreement - PROTOTYPE.docx",
        root / "output" / "320 Rose - Contract for Deed Agreement - DRAFT.docx",
    ]
    for path in paths:
        adjust_payment_table(path)
        print(path)


if __name__ == "__main__":
    main()
