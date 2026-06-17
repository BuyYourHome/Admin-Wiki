from pathlib import Path
import re
import shutil

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt, RGBColor

from package_doc_footer import next_package_version, stamp_docx_footer


PROJECT = Path(r"C:\Codex\Wiki Files\Project Rooms\Contract for Deed")
ENGLISH_CONTRACT = PROJECT / "output" / "320 Rose - Contract for Deed Agreement - DRAFT.docx"
LEGACY_BILINGUAL_CONTRACT = PROJECT / "output" / "320 Rose - Contract for Deed Agreement - BILINGUAL SPANISH DRAFT.docx"
SPANISH_BASE_NAME = "320 Rose - Contract for Deed Agreement - BILINGUAL SPANISH DRAFT.docx"
TEAMS_SPANISH_BASE_NAME = "320 Rose Pl - Contract for Deed Agreement - BILINGUAL SPANISH DRAFT.docx"
TEAMS_SPANISH = Path(
    r"C:\Users\wesbr\Buy Your Home\Buy Your Home - Property\28-SYH-320 Rose Pl"
    r"\Selling\Ever Cardoza\Contract Package\Spanish Package"
)
TEAMS_ROOT = TEAMS_SPANISH.parent
TEAMS_SPANISH_ARCHIVE = Path(
    r"C:\Users\wesbr\Buy Your Home\Buy Your Home - Property\28-SYH-320 Rose Pl"
    r"\Selling\Ever Cardoza\Contract Package\Archive\Spanish Package"
)

SPANISH_STYLE = "Spanish Translation Draft"
SPANISH_BLUE = RGBColor(0x00, 0x66, 0xCC)
VISUAL_RIGHT_SHIFT = Pt(4)
SPANISH_CONTROL_NOTICE = (
    "Prevalece la version en ingles: El texto en espanol de este borrador bilingue "
    "se proporciona solo como una traduccion preliminar de conveniencia. Si existe "
    "algun conflicto entre el texto en ingles y el texto en espanol, prevalece el "
    "texto en ingles. El texto en espanol no es una traduccion legal certificada "
    "salvo que sea aprobado por separado por Wes, un abogado o un traductor calificado."
)
CONTROL_NOTICE_PREFIXES = (
    "English Version Controls:",
    "Prevalece la version en ingles:",
)
VERSION_PATTERN = re.compile(r"^v(?P<num>\d{2}) - 320 Rose(?: Pl)? - Contract for Deed Agreement - BILINGUAL SPANISH DRAFT\.docx$")
LIST_SECTION_ENDS = {
    "Sales Price:": ("ADDITIONAL CHARGES AND FEES:",),
    "ADDITIONAL CHARGES AND FEES:": ("INTEREST RATE:",),
}
MANUAL_TRANSLATIONS = {
    "Seller shall remain responsible for the pending orders or adverse conditions listed above. These matters may remain in place during the term of this Contract. To the extent any such matter must be released, satisfied, or otherwise resolved to convey marketable title upon completion of this Agreement's term, Seller shall cause it to be released, satisfied, or otherwise resolved before or upon payment of all amounts due herein, unless resolved earlier or otherwise agreed in writing.": [
        "El Vendedor seguira siendo responsable de las ordenes pendientes o condiciones adversas indicadas anteriormente. Estos asuntos podran permanecer vigentes durante el plazo de este Contrato. En la medida en que cualquiera de esos asuntos deba ser liberado, satisfecho o resuelto de otro modo para transmitir titulo comerciable al completarse el plazo de este Contrato, el Vendedor hara que sea liberado, satisfecho o resuelto antes o al momento del pago de todas las cantidades adeudadas en este documento, salvo que se resuelva antes o se acuerde lo contrario por escrito."
    ],
    "The parties shall sign the Contract again at closing, or execute any, with any such alterations incorporated. If such changes are material and Purchaser does not agree to proceed, Seller may elect to proceed under the previously signed Contract. If Seller does not so elect, Purchaser's due diligence funds shall be returned to Purchaser.": [
        "Las partes firmaran nuevamente el Contrato al cierre, o ejecutaran cualquier documento requerido, con dichas modificaciones incorporadas. Si tales cambios son materiales y el Comprador no acepta proceder, el Vendedor podra optar por proceder bajo el Contrato firmado anteriormente. Si el Vendedor no lo elige, los fondos de debida diligencia del Comprador seran devueltos al Comprador."
    ],
}


def is_spanish_translation(paragraph):
    if paragraph.style and paragraph.style.name == SPANISH_STYLE:
        return True
    for run in paragraph.runs:
        color = run.font.color
        if color is not None and color.rgb == SPANISH_BLUE:
            return True
    return False


def paragraph_text(paragraph):
    return " ".join(paragraph.text.split())


def is_control_notice_text(text):
    normalized = text.strip()
    return any(normalized.startswith(prefix) for prefix in CONTROL_NOTICE_PREFIXES)


def build_translation_memory(path):
    doc = Document(str(path))
    memory = {}
    last_english = None
    for paragraph in doc.paragraphs:
        text = paragraph_text(paragraph)
        if not text:
            continue
        if is_control_notice_text(text):
            continue
        if is_spanish_translation(paragraph):
            if last_english:
                memory.setdefault(last_english, []).append(paragraph.text)
        else:
            last_english = text
    return memory


def merge_translation_memory(memory, path):
    if not path.exists():
        return
    for english, translations in build_translation_memory(path).items():
        memory.setdefault(english, translations)


def versioned_spanish_files():
    output = PROJECT / "output"
    found = []
    for path in output.glob("v?? - 320 Rose - Contract for Deed Agreement - BILINGUAL SPANISH DRAFT.docx"):
        match = VERSION_PATTERN.match(path.name)
        if match:
            found.append((int(match.group("num")), path))
    return sorted(found)


def translation_memory_paths(output_path):
    candidates = []
    for _, path in reversed(versioned_spanish_files()):
        if path != output_path:
            candidates.append(path)
    output = PROJECT / "output"
    for path in output.glob("*BILINGUAL SPANISH DRAFT*.docx"):
        if path != output_path and path not in candidates:
            candidates.append(path)
    if LEGACY_BILINGUAL_CONTRACT.exists() and LEGACY_BILINGUAL_CONTRACT not in candidates:
        candidates.append(LEGACY_BILINGUAL_CONTRACT)
    return candidates


def next_output_path():
    versions = versioned_spanish_files()
    next_version = (versions[-1][0] + 1) if versions else 1
    return PROJECT / "output" / f"v{next_version:02d} - {SPANISH_BASE_NAME}"


def next_teams_archive_path():
    TEAMS_SPANISH_ARCHIVE.mkdir(parents=True, exist_ok=True)
    versions = []
    for path in TEAMS_SPANISH_ARCHIVE.glob("v?? - *Contract for Deed Agreement - BILINGUAL SPANISH DRAFT.docx"):
        match = VERSION_PATTERN.match(path.name)
        if match:
            versions.append(int(match.group("num")))
    next_version = (max(versions) + 1) if versions else 1
    return TEAMS_SPANISH_ARCHIVE / f"v{next_version:02d} - {TEAMS_SPANISH_BASE_NAME}"


def archive_existing_teams_copy():
    current = TEAMS_SPANISH / TEAMS_SPANISH_BASE_NAME
    if not current.exists():
        return None
    archive_path = next_teams_archive_path()
    shutil.move(str(current), str(archive_path))
    return archive_path


def current_package_docx_files():
    if not TEAMS_ROOT.exists():
        return []
    return sorted(path for path in TEAMS_ROOT.rglob("*.docx") if path.is_file())


def build_combined_translation_memory(output_path):
    memory = {}
    sources = translation_memory_paths(output_path)
    for path in sources:
        merge_translation_memory(memory, path)
    if not memory:
        raise FileNotFoundError(
            "Existing bilingual draft is required as translation source: "
            f"versioned Spanish output or {LEGACY_BILINGUAL_CONTRACT}"
        )
    return memory, sources


def ensure_spanish_style(doc):
    styles = doc.styles
    if SPANISH_STYLE in [style.name for style in styles]:
        style = styles[SPANISH_STYLE]
    else:
        style = styles.add_style(SPANISH_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Arial"
    style.font.size = Pt(9)
    style.font.color.rgb = SPANISH_BLUE
    style.paragraph_format.left_indent = None
    style.paragraph_format.right_indent = None
    return style


def style_chain(paragraph):
    style = paragraph.style
    while style is not None:
        yield style
        style = style.base_style


def effective_alignment(paragraph):
    if paragraph.alignment is not None:
        return paragraph.alignment
    for style in style_chain(paragraph):
        alignment = style.paragraph_format.alignment
        if alignment is not None:
            return alignment
    return None


def effective_left_indent(paragraph):
    if paragraph.paragraph_format.left_indent is not None:
        return paragraph.paragraph_format.left_indent
    for style in style_chain(paragraph):
        left_indent = style.paragraph_format.left_indent
        if left_indent is not None:
            return left_indent
    return None


def leading_tabs(paragraph):
    match = re.match(r"^\t+", paragraph.text or "")
    return match.group(0) if match else ""


def insert_paragraph_after(insert_after, text, style, format_source):
    new_p = OxmlElement("w:p")
    insert_after._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, insert_after._parent)
    new_paragraph.style = style
    new_paragraph.paragraph_format.first_line_indent = None
    new_paragraph.paragraph_format.right_indent = None
    if effective_alignment(format_source) == WD_ALIGN_PARAGRAPH.CENTER:
        new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        new_paragraph.paragraph_format.left_indent = None
    else:
        base_left_indent = effective_left_indent(format_source) or 0
        new_paragraph.paragraph_format.left_indent = base_left_indent + VISUAL_RIGHT_SHIFT
    run_text = text if new_paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER else f"{leading_tabs(format_source)}{text}"
    run = new_paragraph.add_run(run_text)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = SPANISH_BLUE
    return new_paragraph


def insert_control_notice_before(paragraph, style):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    new_paragraph.style = style
    new_paragraph.paragraph_format.left_indent = None
    new_paragraph.paragraph_format.right_indent = None
    new_paragraph.paragraph_format.first_line_indent = None
    new_paragraph.paragraph_format.space_after = Pt(8)
    run = new_paragraph.add_run(SPANISH_CONTROL_NOTICE)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = SPANISH_BLUE
    return new_paragraph


def insert_control_notice_above_paragraph_2(doc, style):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Date of Agreement:"):
            insert_control_notice_before(paragraph, style)
            return True
    return False


def is_list_or_structured_line(text):
    stripped = text.strip()
    if not stripped:
        return True
    if "(Seller)" in stripped or "(Purchaser)" in stripped:
        return True
    if any(token in stripped for token in ("Trustee", "LLC", "Trust dated", "Bolanos", "Sarmiento")):
        return True
    if stripped.startswith(("[X]", "[ ]", "[   ]", "(*)")):
        return True
    if stripped in {"Purchaser Will Pay the following options, if they apply:", "Seller Will Pay:"}:
        return True
    if stripped.startswith(("Property Address:", "Address:", "Legal Description", "County of:")):
        return True
    if stripped.startswith(("First Position Mortgage", "Deferred Balance", "Lien to ")):
        return True
    if "Lien to " in stripped and " amount " in stripped:
        return True
    if stripped.startswith(("STE ", "Suite ")):
        return True
    if re.search(r"\d{3,} .+", stripped):
        return True
    if re.search(r"\b(?:[A-Z]{2}|[A-Z][a-z]+)\s+\d{5}(?:-\d{4})?\b", stripped):
        return True
    if stripped.upper() == stripped and len(stripped.split()) <= 8:
        return True
    return False


def starts_list_section(text):
    stripped = text.strip()
    for start in LIST_SECTION_ENDS:
        if stripped.startswith(start):
            return start
    return None


def ends_list_section(text, active_start):
    stripped = text.strip()
    return any(stripped.startswith(prefix) for prefix in LIST_SECTION_ENDS.get(active_start, ()))


def first_signature_index(doc):
    signature_markers = (
        "Signatures:",
        "PURCHASER:",
        "SELLER:",
        "STATE OF:",
        "COUNTY OF:",
        "Notary",
    )
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if any(marker in text for marker in signature_markers):
            return idx
    return len(doc.paragraphs)


def main():
    if not ENGLISH_CONTRACT.exists():
        raise FileNotFoundError(ENGLISH_CONTRACT)
    bilingual_contract = next_output_path()
    memory, memory_sources = build_combined_translation_memory(bilingual_contract)
    work_path = bilingual_contract.with_suffix(".tmp.docx")
    shutil.copy2(ENGLISH_CONTRACT, work_path)

    doc = Document(str(work_path))
    style = ensure_spanish_style(doc)
    cutoff = first_signature_index(doc)
    inserted = 0
    missing = []

    original_paragraphs = list(doc.paragraphs[:cutoff])
    active_list_section = None
    for paragraph in original_paragraphs:
        key = paragraph_text(paragraph)
        if not key:
            continue
        if active_list_section:
            if ends_list_section(key, active_list_section):
                active_list_section = None
            else:
                continue
        list_section_start = starts_list_section(key)
        translate_heading_before_list = bool(list_section_start)
        if is_control_notice_text(key):
            continue
        if not translate_heading_before_list and is_list_or_structured_line(key):
            continue
        translations = memory.get(key) or MANUAL_TRANSLATIONS.get(key)
        if not translations:
            missing.append(key)
            if translate_heading_before_list:
                active_list_section = list_section_start
            continue
        insert_after = paragraph
        for translation in translations:
            insert_after = insert_paragraph_after(insert_after, translation, style, paragraph)
            inserted += 1
        if translate_heading_before_list:
            active_list_section = list_section_start

    if not insert_control_notice_above_paragraph_2(doc, style):
        raise RuntimeError("Could not find paragraph 2/Date of Agreement location for Spanish control notice.")

    doc.save(str(work_path))
    work_path.replace(bilingual_contract)

    package_version = next_package_version(current_package_docx_files())
    stamp_docx_footer(bilingual_contract, package_version)

    TEAMS_SPANISH.mkdir(parents=True, exist_ok=True)
    archived_teams_copy = archive_existing_teams_copy()
    teams_copy = TEAMS_SPANISH / TEAMS_SPANISH_BASE_NAME
    teams_status = "copied"
    try:
        shutil.copy2(bilingual_contract, teams_copy)
    except PermissionError as exc:
        teams_status = f"blocked: {exc}"

    print(bilingual_contract)
    print("Translation memory sources:")
    for source in memory_sources:
        print(f"- {source}")
    if archived_teams_copy:
        print(f"Archived prior Teams Spanish copy: {archived_teams_copy}")
    print(teams_copy)
    print(f"Teams copy status: {teams_status}")
    print(f"Package footer version: {package_version}")
    print(f"Inserted Spanish/control paragraphs: {inserted}")
    if missing:
        print(f"English paragraphs without Spanish memory: {len(missing)}")
        for item in missing[:10]:
            print(f"- {item[:160]}")


if __name__ == "__main__":
    main()
