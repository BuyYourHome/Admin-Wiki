from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(r"C:\Codex\Wiki Files\Project Rooms\Contract for Deed")
TEAMS_ROOT = Path(
    r"C:\Users\wesbr\Buy Your Home\Buy Your Home - Property\28-SYH-320 Rose Pl"
    r"\Selling\Ever Cardoza\Contract Package"
)
SPANISH_PACKAGE = TEAMS_ROOT / "Spanish Package"

SPANISH_DOC_LABELS = {
    "320 Rose Pl - Contract for Deed Agreement - BILINGUAL SPANISH DRAFT.docx": "Contract for Deed Agreement - Bilingual Spanish Draft",
    "320 Rose Pl - Term Sheet - SPANISH DRAFT.docx": "Term Sheet - Spanish Draft",
    "320 Rose Pl - Buyer Acknowledgment Addendum - SPANISH DRAFT.docx": "Buyer Acknowledgment Addendum - Spanish Draft",
}
SPANISH_DOC_ORDER = {name: index for index, name in enumerate(SPANISH_DOC_LABELS)}

DEFAULT_PATHS = [
    PROJECT_ROOT
    / "transactions"
    / "320 Rose Pl - Ever Cardoza"
    / "output"
    / "closing-checklist"
    / "320 Rose Pl - Ever Cardoza - Closing Package Cover Page.docx",
    PROJECT_ROOT
    / "transactions"
    / "320 Rose Pl - Ever Cardoza"
    / "output"
    / "clean"
    / "320 Rose Pl - Ever Cardoza - Closing Package Cover Page.docx",
    TEAMS_ROOT / "320 Rose Pl - Ever Cardoza - Closing Package Cover Page.docx",
]


def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def paragraph_after(element, paragraph):
    paragraph._p.addnext(element)


def clear_cell(cell):
    for paragraph in cell.paragraphs:
        paragraph.clear()


def set_font(run, size=10, bold=False):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold


def add_prepared_closing_table(doc, prepared_text):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)

    left, right = table.rows[0].cells
    set_cell_width(left, Inches(3.25))
    set_cell_width(right, Inches(3.25))
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    clear_cell(left)
    p_left = left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_left.add_run(prepared_text)
    set_font(run, 10)

    clear_cell(right)
    p_right = right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = p_right.add_run("Closing Date: ")
    set_font(label, 10, bold=True)
    blank = p_right.add_run("________________________")
    set_font(blank, 10)
    blank.font.highlight_color = WD_COLOR_INDEX.YELLOW
    shade_cell(right, "FFF59D")

    return table


def has_closing_date_field(doc):
    for table in doc.tables[:3]:
        table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if "Closing Date:" in table_text:
            return True
    for paragraph in doc.paragraphs[:8]:
        if "Closing Date:" in paragraph.text:
            return True
    return False


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def spanish_package_files():
    if not SPANISH_PACKAGE.exists():
        return []
    paths = [
        path
        for path in SPANISH_PACKAGE.glob("*.docx")
        if path.is_file() and not path.name.startswith("~$")
    ]
    paths.sort(key=lambda path: (SPANISH_DOC_ORDER.get(path.name, 999), path.name.lower()))
    return paths


def find_paragraph(doc, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    return None


def first_document_item_style(doc):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.endswith(".docx") or text.endswith(".pdf") or text.endswith(".zip"):
            return paragraph.style
    return None


def remove_existing_spanish_section(doc):
    start = find_paragraph(doc, "Spanish / Bilingual Drafts")
    if start is None:
        return False
    paragraphs = list(doc.paragraphs)
    start_index = paragraphs.index(start)
    end_index = len(paragraphs)
    for index in range(start_index + 1, len(paragraphs)):
        if paragraphs[index].text.strip() == "Required Closing Deliverables":
            end_index = index
            break
    for paragraph in paragraphs[start_index:end_index]:
        remove_paragraph(paragraph)
    return True


def ensure_spanish_section(doc):
    changed = remove_existing_spanish_section(doc)
    files = spanish_package_files()
    if not files:
        return changed

    required_heading = find_paragraph(doc, "Required Closing Deliverables")
    if required_heading is None:
        raise RuntimeError("Required Closing Deliverables heading not found")

    list_style = first_document_item_style(doc)
    required_heading.insert_paragraph_before(
        "Spanish / Bilingual Drafts",
        style=required_heading.style,
    )
    note = required_heading.insert_paragraph_before(
        "Draft convenience translations; English documents control unless separately approved."
    )
    for run in note.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)

    for path in files:
        paragraph = required_heading.insert_paragraph_before(path.name, style=list_style)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)
    return True


def update_doc(path):
    doc = Document(path)
    changed = False

    if not has_closing_date_field(doc):
        for paragraph in doc.paragraphs[:8]:
            text = paragraph.text.strip()
            if text.startswith("Prepared:"):
                table = add_prepared_closing_table(doc, text)
                paragraph_after(table._tbl, paragraph)
                remove_paragraph(paragraph)
                changed = True
                break
        if not changed:
            raise RuntimeError(f"Prepared date paragraph not found in {path}")

    changed = ensure_spanish_section(doc) or changed

    if changed:
        doc.save(path)
    return changed


def main():
    for path in DEFAULT_PATHS:
        if not path.exists():
            print(f"missing: {path}")
            continue
        changed = update_doc(path)
        status = "updated" if changed else "already-current"
        print(f"{status}: {path}")


if __name__ == "__main__":
    main()
