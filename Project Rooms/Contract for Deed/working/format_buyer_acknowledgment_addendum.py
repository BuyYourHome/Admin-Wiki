from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_cfd_drafts import get_docs_values, money, normalize_values


ROOT = Path(r"C:\Codex\Wiki Files\Project Rooms\Contract for Deed")
OUTPUT = ROOT / "output" / "320 Rose - Buyer Acknowledgment Addendum - DRAFT.docx"


ACKNOWLEDGMENTS = [
    "Buyer understands this is a Contract for Deed, not a deed transfer at signing or closing.",
    "Buyer understands Seller retains legal title until Buyer satisfies the contract terms and title is transferred later as provided in the contract.",
    "Buyer understands Buyer will have payment obligations under both the Contract for Deed and the Promissory Note.",
    "Buyer understands the monthly payment amount, due date, interest rate, loan term, and payment schedule.",
    "Buyer understands earnest money is paid at signing and credited toward the total down payment.",
    "Buyer understands the remaining down payment balance is due at closing.",
    "Buyer understands payments will be made by ACH draft unless otherwise agreed in writing.",
    "Buyer understands the property is being accepted in its stated condition, subject to the contract terms and any written addenda.",
    "Buyer understands Seller has disclosed the listed pending orders, liens, mortgages, or adverse conditions known to Seller.",
    "Buyer understands existing liens may remain during the contract term and that lienholders may have rights against the property if required payments are not made.",
    "Buyer understands the Contract or a Memorandum of Contract for Deed will be recorded with the county Register of Deeds.",
    "Buyer understands Buyer has a statutory right to cancel the Contract for Deed until midnight of the third business day after execution or delivery, whichever is later.",
    "Buyer understands Buyer has a right to cure default before forfeiture as provided by North Carolina law and the contract.",
    "Buyer understands Seller must provide periodic account statements during the contract term.",
    "Buyer understands Buyer may prepay the balance without penalty except as specifically allowed by the contract and applicable law.",
    "Buyer understands responsibilities for taxes, insurance, dues, repairs, and other property charges are governed by the contract.",
    "Buyer understands Seller is not giving Buyer legal advice.",
    "Buyer understands Buyer may hire Buyer's own attorney, at Buyer's expense, before signing and before closing.",
    "Buyer understands Seller's attorney may require changes, and the parties may sign the contract again at closing or execute an amendment.",
    "Buyer acknowledges that Buyer has read the Contract for Deed, Promissory Note, Memorandum, and all addenda before signing.",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, center=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    paragraph.paragraph_format.space_after = Pt(0)
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_document_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.05

    for style_name, size, color in [
        ("Heading 1", 15, "2E74B5"),
        ("Heading 2", 12, "2E74B5"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(3)


def add_label_value(paragraph, label, value):
    paragraph.add_run(label).bold = True
    paragraph.add_run(str(value))


def add_acknowledgment_table(doc, buyer1, buyer2):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    widths = [520, 1150, 1150, 7080]
    headers = ["No.", f"{buyer1}\nInitials", f"{buyer2}\nInitials", "Acknowledgment"]
    for idx, (cell, header, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
        set_cell_width(cell, width)
        set_cell_text(cell, header, bold=True, center=(idx < 3))
        set_cell_shading(cell, "F2F4F7")

    for number, text in enumerate(ACKNOWLEDGMENTS, start=1):
        cells = table.add_row().cells
        values = [str(number), "________", "________", text]
        for idx, (cell, value, width) in enumerate(zip(cells, values, widths)):
            set_cell_width(cell, width)
            set_cell_text(cell, value, center=(idx < 3))


def add_notary_block(doc, county, buyer_names):
    doc.add_heading("Notary Acknowledgment", level=2)
    for line in [
        "STATE OF: NORTH CAROLINA",
        f"COUNTY OF: {str(county).upper()}",
        (
            "I certify that the following person(s) personally appeared before me this day, "
            "each acknowledging to me that he/she/they signed the foregoing document: "
            f"{buyer_names}."
        ),
        "Date: ____________________",
        "Official Signature of Notary: ________________________________________",
        "Notary's printed or typed name: ______________________________, Notary Public",
        "My commission expires: ______________________",
    ]:
        doc.add_paragraph(line)


def main(x=None):
    if x is None:
        x = normalize_values(get_docs_values())
    buyer1 = x.get("buyer1") or "Purchaser 1"
    buyer2 = x.get("buyer2") or "Purchaser 2"
    buyer_names = " and ".join(name for name in [buyer1, buyer2] if name)

    doc = Document()
    set_document_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Buyer Acknowledgment Addendum")
    run.bold = True
    run.font.size = Pt(17)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Contract for Deed - {x['property_address']}, {x['property_city_state']}").bold = True

    intro = doc.add_paragraph()
    intro.add_run("Purpose: ").bold = True
    intro.add_run(
        "This Addendum records Purchaser's acknowledgments concerning the nature of the Contract for Deed transaction. "
        "Each Purchaser should initial each acknowledgment and sign below."
    )

    details = doc.add_paragraph()
    add_label_value(details, "Purchaser: ", buyer_names)
    add_label_value(doc.add_paragraph(), "Property: ", f"{x['property_address']}, {x['property_city_state']}")
    add_label_value(doc.add_paragraph(), "Seller: ", f"{x['trust']}, by and through {x['trustee']}, Trustee")
    add_label_value(doc.add_paragraph(), "Purchase price: ", money(x["sale_price"]))
    add_label_value(doc.add_paragraph(), "Amount financed: ", money(x["loan_amount"]))

    doc.add_heading("Buyer Initialed Acknowledgments", level=2)
    add_acknowledgment_table(doc, buyer1, buyer2)

    doc.add_heading("Purchaser Signatures", level=2)
    doc.add_paragraph(f"{buyer1}: ________________________________________ Date: ____________________")
    doc.add_paragraph(f"{buyer2}: ________________________________________ Date: ____________________")

    add_notary_block(doc, x["county"], buyer_names)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
