from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_cfd_drafts import get_docs_values, money, normalize_values


ROOT = Path(r"C:\Codex\Wiki Files\Project Rooms\Contract for Deed")
OUTPUT = ROOT / "output" / "320 Rose - Term Sheet - DRAFT.docx"


def format_date(value):
    return f"{value.strftime('%B')} {value.day}, {value.year}" if hasattr(value, "strftime") else str(value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    set_cell_width(table.rows[0].cells[0], 3100)
    set_cell_width(table.rows[0].cells[1], 6260)
    set_cell_text(table.rows[0].cells[0], "Term", bold=True)
    set_cell_text(table.rows[0].cells[1], "Amount / Detail", bold=True)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "F2F4F7")
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_width(cells[0], 3100)
        set_cell_width(cells[1], 6260)
        set_cell_text(cells[0], label, bold=True)
        set_cell_text(cells[1], value)
    doc.add_paragraph("")
    return table


def set_document_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def add_heading(doc, text):
    paragraph = doc.add_heading(text, level=2)
    return paragraph


def main():
    x = normalize_values(get_docs_values())
    doc = Document()
    set_document_styles(doc)

    title = doc.add_paragraph()
    title_run = title.add_run("Term Sheet")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Proposed Contract for Deed Terms for {x['property_address']}, {x['property_city_state']}").bold = True
    subtitle.paragraph_format.space_after = Pt(10)

    intro = doc.add_paragraph()
    intro.add_run("Source of values: ").bold = True
    intro.add_run("320 Rose project spreadsheet, Docs worksheet.")

    notice = doc.add_paragraph()
    notice.add_run("Status: ").bold = True
    notice.add_run(
        "This Term Sheet summarizes proposed business terms for review and discussion. "
        "It is not the final contract, does not transfer title, and is subject to final contract documents, required disclosures, buyer qualification review, and attorney review. "
        "If this Term Sheet conflicts with the final signed documents, the final signed documents control."
    )

    add_heading(doc, "Parties and Property")
    seller = f"{x['trust']}, by and through {x['trustee']}, Trustee"
    add_key_value_table(
        doc,
        [
            ("Seller", seller),
            ("Seller / trustee address", x["trustee_address"]),
            ("Purchaser", x["buyer"]),
            ("Purchaser address", x["buyer_address"]),
            ("Property", f"{x['property_address']}, {x['property_city_state']}"),
            ("County", str(x["county"]).title()),
            ("Parcel ID", x["parcel"]),
            ("Legal description summary", x["brief_legal"]),
        ],
    )

    add_heading(doc, "Purchase and Financing Terms")
    add_key_value_table(
        doc,
        [
            ("Proposed purchase price", money(x["sale_price"])),
            ("Earnest money due at signing", money(x["earnest_money"])),
            ("Total down payment", money(x["down_payment"])),
            ("Remaining down payment due at closing", money(x["remaining_down_payment"])),
            ("Amount financed", money(x["loan_amount"])),
            ("Interest rate", x["interest"]),
            ("Financing term", f"{x['term_years']} years; {x['term_months']} monthly installments"),
            ("Proposed first payment date", format_date(x["loan_start"])),
            ("Proposed final payment date", format_date(x["loan_end"])),
        ],
    )

    total_payment = x["monthly_pi"] + x["insurance"] + x["tax"]
    add_heading(doc, "Estimated Monthly Payment")
    add_key_value_table(
        doc,
        [
            ("Principal and interest", money(x["monthly_pi"])),
            ("Property insurance escrow", money(x["insurance"])),
            ("Property tax escrow", money(x["tax"])),
            ("Estimated total monthly payment", money(total_payment)),
        ],
    )

    add_heading(doc, "Known Pending Orders or Adverse Conditions")
    adverse_items = [line.strip() for line in str(x["adverse"]).splitlines() if line.strip()]
    if adverse_items:
        for item in adverse_items:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("None listed on the Docs worksheet.")

    add_heading(doc, "Next Steps")
    for item in [
        "Confirm buyer qualification and final business terms.",
        "Prepare final Contract for Deed, Memorandum of Contract for Deed, Promissory Note, required acknowledgments, and any required disclosures.",
        "Submit documents for attorney review before execution.",
        "Have all required documents signed, notarized where applicable, and recorded as required.",
    ]:
        doc.add_paragraph(item, style="List Number")

    add_heading(doc, "Acknowledgment of Receipt")
    doc.add_paragraph("Purchaser acknowledges receipt of this Term Sheet for review and discussion.")
    doc.add_paragraph("Purchaser 1: ________________________________________ Date: ____________________")
    doc.add_paragraph("Purchaser 2: ________________________________________ Date: ____________________")
    doc.add_paragraph("Seller / Trustee: ___________________________________ Date: ____________________")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
