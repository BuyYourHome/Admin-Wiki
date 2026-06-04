from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:\Codex\Wiki Files\Project Rooms\Contract for Deed")
OUTPUT = ROOT / "output"
REVIEW_MARKERS = ("ATTORNEY REVIEW NOTE:", "MANAGEMENT REVIEW NOTE:")


DOCS = [
    (
        "320 Rose - Contract for Deed Agreement - DRAFT.docx",
        "320 Rose - Contract for Deed Agreement - ATTORNEY REVIEW PACKAGE.docx",
        [
            (
                "Date of Agreement:",
                "ATTORNEY REVIEW NOTE: Please confirm this Contract for Deed complies with North Carolina Chapter 47H, including required contents, signed-and-acknowledged party requirements, delivery of an exact copy to the purchaser, recordation timing, cancellation rights, default/cure rights, periodic-statement obligations, and any additional statutory notices or disclosures that should be included before the document is sent to the seller or buyer.",
            ),
            (
                "Installment Payments",
                "MANAGEMENT REVIEW NOTE: Confirm these business terms match the intended deal before sending to counsel or the parties: purchase price, down payment, loan amount, interest rate, monthly ACH draft/payment amount, first payment date, final payment date, number of payments, escrow amounts, and annual statement process. Do not ask the attorney to compare these values against the spreadsheet because the spreadsheet is not part of the attorney review package.",
            ),
            (
                "PENDING ORDERS OR ADVERSE CONDITIONS:",
                "ATTORNEY REVIEW NOTE: Please confirm whether any liens, subject-to debt, title matters, pending orders, adverse conditions, unpaid taxes, insurance obligations, HOA/special-assessment items, or recording/closing instructions should be disclosed or revised before signing. If the property is encumbered by an existing lien or mortgage, please confirm whether the separate 14-point bold capital-letter disclosure required by N.C. Gen. Stat. 47H-6 is needed.",
            ),
            (
                "RIGHT TO CURE DEFAULT:",
                "ATTORNEY REVIEW NOTE: Please review the default, right-to-cure, forfeiture, eviction, acceleration, and remedy language as a single package across the Contract for Deed and Promissory Note. Confirm whether the cure periods, notice method, and stated remedy are consistent with Chapter 47H and with the intended enforcement process.",
            ),
            (
                "Additional Terms, Conditions or Addenda",
                "ATTORNEY REVIEW NOTE: Please review the attorney-change contingency language, including Seller's option to proceed under the previously signed Contract if material attorney-required changes cause Purchaser not to proceed, and confirm whether the due diligence/earnest money return language is clear and enforceable.",
            ),
            (
                "RIGHT TO CANCEL:",
                "ATTORNEY REVIEW NOTE: Please confirm the right-to-cancel notice is in the required location and formatting, including whether it must appear immediately above the purchaser signature in not less than 14-point boldface type.",
            ),
            (
                "Investment Services LLC, Trustee -",
                "ATTORNEY REVIEW NOTE: Please review the seller/trust/trustee language and signature block. Confirm the trust is properly identified as seller and that Investment Services LLC, Trustee, through Wes Browning, Manager, has proper authority and wording to sign for the trust.",
            ),
        ],
    ),
    (
        "320 Rose - Memorandum of Contract for Deed - DRAFT.docx",
        "320 Rose - Memorandum of Contract for Deed - ATTORNEY REVIEW PACKAGE.docx",
        [
            (
                "2. REAL PROPERTY.",
                "ATTORNEY REVIEW NOTE: Please verify the legal description, county, parcel ID, property address, seller/trustee identity, and buyer names before recording or sending for signature.",
            ),
            (
                "3. TERM OF CONTRACT.",
                "ATTORNEY REVIEW NOTE: Please confirm the Memorandum includes the required applicable time periods for the Contract for Deed and that the term/payment-count information is sufficient for recording without disclosing unnecessary financial terms.",
            ),
            (
                "5. RIGHT TO CANCEL.",
                "ATTORNEY REVIEW NOTE: Please confirm this Memorandum is appropriate for recording in Wake County, North Carolina, and that it contains all information needed to memorialize the Contract for Deed without disclosing more than should be recorded.",
            ),
            (
                "SELLER:",
                "ATTORNEY REVIEW NOTE: Please review the seller/trust/trustee signature block and notary acknowledgment for recordability, including whether the signer capacity and trust authority are stated correctly.",
            ),
            (
                "PURCHASER:",
                "ATTORNEY REVIEW NOTE: Please review the buyer signature lines and notary block, including whether both purchasers must sign and whether the acknowledgment wording is sufficient for recording.",
            ),
        ],
    ),
    (
        "320 Rose - Promissory Note for Contract for Deed - DRAFT.docx",
        "320 Rose - Promissory Note for Contract for Deed - ATTORNEY REVIEW PACKAGE.docx",
        [
            (
                "FOR VALUE RECEIVED",
                "ATTORNEY REVIEW NOTE: Please confirm this Promissory Note is appropriate for the Contract for Deed structure, including whether the note should be separate from or incorporated into the Contract for Deed, and whether it conflicts with the Contract for Deed Agreement or Memorandum.",
            ),
            (
                "8.500% APR",
                "MANAGEMENT REVIEW NOTE: Confirm the principal amount, interest rate, monthly payment, first payment date, final payment date, and maturity/payment count match the intended business terms. Attorney review should focus on legal sufficiency, default remedies, attorney-fee language, and satisfaction language.",
            ),
            (
                "In the event of default",
                "ATTORNEY REVIEW NOTE: Please confirm the note default, cure period, acceleration, and notice provisions are consistent with the Contract for Deed and North Carolina Chapter 47H, including any longer cure period or notice method required before forfeiture of purchaser rights.",
            ),
            (
                "Upon default,",
                "ATTORNEY REVIEW NOTE: Please confirm the attorney-fee clause, default remedies, prepayment treatment, late-charge treatment, satisfaction language, and notary/signature requirements are enforceable and appropriate for this Contract for Deed transaction.",
            ),
        ],
    ),
]


def insert_paragraph_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    new_paragraph.text = text
    return new_paragraph


def style_review_note(paragraph):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.bold = True


def remove_existing_review_notes(doc):
    for paragraph in list(doc.paragraphs):
        if any(marker in paragraph.text for marker in REVIEW_MARKERS):
            paragraph._element.getparent().remove(paragraph._element)


def add_contextual_review_notes(doc, placements):
    remove_existing_review_notes(doc)
    for anchor, note in placements:
        target = None
        for paragraph in doc.paragraphs:
            if anchor in paragraph.text:
                target = paragraph
                break
        if target is None:
            target = doc.paragraphs[0]
        note_paragraph = insert_paragraph_after(target, note)
        style_review_note(note_paragraph)


def main():
    for source_name, output_name, placements in DOCS:
        source = OUTPUT / source_name
        target = OUTPUT / output_name
        doc = Document(str(source))
        add_contextual_review_notes(doc, placements)
        doc.save(target)
        print(target)


if __name__ == "__main__":
    main()
