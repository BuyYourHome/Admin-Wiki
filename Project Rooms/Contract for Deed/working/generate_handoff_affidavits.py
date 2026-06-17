from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(r"C:\Codex\Wiki Files\Project Rooms\Contract for Deed")
TRANSACTION = ROOT / "transactions" / "320 Rose Pl - Ever Cardoza"
OUT = TRANSACTION / "output" / "affidavits"
TEAMS = Path(
    r"C:\Users\wesbr\Buy Your Home\Buy Your Home - Property\28-SYH-320 Rose Pl"
    r"\Selling\Ever Cardoza\Contract Package\Affidavits"
)


PROPERTY = "320 Rose Pl, Wendell, NC 27591"
BUYERS = "Ever Amarildo Cardoza Bolanos and Maria Geraldina Sarmiento"
DATE = "June 9, 2026"
READINESS_NOTE = (
    "This draft support document is prepared for the Contract for Deed closing package. "
    "Final execution remains subject to attorney/compliance review, confirmation of signer "
    "authority and capacity, and notarized signatures."
)
NOTARY_NOTE = (
    "STATE OF: NORTH CAROLINA\n"
    "COUNTY OF: ____________________\n\n"
    "I certify that the following person(s) personally appeared before me this day, "
    "each acknowledging to me that he/she/they signed the foregoing document: "
    "____________________________________________________________.\n\n"
    "Date: ____________________\n\n"
    "Official Signature of Notary: ________________________________________\n\n"
    "Notary's printed or typed name: ______________________________, Notary Public\n\n"
    "My commission expires: ______________________"
)


def set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)


def add_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(13)


def add_labeled_para(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(text)


def add_numbered_facts(doc: Document, facts: list[str]) -> None:
    for fact in facts:
        doc.add_paragraph(fact, style="List Number")


def add_signature_block(doc: Document, signer_lines: list[str]) -> None:
    doc.add_paragraph()
    doc.add_paragraph("Affiant:")
    doc.add_paragraph("________________________________________")
    for line in signer_lines:
        doc.add_paragraph(line)
    doc.add_paragraph()
    doc.add_paragraph(NOTARY_NOTE)


def build_doc(
    filename: str,
    title: str,
    purpose: str,
    status: str,
    signer_capacity: str,
    authority_source: str,
    notary_need: str,
    facts: list[str],
    signer_lines: list[str],
) -> Path:
    doc = Document()
    set_normal_style(doc)
    add_title(doc, title)
    add_labeled_para(doc, "Property: ", PROPERTY)
    add_labeled_para(doc, "Buyer(s): ", BUYERS)
    add_labeled_para(doc, "Prepared: ", DATE)
    add_labeled_para(doc, "Purpose: ", purpose)
    add_labeled_para(doc, "Status: ", status)
    add_labeled_para(doc, "Signer / capacity from CWE handoff: ", signer_capacity)
    add_labeled_para(doc, "Authority / source from CWE handoff: ", authority_source)
    add_labeled_para(doc, "Notary need from CWE handoff: ", notary_need)
    doc.add_paragraph(READINESS_NOTE)
    doc.add_paragraph()
    doc.add_paragraph("Statement of Support", style=None).runs[0].bold = True
    add_numbered_facts(doc, facts)
    doc.add_paragraph(
        "The statements above are intended to support document preparation and closing-file "
        "review. They do not replace attorney review, final underwriting judgment, final funds "
        "verification, or final confirmation of signer authority."
    )
    add_signature_block(doc, signer_lines)
    path = OUT / filename
    doc.save(path)
    return path


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    TEAMS.mkdir(parents=True, exist_ok=True)

    docs = [
        (
            "320 Rose - Related-Company Rent Payment History Affidavit - DRAFT.docx",
            "Related-Company Rent Payment History Affidavit",
            "Supports tenant-buyer payment-history exception and documents that rent history comes from related company Buy Your Home, LLC.",
            "Required if relied on; closing deliverable.",
            "Buy Your Home, LLC manager or authorized representative, signing in company capacity.",
            "Rent verification letter and related-company rent-history affidavit draft.",
            "Recommended if used as underwriting or closing support.",
            [
                "Buy Your Home, LLC is identified in the current credit-worthiness handoff as the related-company source for rent/payment-history support.",
                "The handoff states that related-company rent verification supports the tenant-buyer payment-history exception and should be retained in the closing file if relied on.",
                "The related-company nature of the rent history should be disclosed and should not be described as independent third-party landlord verification.",
                "Final reliance on this support remains subject to attorney/compliance review and confirmation of the signer's authority and company capacity.",
            ],
            [
                "Name: ________________________________________",
                "Title: Manager or Authorized Representative",
                "For: Buy Your Home, LLC",
            ],
        ),
        (
            "320 Rose - Cash Reserves and Receivables Observation Affidavit - DRAFT.docx",
            "Cash Reserves and Receivables Observation Affidavit",
            "Supports funds-to-close and reserve condition through observed cash, receivables, and stated foreign-fund access.",
            "Required or substitute acceptable proof of funds; closing deliverable.",
            "Observing person/entity representative with exact capacity; source drafts identify Buy Your Home, LLC manager for observation facts.",
            "Cash-reserves affidavit materials and due receivable PDFs.",
            "Needed if used as formal proof.",
            [
                "The current handoff states that cash-reserves affidavit materials report $24,000 cash counted.",
                "The current handoff states that affidavit materials reference $14,000 receivables and $10,000 foreign funds in addition to counted cash.",
                "These items are support for funds-to-close and reserves only if accepted by the seller, counsel, and closing process.",
                "Receivables and foreign funds should be treated as conditional until collection, transfer, or closing-approved proof is obtained.",
                "Final funds to close, including earnest money, remaining down payment, and closing costs, must still be verified before closing execution.",
            ],
            [
                "Name: ________________________________________",
                "Title / capacity: ____________________________",
                "Entity, if any: ______________________________",
            ],
        ),
        (
            "320 Rose - Receipt Package Review and Acceptance Affidavit - DRAFT.docx",
            "Receipt Package Review and Acceptance Affidavit",
            "Documents seller/business acceptance of customer receipts as support for gross cash receipts and non-bank cash receipts.",
            "Required if receipts are relied on instead of stronger income proof; closing deliverable.",
            "Sell Your Home, LLC by authorized representative; source draft uses Investment Services LLC, Manager, signed by Al Bennett.",
            "Paid receipt package; Secretary of State filing; receipt-review affidavit draft.",
            "Recommended.",
            [
                "The current handoff states that the receipt package may be used as support for gross receipts and non-bank cash receipts.",
                "Receipt materials should be treated as gross receipt support only and should not be stated as proof of net borrower income.",
                "The handoff identifies the need to confirm the authority and signature block before inclusion in the final closing package.",
                "Final reliance on receipt-package support remains subject to seller business judgment, attorney/compliance review, and confirmation of signer authority.",
            ],
            [
                "Al Bennett, Manager",
                "Investment Services LLC, Manager",
                "For / on behalf of: Sell Your Home, LLC, as authorized representative to be confirmed",
            ],
        ),
        (
            "320 Rose - Business Judgment Approval Direction Affidavit - DRAFT.docx",
            "Business Judgment Approval Direction Affidavit",
            "Documents the business decision to proceed despite nonstandard self-employment proof, high debt ratio, and related-company rent verification.",
            "Recommended / required if seller wants business-side acceptance documented; closing deliverable.",
            "Sell Your Home, LLC by authorized representative; source draft uses Investment Services LLC, Manager, signed by Al Bennett.",
            "Management approval decision; evaluation source evidence; attorney/compliance review.",
            "Recommended.",
            [
                "The current handoff states that the buyer file is supportable for document preparation under a business-judgment and closing-deliverables approach.",
                "The current handoff states that closing execution is not ready until required affidavits, funds proof, final legal-name spelling confirmation, and attorney/compliance review are complete.",
                "The business decision to proceed should account for nonstandard self-employment proof, credit-report debt load, related-company rent verification, and missing or conditional funds proof.",
                "This document should not be treated as final closing approval or legal compliance approval.",
                "Final inclusion remains subject to authority confirmation and attorney/compliance review.",
            ],
            [
                "Al Bennett, Manager",
                "Investment Services LLC, Manager",
                "For / on behalf of: Sell Your Home, LLC, as authorized representative to be confirmed",
            ],
        ),
    ]

    for item in docs:
        generated = build_doc(*item)
        target = TEAMS / generated.name
        target.write_bytes(generated.read_bytes())
        print(generated)
        print(target)


if __name__ == "__main__":
    main()
