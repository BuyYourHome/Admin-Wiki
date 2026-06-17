from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

from build_cfd_drafts import get_docs_values, normalize_values, short_date


ROOT = Path(r"C:\Codex\Wiki Files\Project Rooms\Contract for Deed")
PROTOTYPE = ROOT / "reference" / "Rose memorandum prototype" / "320 Rose - Memorandum of Contract for Deed - PROTOTYPE.docx"
OUTPUT = ROOT / "output" / "320 Rose - Memorandum of Contract for Deed - DRAFT.docx"


def set_paragraph(paragraph, text, blue=False):
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0]
    run.text = str(text)
    if blue:
        run.font.color.rgb = RGBColor(0, 0, 255)


def set_paragraph_bold(paragraph, text):
    set_paragraph(paragraph, text)
    for run in paragraph.runs:
        if run.text:
            run.bold = True


def set_mixed_runs(paragraph, segments):
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.text = ""
    first = paragraph.runs[0]
    for idx, (text, bold) in enumerate(segments):
        run = first if idx == 0 else paragraph.add_run()
        run.text = str(text)
        run.bold = bold


def set_section_paragraph(paragraph, title, body=""):
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.text = ""
    title_run = paragraph.runs[0]
    title_run.text = title
    title_run.bold = True
    if body:
        body_run = paragraph.add_run(body)
        body_run.bold = True


def set_label_value(paragraph, label, value):
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.text = ""
    label_run = paragraph.runs[0]
    label_run.text = label
    value_run = paragraph.add_run(str(value))
    value_run.bold = True


def split_address(address):
    parts = [part.strip() for part in str(address).split(",") if part.strip()]
    if len(parts) >= 3:
        return parts[0], f"{parts[1]}, {parts[2]}"
    if len(parts) == 2:
        return parts[0], parts[1]
    return str(address), ""


def long_date(value):
    return f"{value.strftime('%B')} {value.day}, {value.year}"


def paragraph_startswith(doc, startswith):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(startswith):
            return paragraph
    return None


def set_run_text(paragraph, run_index, text, underline=None, bold=None):
    while len(paragraph.runs) <= run_index:
        paragraph.add_run("")
    run = paragraph.runs[run_index]
    run.text = str(text)
    if underline is not None:
        run.underline = underline
    if bold is not None:
        run.bold = bold
    return run


def clear_runs(paragraph, start_index):
    for run in paragraph.runs[start_index:]:
        run.text = ""


def replace_runs(paragraph, segments):
    if not paragraph.runs:
        paragraph.add_run("")
    base = paragraph.runs[0]
    for run in paragraph.runs:
        run.text = ""
    for idx, segment in enumerate(segments):
        run = base if idx == 0 else paragraph.add_run()
        run.text = str(segment.get("text", ""))
        if "underline" in segment:
            run.underline = segment["underline"]
        if "bold" in segment:
            run.bold = segment["bold"]


def set_kyle_top_table(doc, x):
    if not doc.tables:
        return
    table = doc.tables[0]
    if len(table.rows) < 1 or len(table.rows[0].cells) < 2:
        return
    state_cell = table.rows[0].cells[0]
    title_cell = table.rows[0].cells[1]
    state_cell.text = f"STATE OF NORTH CAROLINA\nCOUNTY OF {x['county']}"
    title_cell.text = "MEMORANDUM OF A\nCONTRACT FOR DEED\n(N.C. Gen. Stat. \u00a7 47H-2(d))"
    for cell in (state_cell, title_cell):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def update_kyle_parties(doc, x):
    paragraph = paragraph_startswith(doc, "1. PARTIES.")
    if not paragraph:
        return
    trustee_capacity = f"{x['trustee']}, a Wyoming limited liability company, as Trustee of the"
    parts = [
        "\n",
        "1. PARTIES.",
        " ",
        "This ",
        "Contract for Deed is ",
        "by and between ",
        trustee_capacity,
        " ",
        x["trust"],
        " whose principal address is ",
        x["trustee_address"] or "[SELLER ADDRESS TO BE INSERTED]",
        " (\u201cSeller\u201d) and",
        " the following ",
        "Purchaser",
        "(s)",
        ":",
    ]
    for idx, text in enumerate(parts):
        set_run_text(paragraph, idx, text)
    set_run_text(paragraph, 6, trustee_capacity, underline=True)
    set_run_text(paragraph, 8, x["trust"], underline=True)
    set_run_text(paragraph, 10, x["trustee_address"] or "[SELLER ADDRESS TO BE INSERTED]", underline=True)
    clear_runs(paragraph, len(parts))


def update_kyle_value_lines(doc, x):
    replacements = [
        ("Full Legal Name(s):", 1, x["buyer"]),
        ("Address:", 1, x["buyer_address"]),
        ("Street Address (if applicable):", 1, f"{x['property_address']}, {x['property_city_state']}".strip(", ")),
        ("Legal Description:", 1, x["brief_legal"] or x["legal"]),
        ("Due Date of the First Installment Payment:", 1, long_date(x["loan_start"])),
        ("Due Date of the Last Installment Payment:", 1, long_date(x["loan_end"])),
        ("Total Number of Installment Payments:", 2, x["term_months"]),
    ]
    for startswith, run_index, value in replacements:
        paragraph = paragraph_startswith(doc, startswith)
        if paragraph:
            set_run_text(paragraph, run_index, value, underline=True)

    county_line = paragraph_startswith(doc, "County:")
    if county_line:
        replace_runs(
            county_line,
            [
                {"text": "\tCounty: "},
                {"text": x["county"], "underline": True},
                {"text": "\t\t\t\t", "underline": True},
                {"text": " Parcel ID: "},
                {"text": x["parcel"], "underline": True},
                {"text": "\t\t\t\t", "underline": True},
            ],
        )


def update_return_to_line(doc, x):
    trustee_return_to = (
        f"{x['trustee']}, Trustee, {x['trustee_address']}"
        if x["trustee_address"]
        else f"{x['trustee']}, Trustee"
    )
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if "Prepared By:" in text and "Return To:" in text:
            set_mixed_runs(
                paragraph,
                [
                    ("Prepared By: ", False),
                    ("\t", False),
                    (x["prepared_by"], False),
                    ("\nReturn To:", False),
                    ("\t", False),
                    (trustee_return_to, False),
                ],
            )
            return


def update_kyle_seller_signature(doc, x):
    seller_seen = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("SELLER:"):
            seller_seen = True
            continue
        if not seller_seen:
            continue
        if text.startswith("STATE OF"):
            break
        if x["trust"] in text or "Real Estate Trust" in text:
            set_paragraph_bold(paragraph, x["trust"])
            continue
        if text.startswith("By:") and ("trustee" in text.lower() or x["trustee"] in text):
            set_run_text(paragraph, 0, "By:", bold=True)
            set_run_text(paragraph, 1, "\t", bold=True)
            set_run_text(paragraph, 2, f"{x['trustee']}, a Wyoming limited liability company", bold=True)
            set_run_text(paragraph, 3, ", trustee", bold=True)
            clear_runs(paragraph, 4)
            continue
        if "Printed Name:" in text and "Title:" in text and "Date:" in text:
            # Preserve Kyle's underlined signature-line layout while filling spreadsheet-sourced capacity.
            set_run_text(paragraph, 22, f"{x['manager']}\t", underline=True)
            set_run_text(paragraph, 23, "\t", underline=True)
            set_run_text(paragraph, 24, "", underline=True)
            set_run_text(paragraph, 25, "", underline=True)
            set_run_text(paragraph, 34, "Manager\t", underline=True)
            set_run_text(paragraph, 35, "\t", underline=True)
            set_run_text(paragraph, 36, "", underline=True)
            set_run_text(paragraph, 37, "", underline=True)
            set_run_text(paragraph, 38, "", underline=True)
            set_run_text(paragraph, 39, "", underline=True)
            return


def update_kyle_state_county(paragraph, x):
    if not paragraph:
        return
    set_run_text(paragraph, 0, "STATE OF ")
    set_run_text(paragraph, 1, "NORTH CAROLINA", underline=True)
    set_run_text(paragraph, 2, "", underline=True)
    set_run_text(paragraph, 3, "", underline=True)
    set_run_text(paragraph, 4, "", underline=True)
    set_run_text(paragraph, 5, "\nCOUNTY OF ")
    set_run_text(paragraph, 6, x["county"], underline=True)
    for idx in range(7, len(paragraph.runs)):
        paragraph.runs[idx].text = ""


def update_kyle_notaries(doc, x):
    state_blocks = []
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith("STATE OF"):
            state_blocks.append((idx, paragraph))

    if len(state_blocks) >= 1:
        update_kyle_state_county(state_blocks[0][1], x)
    if len(state_blocks) >= 2:
        update_kyle_state_county(state_blocks[1][1], x)

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if "certify that" in text and ("as Manager of Investment Services LLC" in text or "Trustee of the 320 Rose" in text):
            set_run_text(paragraph, 10, x["manager"], bold=True)
            set_run_text(paragraph, 11, "")
            set_run_text(paragraph, 12, f", as Manager of {x['trustee']}, Trustee of the {x['trust']}")
            continue
        if "personally appeared before me this day" in text and ("Ever" in text or "Maria" in text):
            set_run_text(paragraph, 10, x["buyer1"] or x["buyer"], bold=True)
            set_run_text(paragraph, 11, " and ", bold=True)
            set_run_text(paragraph, 12, x["buyer2"], bold=True)
            set_run_text(paragraph, 13, " ")
            set_run_text(paragraph, 14, "")
            continue


def update_kyle_purchaser_signature(doc, x):
    buyer_lines = []
    in_purchaser = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("PURCHASER"):
            in_purchaser = True
            continue
        if in_purchaser and text.startswith("STATE OF"):
            break
        if in_purchaser and text.startswith("Printed Name:"):
            buyer_lines.append(paragraph)

    buyers = [x.get("buyer1") or x["buyer"], x.get("buyer2") or ""]
    for paragraph, buyer in zip(buyer_lines, buyers):
        if not buyer:
            continue
        set_run_text(paragraph, 9, buyer, underline=True)
        for idx in range(10, min(len(paragraph.runs), 14)):
            paragraph.runs[idx].text = ""


def replace_legacy_buyer_name_variants(doc, x):
    buyer2 = (x.get("buyer2") or "").strip()
    if not buyer2:
        return
    legacy_buyer2_names = [
        "Maria Sarmjento",
        "Maria Sarmiento",
        "Maria Geraldine Sarmiento",
        "Maria Geraldina Sarmiento",
    ]
    replacements = {}
    for old in legacy_buyer2_names:
        if old != buyer2:
            replacements[old] = buyer2
            replacements[old.upper()] = buyer2.upper()

    def update_paragraph(paragraph):
        text = paragraph.text
        new_text = text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            set_paragraph(paragraph, new_text)

    for paragraph in doc.paragraphs:
        update_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    update_paragraph(paragraph)


def replace_or_update(doc, startswith, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(startswith):
            set_paragraph(paragraph, text)
            return True
    return False


def replace_or_update_bold(doc, startswith, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(startswith):
            set_paragraph_bold(paragraph, text)
            return True
    return False


def replace_section_or_update(doc, startswith, title, body=""):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(startswith):
            set_section_paragraph(paragraph, title, body)
            return True
    return False


def replace_label_or_update(doc, startswith, label, value):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(startswith):
            set_label_value(paragraph, label, value)
            return True
    return False


def remove_page_breaks_from_section(doc, startswith):
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip().startswith(startswith):
            continue
        for br in list(paragraph._p.xpath(".//w:br")):
            if br.get(qn("w:type")) == "page":
                br.getparent().remove(br)
        return


def fill_county_and_date(doc, x):
    county = x["county"]
    date_value = x["loan_start"]
    month = date_value.strftime("%B")
    year = str(date_value.year)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("STATE OF NORTH CAROLINA") and "COUNTY OF" in text:
            set_mixed_runs(
                paragraph,
                [
                    ("STATE OF ", False),
                    ("NORTH CAROLINA", True),
                    ("    COUNTY OF ", False),
                    (county, True),
                ],
            )
            break

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("this ____ day of"):
            set_mixed_runs(
                paragraph,
                [
                    ("this ____ day of ", False),
                    (month, True),
                    (", ", False),
                    (year, True),
                    (", to memorialize the Contract for Deed of even date hereto.", False),
                ],
            )
            return


def fill_notary_counties(doc, x):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("COUNTY OF") and "STATE OF NORTH CAROLINA" not in text:
            set_mixed_runs(paragraph, [("COUNTY OF: ", False), ("\t", False), (x["county"], True)])
        elif text.startswith("STATE OF:"):
            set_mixed_runs(paragraph, [("STATE OF", False), (":", False), (" ", False), ("\t", False), ("NORTH CAROLINA", True)])


def update_notary_acknowledgments(doc, x):
    buyer = x["buyer"]
    buyer2 = x.get("buyer2") or ""
    legacy_buyer_names = ["Ever Cardoza", "Maria Sarmjento", "Maria Sarmiento", "Maria Geraldine Sarmiento"]
    seller_segments = [
        ("I certify that the following person(s) personally appeared before me this day, each acknowledging to me that he/she/they signed the foregoing document: ", False),
        (x["manager"], True),
        (", Manager of ", False),
        (x["trustee"], True),
        (", Trustee of ", False),
        (x["trust"], True),
        (".", False),
    ]
    buyer_segments = [
        ("I certify that the following person(s) personally appeared before me this day, each acknowledging to me that he/she/they signed the foregoing document: ", False),
        (buyer, True),
        (".", False),
    ]
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if "personally appeared before me this day" not in text:
            continue
        if x["manager"] in text or x["trustee"] in text or x["trust"] in text:
            set_mixed_runs(paragraph, seller_segments)
        elif any(name and name in text for name in [buyer, buyer2, *legacy_buyer_names]):
            set_mixed_runs(paragraph, buyer_segments)


def notary_block_lines(county, signer_text):
    return [
        "STATE OF: NORTH CAROLINA",
        f"COUNTY OF: {str(county).upper()}",
        (
            "I certify that the following person(s) personally appeared before me this day, "
            "each acknowledging to me that he/she/they signed the foregoing document: "
            f"{signer_text}."
        ),
        "Date: ____________________",
        "Official Signature of Notary: ________________________________________",
        "Notary's printed or typed name: ______________________________, Notary Public",
        "My commission expires: ______________________",
    ]


def replace_paragraph_range_with_lines(doc, start_idx, end_idx, lines):
    anchor = doc.paragraphs[start_idx]
    inserted = []
    for text in lines:
        inserted.append(anchor.insert_paragraph_before(text))
    body = doc._body._element
    for paragraph in doc.paragraphs[start_idx + len(inserted) : end_idx + 1 + len(inserted)]:
        body.remove(paragraph._element)


def standardize_notary_blocks(doc, x):
    signer_order = [
        f"{x['manager']}, Manager of {x['trustee']}, Trustee of {x['trust']}",
        x["buyer"],
    ]
    starts = []
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip().upper()
        if not text.startswith("STATE OF"):
            continue
        window = "\n".join(p.text for p in doc.paragraphs[idx : min(len(doc.paragraphs), idx + 10)])
        if "personally appeared before me this day" in window:
            starts.append(idx)
    for block_number, start_idx in reversed(list(enumerate(starts[:2]))):
        end_idx = None
        for idx in range(start_idx, min(len(doc.paragraphs), start_idx + 12)):
            if doc.paragraphs[idx].text.strip().lower().startswith("my commission expires"):
                end_idx = idx
                break
        if end_idx is None:
            continue
        replace_paragraph_range_with_lines(
            doc,
            start_idx,
            end_idx,
            notary_block_lines(x["county"], signer_order[block_number]),
        )


def update_seller_signature_fields(doc, x):
    in_seller_block = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("SELLER:"):
            in_seller_block = True
            continue
        if in_seller_block and text.startswith("STATE OF"):
            break
        if not in_seller_block:
            continue
        if x["trust"] in text:
            set_paragraph_bold(paragraph, x["trust"])
        elif text.startswith("By:") and x["trustee"] in text:
            set_mixed_runs(paragraph, [("By: ", False), (x["trustee"], True), (", Trustee", True)])
        elif text.startswith("Printed Name:"):
            set_mixed_runs(paragraph, [("Printed Name: ", False), (f"{x['manager']}, Manager", True)])
        elif text.startswith("Title:") and x["trustee"] in text:
            set_mixed_runs(paragraph, [("Title: Manager of ", False), (x["trustee"], True), (", Trustee", True)])


def ensure_two_buyer_signature_lines(doc, x):
    buyers = [name.strip() for name in (x.get("buyer1"), x.get("buyer2")) if name and name.strip()]
    if not buyers:
        buyers = [name.strip() for name in str(x["buyer"]).split(" and ") if name.strip()]
    if len(buyers) == 1:
        buyers.append("______________________________________")

    replacement = [
        buyers[0],
        "By: ______________________________________ (SEAL)",
        f"Printed Name: {buyers[0]}",
        "Date: ______________________",
        buyers[1],
        "By: ______________________________________ (SEAL)",
        f"Printed Name: {buyers[1]}",
        "Date: ______________________",
    ]

    paragraphs = list(doc.paragraphs)
    for idx, paragraph in enumerate(paragraphs):
        if not paragraph.text.strip().startswith("PURCHASER:"):
            continue

        insert_after = paragraph._p
        remove = []
        for following in paragraphs[idx + 1 :]:
            text = following.text.strip()
            if text.startswith("STATE OF ") or text.startswith("COUNTY OF "):
                break
            remove.append(following)

        for old in remove:
            old._element.getparent().remove(old._element)

        for text in reversed(replacement):
            new_p = paragraph.insert_paragraph_before(text)
            if text in buyers:
                for run in new_p.runs:
                    run.bold = True
            elif text.startswith("Printed Name:"):
                name = text.replace("Printed Name:", "").strip()
                set_mixed_runs(new_p, [("Printed Name:", False), (" ", False), (name, True)])
            insert_after.addnext(new_p._p)
        return

    doc.add_paragraph("")
    doc.add_paragraph("PURCHASER:")
    for text in replacement:
        p = doc.add_paragraph(text)
        if text in buyers:
            for run in p.runs:
                run.bold = True
        elif text.startswith("Printed Name:"):
            name = text.replace("Printed Name:", "").strip()
            set_mixed_runs(p, [("Printed Name:", False), (" ", False), (name, True)])


def update_buyer_signature_fields(doc, x):
    buyers = [name.strip() for name in str(x["buyer"]).split(" and ") if name.strip()]
    if not buyers:
        return
    in_purchaser_block = False
    buyer_index = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("PURCHASER:"):
            in_purchaser_block = True
            continue
        if in_purchaser_block and text.startswith("STATE OF"):
            break
        if not in_purchaser_block:
            continue
        if text.startswith("Printed Name:") and buyer_index < len(buyers):
            set_mixed_runs(paragraph, [("Printed Name:", False), (" ", False), (buyers[buyer_index], True)])
            buyer_index += 1
            continue
        if text in buyers or text in ("Ever Cardoza", "Maria Sarmjento", "Maria Sarmiento", "Maria Geraldine Sarmiento"):
            match_index = min(buyer_index, len(buyers) - 1)
            set_paragraph_bold(paragraph, buyers[match_index])


def main(x=None):
    if x is None:
        x = normalize_values(get_docs_values())
    doc = Document(str(PROTOTYPE))
    update_return_to_line(doc, x)
    set_kyle_top_table(doc, x)
    update_kyle_parties(doc, x)
    update_kyle_value_lines(doc, x)
    update_kyle_seller_signature(doc, x)
    update_kyle_purchaser_signature(doc, x)
    update_kyle_notaries(doc, x)
    replace_legacy_buyer_name_variants(doc, x)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
