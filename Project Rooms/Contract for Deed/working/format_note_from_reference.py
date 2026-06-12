from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

from build_cfd_drafts import get_docs_values, normalize_values, money, short_date


ROOT = Path(r"C:\Codex\Wiki Files\Project Rooms\Contract for Deed")
REFERENCE = ROOT / "reference" / "Cool Springs selling docs" / "_Promissory Note for Contract for Deed.docx"
PROTOTYPE = ROOT / "reference" / "Rose note prototype" / "320 Rose - Promissory Note for Contract for Deed - PROTOTYPE.docx"
OUTPUT = ROOT / "output" / "320 Rose - Promissory Note for Contract for Deed - DRAFT.docx"


def set_paragraph(paragraph, text, red=False):
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0]
    run.text = str(text)
    if red:
        run.font.color.rgb = RGBColor(192, 0, 0)


def set_mixed_runs(paragraph, segments):
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.text = ""
    first = paragraph.runs[0]
    for idx, segment in enumerate(segments):
        if len(segment) == 2:
            text, bold = segment
            underline = None
        else:
            text, bold, underline = segment
        run = first if idx == 0 else paragraph.add_run()
        run.text = str(text)
        run.bold = bold
        if underline is not None:
            run.underline = underline


def bold_matching_runs(paragraph, names):
    text = paragraph.text
    for name in names:
        if name and name in text:
            before, rest = text.split(name, 1)
            after = rest
            set_mixed_runs(paragraph, [(before, None), (name, True), (after, None)])
            text = paragraph.text


def replace_legacy_buyer_name_variants(doc, x):
    buyer2 = (x.get("buyer2") or "").strip()
    if not buyer2:
        return
    legacy_buyer2_names = [
        "Maria Sarmjento",
        "Maria Sarmiento",
        "Maria Geraldine Sarmiento",
        "Maria Geraldina Sarmiento",
    ]
    replacements = {}
    for old in legacy_buyer2_names:
        if old != buyer2:
            replacements[old] = buyer2
            replacements[old.upper()] = buyer2.upper()

    def update_paragraph(paragraph):
        text = paragraph.text
        new_text = text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            set_paragraph(paragraph, new_text)

    for paragraph in doc.paragraphs:
        update_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    update_paragraph(paragraph)


ONES = [
    "Zero",
    "One",
    "Two",
    "Three",
    "Four",
    "Five",
    "Six",
    "Seven",
    "Eight",
    "Nine",
    "Ten",
    "Eleven",
    "Twelve",
    "Thirteen",
    "Fourteen",
    "Fifteen",
    "Sixteen",
    "Seventeen",
    "Eighteen",
    "Nineteen",
]
TENS = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]


def number_under_1000_to_words(number):
    words = []
    hundreds, remainder = divmod(number, 100)
    if hundreds:
        words.extend([ONES[hundreds], "Hundred"])
    if remainder:
        if remainder < 20:
            words.append(ONES[remainder])
        else:
            tens, ones = divmod(remainder, 10)
            words.append(TENS[tens] if not ones else f"{TENS[tens]}-{ONES[ones]}")
    return " ".join(words)


def number_to_words(number):
    number = int(number)
    if number == 0:
        return ONES[0]
    groups = [
        (1_000_000_000, "Billion"),
        (1_000_000, "Million"),
        (1_000, "Thousand"),
        (1, ""),
    ]
    words = []
    remainder = number
    for value, label in groups:
        count, remainder = divmod(remainder, value)
        if count:
            words.append(number_under_1000_to_words(count))
            if label:
                words.append(label)
    return " ".join(words)


def money_words(value):
    cents_total = int(round(float(value) * 100))
    dollars, cents = divmod(cents_total, 100)
    return f"{number_to_words(dollars)} and {cents:02d}/100 Dollars ({money(value)})"


def set_cell_text_preserve(cell, text):
    lines = str(text).split("\n")
    while len(cell.paragraphs) < len(lines):
        cell.add_paragraph()
    for idx, line in enumerate(lines):
        set_paragraph(cell.paragraphs[idx], line)
    for idx in range(len(lines), len(cell.paragraphs)):
        set_paragraph(cell.paragraphs[idx], "")


def set_cell_paragraph_runs(cell, paragraph_index, segments, alignment=None):
    while len(cell.paragraphs) <= paragraph_index:
        cell.add_paragraph()
    paragraph = cell.paragraphs[paragraph_index]
    set_mixed_runs(paragraph, segments)
    if alignment is not None:
        paragraph.alignment = alignment


def set_note_header_table(table, x, principal):
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.text = ""
    right.text = ""
    set_cell_paragraph_runs(
        left,
        0,
        [("STATE OF:     ", None, False), ("NORTH CAROLINA", True, True)],
        WD_ALIGN_PARAGRAPH.LEFT,
    )
    set_cell_paragraph_runs(
        left,
        1,
        [("COUNTY OF: ", None, False), (x["county"], True, True)],
        WD_ALIGN_PARAGRAPH.LEFT,
    )
    set_cell_paragraph_runs(
        right,
        0,
        [("Date of Note: ", None, False), (short_date(x["loan_start"]), True, True)],
        WD_ALIGN_PARAGRAPH.RIGHT,
    )
    set_cell_paragraph_runs(
        right,
        1,
        [("Amount: ", None, False), (principal, True, True)],
        WD_ALIGN_PARAGRAPH.RIGHT,
    )


def split_address(address):
    parts = [part.strip() for part in str(address).split(",") if part.strip()]
    if len(parts) >= 3:
        return parts[0], f"{parts[1]}, {parts[2]}"
    if len(parts) == 2:
        return parts[0], parts[1]
    return str(address), ""


def ensure_notary_block(doc, x):
    buyer1 = x.get("buyer1") or "Ever Cardoza"
    buyer2 = x.get("buyer2") or "Purchaser 2"
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    if "Notary Public" in full_text and "personally appeared before me" in full_text:
        return
    doc.add_paragraph("")
    doc.add_paragraph("STATE OF: NORTH CAROLINA")
    doc.add_paragraph(f"COUNTY OF: {str(x['county']).upper()}")
    doc.add_paragraph("")
    doc.add_paragraph(
        "I certify that the following person(s) personally appeared before me this day, "
        "each acknowledging to me that he/she/they signed the foregoing document: "
        f"{x['buyer']}."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Date: ____________________")
    doc.add_paragraph("")
    doc.add_paragraph("Official Signature of Notary: ________________________________________")
    doc.add_paragraph("Notary's printed or typed name: ______________________________, Notary Public")
    doc.add_paragraph("My commission expires: ______________________")
    doc.add_paragraph("")
    doc.add_paragraph("Maker Signature: ______________________________")
    doc.add_paragraph(f"Printed Name: {buyer1}")
    doc.add_paragraph("")
    doc.add_paragraph("Maker Signature: ______________________________")
    doc.add_paragraph(f"Printed Name: {buyer2}")


def update_signature_and_notary_names(doc, x):
    buyer1 = x.get("buyer1") or "Ever Cardoza"
    buyer2 = x.get("buyer2") or "Purchaser 2"
    buyers = [name for name in [buyer1, buyer2] if name]
    combined = " and ".join(buyers)
    signature_index = 0
    printed_index = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Printed Name:"):
            name = buyers[min(printed_index, len(buyers) - 1)]
            set_mixed_runs(paragraph, [("Printed Name: ", None), (name, True)])
            printed_index += 1
            continue
        if text in {buyer1.upper(), buyer2.upper(), combined.upper(), x["buyer"].upper()}:
            name = buyers[min(signature_index, len(buyers) - 1)]
            set_paragraph(paragraph, name.upper())
            for run in paragraph.runs:
                run.bold = True
            signature_index += 1
            continue
        if "personally appeared before me" in text and (
            "Cardoza" in text
            or "Sarm" in text
            or (x.get("buyer") and x["buyer"] in text)
            or (x.get("buyer2") and x["buyer2"] in text)
        ):
            set_mixed_runs(
                paragraph,
                [
                    ("I certify that the following person(s) personally appeared before me this day, each acknowledging to me that he/she/they signed the foregoing document: ", None),
                    (combined, True),
                    (".", None),
                ],
            )
            continue
        bold_matching_runs(paragraph, buyers)


def notary_block_lines(county, signer_text):
    return [
        "STATE OF: NORTH CAROLINA",
        f"COUNTY OF: {str(county).upper()}",
        (
            "I certify that the following person(s) personally appeared before me this day, "
            "each acknowledging to me that he/she/they signed the foregoing document: "
            f"{signer_text}."
        ),
        "Date: ____________________",
        "Official Signature of Notary: ________________________________________",
        "Notary's printed or typed name: ______________________________, Notary Public",
        "My commission expires: ______________________",
    ]


def replace_paragraph_range_with_lines(doc, start_idx, end_idx, lines):
    anchor = doc.paragraphs[start_idx]
    inserted = []
    for text in lines:
        inserted.append(anchor.insert_paragraph_before(text))
    body = doc._body._element
    for paragraph in doc.paragraphs[start_idx + len(inserted) : end_idx + 1 + len(inserted)]:
        body.remove(paragraph._element)


def standardize_notary_blocks(doc, x):
    buyer1 = x.get("buyer1") or "Ever Cardoza"
    buyer2 = x.get("buyer2") or "Purchaser 2"
    buyers = [name for name in [buyer1, buyer2] if name]
    combined = " and ".join(buyers)
    for start_idx in reversed(
        [
            idx
            for idx, paragraph in enumerate(doc.paragraphs)
            if paragraph.text.strip().upper().startswith("STATE OF")
        ]
    ):
        end_idx = None
        for idx in range(start_idx, min(len(doc.paragraphs), start_idx + 12)):
            if doc.paragraphs[idx].text.strip().lower().startswith("my commission expires"):
                end_idx = idx
                break
        if end_idx is None:
            continue
        replace_paragraph_range_with_lines(
            doc,
            start_idx,
            end_idx,
            notary_block_lines(x["county"], combined),
        )


def ensure_page_break_before_in_testimony(doc):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("IN TESTIMONY"):
            if 'w:type="page"' in paragraph._p.xml:
                return
            first_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            first_run.add_break(WD_BREAK.PAGE)
            return


def main(x=None):
    if x is None:
        x = normalize_values(get_docs_values())
    buyer_line1, buyer_line2 = split_address(x["buyer_address"])
    holder = f"{x['trust']}, by and through {x['trustee']}, Trustee"
    holder_address = x["trustee_address"] or "[HOLDER ADDRESS TO BE INSERTED]"
    principal = money(x["loan_amount"])
    principal_words = money_words(x["loan_amount"])

    source = PROTOTYPE if PROTOTYPE.exists() else REFERENCE
    doc = Document(str(source))

    # Keep the prototype/reference format and replace only content.
    if len(doc.tables) >= 3:
        set_note_header_table(doc.tables[1], x, principal)
        set_cell_text_preserve(doc.tables[2].cell(0, 1), holder)
        set_cell_text_preserve(doc.tables[2].cell(1, 1), holder_address)
        set_cell_text_preserve(doc.tables[2].cell(2, 1), principal_words)

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text.startswith("7.500% APR") or "APR with a Security Interest" in text:
            set_paragraph(
                paragraph,
                f"{x['interest']} APR with a Security Interest being the Contract for Deed and related security documents for {x['property_address']}, {x['property_city_state']}. Payments begin {short_date(x['loan_start'])} and continue through {short_date(x['loan_end'])}. Monthly principal and interest payment is {money(x['monthly_pi'])}.",
            )
        elif text.startswith("This debt instrument is made in connection"):
            set_paragraph(
                paragraph,
                f"This debt instrument is made in connection with the Contract for Deed for the property located at {x['property_address']}, {x['property_city_state']}. The Holder of this Note is the Seller under the Contract for Deed, acting by and through the Trustee identified above.",
            )
        elif text.strip() == text.strip().upper() and ("CARDOZA" in text or "SARM" in text):
            pass
        elif text.strip().startswith("4121 ") or text.strip().endswith("Tensity Dr"):
            set_paragraph(paragraph, buyer_line1)
        elif "Raleigh, NC" in text:
            set_paragraph(paragraph, buyer_line2)

    ensure_notary_block(doc, x)
    update_signature_and_notary_names(doc, x)
    standardize_notary_blocks(doc, x)
    replace_legacy_buyer_name_variants(doc, x)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
