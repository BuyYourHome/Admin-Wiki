from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import RGBColor
from docx.table import Table

from build_rose_drafts import get_docs_values, normalize_values, money, short_date


ROOT = Path(r"C:\Codex\Wiki Files\Project Rooms\Contract for Deed")
PROTOTYPE = ROOT / "reference" / "Rose contract prototype" / "320 Rose - Contract for Deed Agreement - PROTOTYPE.docx"
REFERENCE = ROOT / "reference" / "Cool Springs selling docs" / "25-02-21 Seller Docs.docx"
OUTPUT = ROOT / "output" / "320 Rose - Contract for Deed Agreement - DRAFT.docx"
ATTORNEY_CONTINGENCY_TEXT = (
    "This Contract is signed subject to any changes Seller's attorney may make or require. "
    "The parties shall sign the Contract again at closing, or execute any required amendment, "
    "with any such attorney-required alterations incorporated. If such changes are material "
    "and Purchaser does not agree to proceed, Seller may elect to proceed under the previously "
    "signed Contract. If Seller does not so elect, Purchaser's due diligence funds shall be "
    "returned to Purchaser."
)
ADVERSE_CONDITION_RESPONSIBILITY_TEXT = (
    "Seller shall remain responsible for the pending orders or adverse conditions listed below. "
    "These matters may remain in place during the term of this Contract. To the extent any such "
    "matter must be released, satisfied, or otherwise resolved to convey marketable title, Seller "
    "shall cause it to be released, satisfied, or otherwise resolved before or at the closing of "
    "Purchaser's future sale of the Property, unless resolved earlier or otherwise agreed in writing."
)
INDEPENDENT_COUNSEL_TEXT = (
    "Purchaser acknowledges that Seller is not providing legal advice to Purchaser. "
    "Purchaser has been advised that Purchaser may, at Purchaser's own expense, consult with "
    "an attorney of Purchaser's choice before signing this Contract and before closing."
)


def split_address(address):
    parts = [part.strip() for part in str(address).split(",") if part.strip()]
    if len(parts) >= 2:
        return parts[0], ", ".join(parts[1:])
    return str(address), ""


def set_paragraph(paragraph, text, red=False):
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0]
    run.text = str(text)
    if red:
        run.font.color.rgb = RGBColor(192, 0, 0)


def set_paragraph_runs(paragraph, parts):
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.text = ""
    first = True
    for part in parts:
        text = str(part.get("text", ""))
        run = paragraph.runs[0] if first else paragraph.add_run()
        first = False
        run.text = text
        if "bold" in part:
            run.bold = part["bold"]
        if "underline" in part:
            run.underline = part["underline"]
        if part.get("red"):
            run.font.color.rgb = RGBColor(192, 0, 0)


def set_notary_acknowledgment(paragraph, signer_text, red_signer=False):
    set_paragraph_runs(
        paragraph,
        [
            {"text": "I certify that the following person(s) personally appeared before me this day, each acknowledging to me that he/she/they signed the foregoing document: "},
            {"text": signer_text, "red": red_signer},
            {"text": "."},
        ],
    )


def set_notary_certification_prefix(paragraph):
    set_paragraph_runs(
        paragraph,
        [
            {"text": "I certify that the following person(s) personally appeared before me this day, each acknowledging to me that he/she/they signed the foregoing document: "},
        ],
    )


def set_notary_signer_continuation(paragraph, signer_text, red_signer=False):
    set_paragraph_runs(
        paragraph,
        [
            {"text": signer_text, "red": red_signer},
            {"text": "."},
        ],
    )


def notary_block_lines(county, signer_text):
    return [
        "STATE OF: NORTH CAROLINA",
        f"COUNTY OF: {str(county).title()}",
        (
            "I certify that the following person(s) personally appeared before me this day, "
            "each acknowledging to me that he/she/they signed the foregoing document: "
            f"{signer_text}."
        ),
        "Date: ____________________",
        "Official Signature of Notary: ________________________________________",
        "Notary's printed or typed name: ______________________________, Notary Public",
        "My commission expires: ______________________",
    ]


def replace_paragraph_range_with_lines(doc, start_idx, end_idx, lines):
    anchor = doc.paragraphs[start_idx]
    inserted = []
    for text in lines:
        inserted.append(anchor.insert_paragraph_before(text))
    body = doc._body._element
    for paragraph in doc.paragraphs[start_idx + len(inserted) : end_idx + 1 + len(inserted)]:
        body.remove(paragraph._element)
    return inserted


def standardize_contract_notary_blocks(doc, x):
    signer_order = [
        x["buyer"],
        f"{x['manager']}, Manager of {x['trustee']}, Trustee of {x['trust']}",
    ]
    starts = []
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip().upper()
        if idx > 100 and (text.startswith("STATE:") or text.startswith("STATE OF")):
            starts.append(idx)
    for block_number, start_idx in reversed(list(enumerate(starts[:2]))):
        end_idx = None
        for idx in range(start_idx, min(len(doc.paragraphs), start_idx + 16)):
            text = doc.paragraphs[idx].text.strip()
            if text.startswith("Notary Public") or text.startswith("My commission expires") or text.startswith("Notary Public My Commission expires"):
                end_idx = idx
        if end_idx is None:
            continue
        replace_paragraph_range_with_lines(
            doc,
            start_idx,
            end_idx,
            notary_block_lines(x["county"], signer_order[block_number]),
        )


def insert_underlined_signature_line_before(paragraph):
    prev = paragraph._p.getprevious()
    if prev is not None:
        prev_text = "".join(node.text or "" for node in prev.iter() if node.tag.endswith("}t")).strip()
        if prev_text and set(prev_text) <= {"_"}:
            return
    signature_line = paragraph.insert_paragraph_before()
    set_paragraph_runs(signature_line, [{"text": "____________________________________________", "underline": False}])


def ensure_notary_signature_lines(doc):
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip().startswith("Notary Public") and "My Commission expires" in paragraph.text:
            insert_underlined_signature_line_before(paragraph)


def remove_after_contract(doc):
    body = doc._body._element
    children = list(body)
    start_remove = None
    for idx, child in enumerate(children):
        text = "".join(node.text or "" for node in child.iter() if node.tag.endswith("}t"))
        if text.strip().startswith("Prepared By:"):
            start_remove = idx
            break
    if start_remove is not None:
        for child in children[start_remove:]:
            body.remove(child)


def replace_text_in_paragraphs(doc, replacements):
    for paragraph in doc.paragraphs:
        for old, new in replacements.items():
            if old in paragraph.text:
                # Preserve paragraph formatting; direct replacement is safe when the phrase is isolated.
                text = paragraph.text.replace(old, new)
                set_paragraph(paragraph, text)


def find_paragraph(doc, startswith):
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith(startswith):
            return idx
    return None


def find_paragraph_containing(doc, needle):
    needle = needle.lower()
    for idx, paragraph in enumerate(doc.paragraphs):
        if needle in paragraph.text.strip().lower():
            return idx
    return None


def set_first_paragraph_starting(doc, startswith, text):
    idx = find_paragraph(doc, startswith)
    if idx is not None:
        set_paragraph(doc.paragraphs[idx], text)
    return idx


def set_county_of_paragraph(paragraph, county):
    original = paragraph.text
    prefix = original[: len(original) - len(original.lstrip())]
    set_paragraph(paragraph, f"{prefix}County of: {str(county).title()}")


def set_first_paragraph_starting_any(doc, startswith_options, text, red=False):
    for startswith in startswith_options:
        idx = find_paragraph(doc, startswith)
        if idx is not None:
            set_paragraph(doc.paragraphs[idx], text, red=red)
            return idx
    return None


def set_cell_text_preserve_formatting(cell, text, red=False):
    if not cell.paragraphs:
        cell.add_paragraph()
    set_paragraph(cell.paragraphs[0], text, red=red)
    for paragraph in cell.paragraphs[1:]:
        set_paragraph(paragraph, "")


def replace_legacy_buyer_name_variants(doc, x):
    buyer2 = (x.get("buyer2") or "").strip()
    if not buyer2:
        return
    legacy_buyer2_names = [
        "Maria Sarmjento",
        "Maria Sarmiento",
        "Maria Geraldine Sarmiento",
        "Maria Geraldina Sarmiento",
    ]
    replacements = {}
    for old in legacy_buyer2_names:
        if old != buyer2:
            replacements[old] = buyer2
            replacements[old.upper()] = buyer2.upper()

    def update_paragraph(paragraph):
        text = paragraph.text
        new_text = text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            set_paragraph(paragraph, new_text)

    for paragraph in doc.paragraphs:
        update_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    update_paragraph(paragraph)


def copy_paragraph_format(source, target):
    target.style = source.style
    target.alignment = source.alignment
    target.paragraph_format.left_indent = source.paragraph_format.left_indent
    target.paragraph_format.right_indent = source.paragraph_format.right_indent
    target.paragraph_format.first_line_indent = source.paragraph_format.first_line_indent
    target.paragraph_format.space_before = source.paragraph_format.space_before
    target.paragraph_format.space_after = source.paragraph_format.space_after
    target.paragraph_format.line_spacing = source.paragraph_format.line_spacing


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def insert_table_after_paragraph(doc, paragraph, rows, cols):
    table = doc.add_table(rows=rows, cols=cols)
    paragraph._p.addnext(table._tbl)
    return table


def cell_text(cell):
    return "\n".join(paragraph.text for paragraph in cell.paragraphs).strip()


def fill_cell_lines_preserve_formatting(cell, lines):
    if not cell.paragraphs:
        cell.add_paragraph()
    template = cell.paragraphs[0]
    while len(cell.paragraphs) < len(lines):
        paragraph = cell.add_paragraph()
        copy_paragraph_format(template, paragraph)
    for idx, line in enumerate(lines):
        set_paragraph(cell.paragraphs[idx], line)
    for paragraph in list(cell.paragraphs[len(lines) :]):
        remove_paragraph(paragraph)


def fill_value_table_preserve_formatting(table, lines):
    while len(table.rows) < len(lines):
        table.add_row()
    for idx, line in enumerate(lines):
        cell = table.cell(idx, 0)
        set_cell_text_preserve_formatting(cell, line)
    for row in table.rows[len(lines) :]:
        for cell in row.cells:
            set_cell_text_preserve_formatting(cell, "")


def table_after_heading_before(doc, heading_idx, end_idx):
    body = doc._body._element
    children = list(body)
    start_pos = children.index(doc.paragraphs[heading_idx]._p)
    end_pos = children.index(doc.paragraphs[end_idx]._p)
    for child in children[start_pos + 1 : end_pos]:
        if child.tag.endswith("}tbl"):
            return Table(child, doc)
    return None


def replace_installment_block_with_table(doc, x):
    heading_idx = find_paragraph_containing(doc, "Installment Payments")
    if heading_idx is None:
        return
    end_idx = find_paragraph_containing(doc, "ACH DRAFT PAYMENTS:")
    if end_idx is None:
        end_idx = find_paragraph_containing(doc, "AMORTIZATION SCHEDULE:")
    if end_idx is None or end_idx <= heading_idx:
        return

    value_rows = [
        str(x["term_months"]),
        short_date(x["loan_start"]),
        short_date(x["loan_end"]),
        money(x["monthly_pi"]).replace("$", ""),
        money(x["insurance"]).replace("$", ""),
        money(x["tax"]).replace("$", ""),
        money(x["total_payment"]).replace("$", ""),
    ]
    label_rows = [
        "For the first instalments:",
        "Due Date of the First Installment Payment:",
        "Due Date of the Last Installment Payment:",
        "Principal and Interest:",
        "Property Insurance *:",
        "Property Tax *:",
        "Total Monthly Payment:",
    ]

    existing_table = table_after_heading_before(doc, heading_idx, end_idx)
    if existing_table is not None:
        target_cell = None
        label_markers = ("Number of Installments", "For the first", "Due Date of the First", "Principal and Interest")
        for row in existing_table.rows:
            cells = row.cells
            for col_idx, cell in enumerate(cells):
                text = cell_text(cell)
                if any(marker in text for marker in label_markers):
                    if col_idx == 0 and len(cells) > 1:
                        target_cell = cells[1]
                    elif col_idx > 0:
                        target_cell = cells[col_idx - 1]
                    break
            if target_cell is not None:
                break
        if target_cell is None:
            for row in existing_table.rows:
                for cell in row.cells:
                    if not cell_text(cell):
                        target_cell = cell
                        break
                if target_cell is not None:
                    break
        if target_cell is None and len(existing_table.columns) > 1:
            target_cell = existing_table.cell(0, 1)
        if target_cell is not None:
            if target_cell.tables:
                fill_value_table_preserve_formatting(target_cell.tables[0], value_rows)
            else:
                fill_cell_lines_preserve_formatting(target_cell, value_rows)
            return

    template_paragraphs = doc.paragraphs[heading_idx + 1 : end_idx]
    while template_paragraphs and not template_paragraphs[-1].text.strip():
        template_paragraphs.pop()
    base_paragraph = template_paragraphs[0] if template_paragraphs else doc.paragraphs[heading_idx]

    body = doc._body._element
    children = list(body)
    start_pos = children.index(doc.paragraphs[heading_idx]._p)
    end_pos = children.index(doc.paragraphs[end_idx]._p)
    for child in children[start_pos + 1 : end_pos]:
        body.remove(child)

    outer = insert_table_after_paragraph(doc, doc.paragraphs[heading_idx], 1, 2)
    outer.alignment = WD_TABLE_ALIGNMENT.CENTER
    left_cell = outer.cell(0, 0)
    right_cell = outer.cell(0, 1)
    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    nested = left_cell.add_table(rows=7, cols=1)
    nested.alignment = WD_TABLE_ALIGNMENT.CENTER
    nested.style = "Table Grid"
    if left_cell.paragraphs and not left_cell.paragraphs[0].text.strip():
        remove_paragraph(left_cell.paragraphs[0])
    for idx, value in enumerate(value_rows):
        nested_cell = nested.cell(idx, 0)
        nested_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_text_preserve_formatting(nested_cell, value)
        nested_cell.paragraphs[0].alignment = base_paragraph.alignment

    set_paragraph(right_cell.paragraphs[0], label_rows[0])
    copy_paragraph_format(base_paragraph, right_cell.paragraphs[0])
    for label in label_rows[1:]:
        p = right_cell.add_paragraph()
        copy_paragraph_format(base_paragraph, p)
        set_paragraph(p, label)

    p = right_cell.add_paragraph()
    copy_paragraph_format(base_paragraph, p)
    set_paragraph(p, "(*) The property Insurance and Property Tax values are correct for the contract year only and will adjust annually.")


def ensure_attorney_contingency_clause(doc):
    existing_idx = find_paragraph_containing(doc, "Seller's attorney may make or require")
    if existing_idx is not None:
        set_paragraph(doc.paragraphs[existing_idx], ATTORNEY_CONTINGENCY_TEXT)
        return
    heading_idx = find_paragraph(doc, "Additional Terms")
    if heading_idx is None:
        return
    anchor_idx = None
    for idx in range(heading_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[idx].text.strip():
            anchor_idx = idx
            break
    if anchor_idx is None:
        return
    anchor = doc.paragraphs[anchor_idx]
    paragraph = anchor.insert_paragraph_before()
    copy_paragraph_format(anchor, paragraph)
    set_paragraph(paragraph, ATTORNEY_CONTINGENCY_TEXT)
    source_bold = next((run.bold for run in anchor.runs if run.text), None)
    if source_bold is not None:
        for run in paragraph.runs:
            run.bold = source_bold


def ensure_independent_counsel_clause(doc):
    existing_idx = find_paragraph_containing(doc, "Seller is not providing legal advice to Purchaser")
    if existing_idx is not None:
        set_paragraph(doc.paragraphs[existing_idx], INDEPENDENT_COUNSEL_TEXT)
        return
    contingency_idx = find_paragraph_containing(doc, "Seller's attorney may make or require")
    if contingency_idx is None:
        return
    anchor = doc.paragraphs[contingency_idx]
    paragraph = insert_paragraph_after(anchor, INDEPENDENT_COUNSEL_TEXT)
    copy_paragraph_format(anchor, paragraph)
    source_bold = next((run.bold for run in anchor.runs if run.text), None)
    if source_bold is not None:
        for run in paragraph.runs:
            run.bold = source_bold


def set_adverse_condition(doc, text):
    pending_idx = find_paragraph(doc, "PENDING ORDERS OR ADVERSE CONDITIONS:")
    cure_idx = find_paragraph(doc, "RIGHT TO CURE DEFAULT")
    if pending_idx is None:
        return
    text = str(text or "").strip()
    if not text:
        text = "NOTE FOR ATTORNEY REVIEW: Confirm whether any pending orders, liens, adverse conditions, title matters, or required disclosures should be listed here before signing."
    use_responsibility_text = not text.startswith("NOTE FOR ATTORNEY REVIEW:")
    inserted_responsibility = False
    inserted_adverse = False
    for idx in range(pending_idx + 1, cure_idx or len(doc.paragraphs)):
        if doc.paragraphs[idx].text.strip():
            if use_responsibility_text and not inserted_responsibility:
                set_paragraph(doc.paragraphs[idx], ADVERSE_CONDITION_RESPONSIBILITY_TEXT)
                inserted_responsibility = True
            elif not inserted_adverse:
                set_paragraph(doc.paragraphs[idx], text)
                inserted_adverse = True
            else:
                set_paragraph(doc.paragraphs[idx], "")
    if inserted_adverse:
        return
    if use_responsibility_text and inserted_responsibility:
        anchor = doc.paragraphs[pending_idx + 1]
        paragraph = insert_paragraph_after(anchor, text)
        copy_paragraph_format(anchor, paragraph)
        return
    if pending_idx + 1 < len(doc.paragraphs):
        set_paragraph(doc.paragraphs[pending_idx + 1], text)


def set_seller_signature_manager(doc, x):
    signature_text = f"{x['trustee']}, Trustee - {x['manager']}, Manager"
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        if "[MANAGER NAME]" in text or (x["trustee"] in text and "Trustee" in text and "Manager" in text):
            set_paragraph(paragraph, signature_text)
            if idx + 1 < len(doc.paragraphs) and doc.paragraphs[idx + 1].text.strip() == f"{x['manager']}, Manager":
                set_paragraph(doc.paragraphs[idx + 1], "")
            return


def set_seller_notary_manager(doc, x):
    purchaser_sig_idx = find_paragraph(doc, "Purchaser’s Signature")
    if purchaser_sig_idx is None:
        purchaser_sig_idx = find_paragraph(doc, "Purchaser's Signature")
    for idx, paragraph in enumerate(doc.paragraphs):
        if purchaser_sig_idx is not None and idx >= purchaser_sig_idx:
            return
        text = paragraph.text
        if "personally appeared before me this day and acknowledged the execution" in text and (
            "[MANAGER NAME]" in text or x["manager"] not in text
        ):
            set_notary_acknowledgment(paragraph, x["manager"])
            return


def trim_after_last_notary(doc):
    body = doc._body._element
    paragraphs = list(body.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"))
    last_notary = None
    for paragraph in paragraphs:
        text = "".join(node.text or "" for node in paragraph.iter() if node.tag.endswith("}t"))
        if "Notary Public" in text and "My Commission expires" in text:
            last_notary = paragraph
    if last_notary is None:
        return
    children = list(body)
    try:
        last_idx = children.index(last_notary)
    except ValueError:
        return
    for child in children[last_idx + 1 :]:
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def set_receipt_acknowledgment(doc, x):
    text = f"[ ] Agent   [X] Seller, by the signature below, acknowledge receipt of {money(x['earnest_money'])} [  ] Cash   [   ] Check, as earnest money deposit, which is credited toward the total down payment mentioned in paragraph 1A of this Agreement. The remaining down payment balance of {money(x['remaining_down_payment'])} shall be paid at closing."
    set_first_paragraph_starting_any(doc, ["[ ] Agent"], text)


def set_earnest_money_terms(doc, x):
    set_first_paragraph_starting_any(
        doc,
        [
            "Earnest money deposit",
            "Binder deposit which will remain as a binder",
        ],
        "Earnest money deposit to be paid by Purchaser upon execution of this Contract. This earnest money shall remain as a binder until closing, unless sooner forfeited or returned according to the provisions in this Agreement, and shall be credited toward the total down payment.",
    )
    set_first_paragraph_starting_any(
        doc,
        [
            "Remaining balance of the down payment",
            "Additional binder deposit due",
        ],
        "Remaining balance of the down payment due at closing.",
    )


def set_purchaser_notary(doc, x):
    seen_seller_notary = False
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if x["manager"] in text and "personally appeared before me" in text:
            seen_seller_notary = True
            continue
        if seen_seller_notary and text.strip().startswith("COUNTY:"):
            set_paragraph(paragraph, f"COUNTY:  {x['county'].title()}")
            continue
        if seen_seller_notary and "personally appeared before me this day and acknowledged the execution" in text:
            set_notary_acknowledgment(paragraph, x["buyer"])
            return


def apply_notary_block_revisions(doc, x):
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        if x["buyer"] in text and "personally appeared before me this day and acknowledged the execution" in text:
            has_state = False
            has_county = False
            for prev_idx in range(idx - 1, max(-1, idx - 6), -1):
                prev_text = doc.paragraphs[prev_idx].text.strip()
                if prev_text.startswith("COUNTY"):
                    set_paragraph(doc.paragraphs[prev_idx], f"COUNTY:  {x['county'].title()}")
                    has_county = True
                if prev_text.startswith("STATE"):
                    has_state = True
            if not has_county:
                if not has_state:
                    paragraph.insert_paragraph_before("STATE:   North Carolina")
                    has_state = True
                paragraph.insert_paragraph_before(f"COUNTY:  {x['county'].title()}")
            if not has_state:
                paragraph.insert_paragraph_before("STATE:   North Carolina")
            set_notary_acknowledgment(paragraph, x["buyer"])

        if x["manager"] in text and "personally appeared before me this day and acknowledged the execution" in text:
            for prev_idx in range(idx - 1, max(-1, idx - 8), -1):
                if doc.paragraphs[prev_idx].text.strip().startswith("COUNTY:"):
                    set_paragraph(doc.paragraphs[prev_idx], f"COUNTY:  {x['county'].title()}")
                    break
            signer_text = f"{x['manager']}, Manager of {x['trustee']}, Trustee of {x['trust']},"
            prefix_idx = None
            for prev_idx in range(idx - 1, max(-1, idx - 5), -1):
                prev_text = doc.paragraphs[prev_idx].text
                if prev_text.startswith("I,") and "do hereby certify that" in prev_text:
                    prefix_idx = prev_idx
                    break
            if prefix_idx is not None:
                set_notary_certification_prefix(doc.paragraphs[prefix_idx])
                set_notary_signer_continuation(paragraph, signer_text)
            else:
                set_notary_acknowledgment(paragraph, signer_text)


def insert_paragraph_after(paragraph, text):
    new_paragraph = paragraph.insert_paragraph_before(text)
    paragraph._p.addnext(new_paragraph._p)
    return new_paragraph


def ensure_ach_terms(doc):
    ach_paragraph = (
        "ACH DRAFT PAYMENTS: Purchaser shall make each monthly payment by automatic ACH draft. "
        "Seller shall set up the ACH draft using the bank account and routing information provided by Purchaser on the signature page."
    )
    ach_fields = [
        "ACH Draft Authorization:",
        "Bank Account Number: ________________________________________________",
        "Routing Number: _____________________________________________________",
    ]

    # If the prototype already contains these blocks, leave them in place so user-set
    # numbering/lettering, spacing, and section placement remain intact.
    ach_blocks = [p for p in doc.paragraphs if p.text.strip().startswith("ACH DRAFT PAYMENTS:")]
    if not ach_blocks:
        insert_after_idx = find_paragraph(doc, "Total Number of Installment Payments:")
        if insert_after_idx is not None:
            p = insert_paragraph_after(doc.paragraphs[insert_after_idx], ach_paragraph)
            for run in p.runs:
                run.bold = False
    for paragraph in ach_blocks[1:]:
        paragraph._element.getparent().remove(paragraph._element)

    for field in ach_fields:
        matches = [p for p in doc.paragraphs if p.text.strip().startswith(field.split(":")[0] + ":")]
        for paragraph in matches[1:]:
            paragraph._element.getparent().remove(paragraph._element)

    if not any(p.text.strip().startswith("ACH Draft Authorization:") for p in doc.paragraphs):
        state_idx = None
        for idx, paragraph in enumerate(doc.paragraphs):
            if idx > 100 and paragraph.text.strip().startswith("STATE:"):
                state_idx = idx
                break
        if state_idx is not None:
            anchor = doc.paragraphs[state_idx]
            for text in ach_fields + [""]:
                p = anchor.insert_paragraph_before(text)
                if text == "ACH Draft Authorization:":
                    for run in p.runs:
                        run.bold = True


def main():
    x = normalize_values(get_docs_values())
    buyer_line1, buyer_line2 = split_address(x["buyer_address"])
    trustee_line1, trustee_line2 = split_address(x["trustee_address"])

    source = PROTOTYPE if PROTOTYPE.exists() else REFERENCE
    doc = Document(str(source))
    remove_after_contract(doc)

    # Direct paragraph substitutions at the same paragraph positions as the Cool Springs contract.
    set_paragraph(doc.paragraphs[3], f"{x['trust']}, by and through {x['trustee']}, Trustee, (Seller)")
    set_paragraph(doc.paragraphs[5], trustee_line1)
    set_paragraph(doc.paragraphs[6], trustee_line2)
    set_paragraph(doc.paragraphs[8], f"{x['buyer']}, (Purchaser)")
    set_paragraph(doc.paragraphs[10], buyer_line1)
    set_paragraph(doc.paragraphs[11], buyer_line2)
    set_paragraph(doc.paragraphs[17], x["property_address"])
    set_paragraph(doc.paragraphs[18], x["property_city_state"])
    set_paragraph(doc.paragraphs[21], x["brief_legal"] or x["legal"])
    set_county_of_paragraph(doc.paragraphs[23], x["county"])

    set_first_paragraph_starting(
        doc,
        "INTEREST RATE:",
        f"INTEREST RATE: Unpaid principal shall accrue interest at a rate of {x['interest']}(APR), with interest to be amortized over the term of this Contract for Deed. The Amount of Each Installment Payment above includes any interest payable, property Insurance, and escrowed Property Tax. Late payments shall accrue interest at a rate of four percent (4%) on any delinquent monies owed.",
    )
    set_first_paragraph_starting(doc, "For the first", f"For the first {x['term_months']} instalments:")
    set_first_paragraph_starting(doc, "Due Date of the First Installment Payment:", f"Due Date of the First Installment Payment:\t{short_date(x['loan_start'])}")
    set_first_paragraph_starting(doc, "Due Date of the Last Installment Payment:", f"Due Date of the Last Installment Payment:\t{short_date(x['loan_end'])}")
    set_first_paragraph_starting(doc, "Principal and Interest:", f"Principal and Interest:\t{money(x['monthly_pi']).replace('$', '')}")
    set_first_paragraph_starting(doc, "Property Insurance:", f"Property Insurance:\t\t {money(x['insurance']).replace('$', '')}")
    set_first_paragraph_starting(doc, "Property Tax:", f"Property Tax:\t\t {money(x['tax']).replace('$', '')}")
    set_first_paragraph_starting(doc, "Total Monthly Payment:", f"Total Monthly Payment:\t{money(x['total_payment']).replace('$', '')}")
    set_first_paragraph_starting(doc, "Total Number of Installment Payments:", f"Total Number of Installment Payments:\t{x['term_months']}")
    replace_installment_block_with_table(doc, x)
    set_adverse_condition(doc, x["adverse"])
    set_earnest_money_terms(doc, x)

    set_seller_signature_manager(doc, x)
    set_receipt_acknowledgment(doc, x)
    set_seller_notary_manager(doc, x)
    set_purchaser_notary(doc, x)
    apply_notary_block_revisions(doc, x)
    ensure_notary_signature_lines(doc)
    standardize_contract_notary_blocks(doc, x)
    ensure_ach_terms(doc)
    ensure_attorney_contingency_clause(doc)
    ensure_independent_counsel_clause(doc)

    # Table 0 is the price table in the contract. Keep the table formatting and replace only values.
    table = doc.tables[0]
    set_cell_text_preserve_formatting(table.cell(0, 0), "$")
    set_cell_text_preserve_formatting(table.cell(0, 1), money(x["earnest_money"]).replace("$", ""))
    if len(table.rows) > 1:
        set_cell_text_preserve_formatting(table.cell(1, 0), "$")
        set_cell_text_preserve_formatting(table.cell(1, 1), money(x["remaining_down_payment"]).replace("$", ""))
    loan_row = 6 if len(table.rows) > 6 else len(table.rows) - 2
    total_row = 8 if len(table.rows) == 9 else 9
    if len(table.rows) > 2:
        set_cell_text_preserve_formatting(table.cell(2, 1), "")
    set_cell_text_preserve_formatting(table.cell(loan_row, 1), money(x["loan_amount"]).replace("$", ""))
    set_cell_text_preserve_formatting(table.cell(total_row, 1), money(x["sale_price"]).replace("$", ""))

    replace_text_in_paragraphs(
        doc,
        {
            "02/    /2025.": "06/    /2026.",
            "2025.": "2026.",
            "Lee": x["county"].title(),
        },
    )

    standardize_contract_notary_blocks(doc, x)
    trim_after_last_notary(doc)
    replace_legacy_buyer_name_variants(doc, x)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
