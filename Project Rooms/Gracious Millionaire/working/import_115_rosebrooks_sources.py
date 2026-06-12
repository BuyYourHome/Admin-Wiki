from __future__ import annotations

import csv
import hashlib
import html
import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


ROOT = Path(r"C:\Codex\Wiki Files")
SOURCE_ROOT = Path(r"C:\Users\wesbr\Buy Your Home\Buy Your Home - Property\19-HM-115 Rosebrooks Dr")
OUT_ROOT = ROOT / "Project Rooms" / "Gracious Millionaire" / "sources" / "115-rosebrooks-project"
INVENTORY = OUT_ROOT / "_source-inventory.md"

TEXT_EXTS = {".txt", ".csv", ".html"}
DOCX_EXTS = {".docx"}
XLSX_EXTS = {".xlsx", ".xls"}
PDF_EXTS = {".pdf"}
MEDIA_EXTS = {".jpg", ".jpeg", ".png", ".mp3", ".mp4", ".3gp"}
OTHER_EXTS = {".doc", ".msg"}


def slug(text: str, limit: int = 90) -> str:
    text = re.sub(r"[^A-Za-z0-9]+", "-", text).strip("-").lower()
    return text[:limit].strip("-") or "file"


def rel_slug(path: Path) -> str:
    rel = path.relative_to(SOURCE_ROOT)
    parts = [slug(p, 48) for p in rel.parts]
    return "__".join(parts)


def file_hash(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def strip_html(raw: str) -> str:
    raw = re.sub(r"(?is)<(script|style).*?>.*?</\\1>", " ", raw)
    raw = re.sub(r"(?i)<br\\s*/?>", "\n", raw)
    raw = re.sub(r"(?i)</p\\s*>", "\n\n", raw)
    raw = re.sub(r"<[^>]+>", " ", raw)
    raw = html.unescape(raw)
    return normalize_text(raw)


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \\t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def read_text(path: Path) -> str:
    raw = None
    for enc in ("utf-8-sig", "utf-16", "cp1252", "latin-1"):
        try:
            raw = path.read_text(encoding=enc)
            break
        except UnicodeError:
            continue
    if raw is None:
        raw = path.read_text(errors="replace")
    if path.suffix.lower() == ".html":
        return strip_html(raw)
    if path.suffix.lower() == ".csv":
        rows = []
        with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as f:
            reader = csv.reader(f)
            for i, row in enumerate(reader):
                if i >= 250:
                    rows.append(["[truncated after 250 rows]"])
                    break
                rows.append(row)
        if not rows:
            return ""
        width = max(len(r) for r in rows)
        padded = [r + [""] * (width - len(r)) for r in rows]
        return "\n".join("| " + " | ".join(c.replace("\n", " ") for c in r) + " |" for r in padded)
    return normalize_text(raw)


def read_docx(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append("| " + " | ".join(cells) + " |")
    return normalize_text("\n\n".join(parts))


def read_xlsx(path: Path) -> str:
    wb = load_workbook(path, data_only=False, read_only=True)
    sections: list[str] = []
    for ws in wb.worksheets:
        sections.append(f"## Sheet: {ws.title}")
        max_row = min(ws.max_row or 0, 200)
        max_col = min(ws.max_column or 0, 40)
        if max_row == 0 or max_col == 0:
            sections.append("[empty sheet]")
            continue
        for row in ws.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col, values_only=True):
            values = ["" if v is None else str(v).replace("\n", " ") for v in row]
            if any(values):
                sections.append("| " + " | ".join(values) + " |")
        if (ws.max_row or 0) > max_row or (ws.max_column or 0) > max_col:
            sections.append(f"[truncated to {max_row} rows x {max_col} columns]")
    return "\n".join(sections).strip()


def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts: list[str] = []
    total_pages = len(reader.pages)
    for idx, page in enumerate(reader.pages[:40], start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001
            text = f"[page {idx} extraction failed: {exc}]"
        text = normalize_text(text)
        if text:
            parts.append(f"## Page {idx}\n\n{text}")
    if total_pages > 40:
        parts.append(f"[truncated after 40 of {total_pages} pages]")
    return "\n\n".join(parts).strip()


def write_markdown(path: Path, content: str, status: str, notes: str, used_names: set[str]) -> dict[str, str]:
    rel = path.relative_to(SOURCE_ROOT)
    base_name = rel_slug(path)
    stat = path.stat()
    digest = file_hash(path)
    out_name = f"{base_name}.md"
    if out_name in used_names:
        out_name = f"{base_name}-{digest[:8]}.md"
    suffix = 2
    while out_name in used_names:
        out_name = f"{base_name}-{digest[:8]}-{suffix}.md"
        suffix += 1
    used_names.add(out_name)
    out = OUT_ROOT / out_name
    body = [
        f"# {path.name}",
        "",
        "## Source Metadata",
        "",
        f"- Original path: `{path}`",
        f"- Relative path: `{rel}`",
        f"- File type: `{path.suffix.lower() or '[none]'}`",
        f"- Size bytes: {stat.st_size}",
        f"- Last modified: {datetime.fromtimestamp(stat.st_mtime).isoformat(timespec='seconds')}",
        f"- SHA256: `{digest}`",
        f"- Extraction status: {status}",
        f"- Notes: {notes}",
        "",
        "## Extracted Content",
        "",
        content.strip() or "[No extractable text content.]",
        "",
    ]
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text("\n".join(body), encoding="utf-8")
    return {
        "relative": str(rel).replace("\\", "/"),
        "markdown": str(out.relative_to(OUT_ROOT)).replace("\\", "/"),
        "status": status,
        "notes": notes,
        "size": str(stat.st_size),
    }


def main() -> None:
    if OUT_ROOT.exists():
        shutil.rmtree(OUT_ROOT)
    OUT_ROOT.mkdir(parents=True, exist_ok=True)

    rows = []
    used_names: set[str] = set()
    files = sorted(p for p in SOURCE_ROOT.rglob("*") if p.is_file())
    for path in files:
        ext = path.suffix.lower()
        status = "metadata-only"
        notes = "Unsupported or binary file; source metadata preserved."
        content = ""
        try:
            if ext in TEXT_EXTS:
                content = read_text(path)
                status = "text-extracted"
                notes = "Text extracted directly from source file."
            elif ext in DOCX_EXTS:
                content = read_docx(path)
                status = "text-extracted"
                notes = "Text extracted from DOCX paragraphs and tables."
            elif ext in XLSX_EXTS:
                content = read_xlsx(path)
                status = "text-extracted"
                notes = "Workbook sheets extracted to markdown-like tables; capped for very large sheets."
            elif ext in PDF_EXTS:
                content = read_pdf(path)
                status = "text-extracted" if content else "metadata-only"
                notes = "PDF text extracted with pypdf; scanned pages may have no text layer."
            elif ext in MEDIA_EXTS:
                notes = "Media file; no OCR/transcription performed in this pass."
            elif ext in OTHER_EXTS:
                notes = "Legacy Office or Outlook message file; metadata preserved in this pass."
        except Exception as exc:  # noqa: BLE001
            content = f"[Extraction failed: {exc}]"
            status = "extraction-failed"
            notes = str(exc).replace("\n", " ")
        rows.append(write_markdown(path, content, status, notes, used_names))

    inv_lines = [
        "# 115 Rosebrooks Source Inventory",
        "",
        f"- Source root: `{SOURCE_ROOT}`",
        f"- Files processed: {len(rows)}",
        "",
        "| Original file | Markdown file | Status | Size bytes | Notes |",
        "| --- | --- | --- | ---: | --- |",
    ]
    for row in rows:
        inv_lines.append(
            f"| `{row['relative']}` | `{row['markdown']}` | {row['status']} | {row['size']} | {row['notes']} |"
        )
    INVENTORY.write_text("\n".join(inv_lines) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()
