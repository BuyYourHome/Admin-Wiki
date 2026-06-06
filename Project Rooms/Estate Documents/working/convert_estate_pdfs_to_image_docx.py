from io import BytesIO
from pathlib import Path

import fitz
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches


ROOT = Path(r"C:\Codex\Wiki Files\Project Rooms\Estate Documents")
SOURCE_DIR = ROOT / "sources" / "Office Admin Estate"
OUTPUT_DIR = ROOT / "outputs" / "word-conversions-no-signatures"


def convert_pdf(pdf_path: Path) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.25)
    section.bottom_margin = Inches(0.25)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)

    pdf = fitz.open(pdf_path)
    for index, page in enumerate(pdf):
        if index:
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            section.top_margin = Inches(0.25)
            section.bottom_margin = Inches(0.25)
            section.left_margin = Inches(0.25)
            section.right_margin = Inches(0.25)

        pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
        image_bytes = pix.tobytes("png")
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0
        run = paragraph.add_run()
        run.add_picture(BytesIO(image_bytes), width=Inches(8.0))

    out_path = OUTPUT_DIR / f"{pdf_path.stem}.docx"
    doc.save(out_path)
    return out_path


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for pdf in sorted(SOURCE_DIR.glob("*.pdf")):
        print(f"Converting {pdf.name}")
        print(f"  -> {convert_pdf(pdf)}")


if __name__ == "__main__":
    main()
