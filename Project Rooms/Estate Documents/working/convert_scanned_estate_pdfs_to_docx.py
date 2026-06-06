from pathlib import Path
import re

import fitz
import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from rapidocr_onnxruntime import RapidOCR
from PIL import Image


ROOT = Path(r"C:\Codex\Wiki Files\Project Rooms\Estate Documents")
SOURCE_DIR = ROOT / "sources" / "Office Admin Estate"
OUTPUT_DIR = ROOT / "outputs" / "word-conversions-no-signatures"


SIGNATURE_LINE_HINTS = (
    "signature",
    "signed",
    "initials",
    "witness",
    "notary public",
    "my commission expires",
    "subscribed",
    "sworn",
    "acknowledged",
)


def clean_text(text: str) -> str:
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\u2018", "'").replace("\u2019", "'")
    text = text.replace("?", "'")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def likely_handwriting_artifact(text: str, score: float) -> bool:
    cleaned = re.sub(r"[^A-Za-z0-9]", "", text)
    if not cleaned:
        return True
    if score < 0.62:
        return True
    if score < 0.82 and len(cleaned) <= 8:
        return True
    return False


def should_add_blank_line(text: str) -> bool:
    lower = text.lower()
    return any(hint in lower for hint in SIGNATURE_LINE_HINTS)


def add_blank_signature_line(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("_" * 42)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def add_ocr_line(doc: Document, text: str, score: float) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    if text.isupper() and len(text) < 70:
        run.bold = True
        p.alignment = 1
    if should_add_blank_line(text):
        add_blank_signature_line(doc)


def convert_pdf(pdf_path: Path, ocr: RapidOCR) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    source = fitz.open(pdf_path)
    total_lines = 0
    for page_index, page in enumerate(source):
        if page_index:
            doc.add_page_break()
        pix = page.get_pixmap(matrix=fitz.Matrix(1.75, 1.75), alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        result, _ = ocr(np.array(img))
        rows = []
        for row in result or []:
            box, text, score = row
            text = clean_text(text)
            if likely_handwriting_artifact(text, score):
                continue
            y = sum(point[1] for point in box) / 4
            x = sum(point[0] for point in box) / 4
            rows.append((y, x, text, score))
        rows.sort(key=lambda item: (item[0], item[1]))

        last_y = None
        for y, _x, text, score in rows:
            if last_y is not None and y - last_y > 42:
                doc.add_paragraph()
            add_ocr_line(doc, text, score)
            total_lines += 1
            last_y = y

    if total_lines == 0:
        p = doc.add_paragraph()
        run = p.add_run("No readable text detected in scanned source PDF.")
        run.italic = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    out_path = OUTPUT_DIR / f"{pdf_path.stem}.docx"
    doc.save(out_path)
    return out_path


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ocr = RapidOCR()
    for pdf in sorted(SOURCE_DIR.glob("*.pdf")):
        print(f"Converting {pdf.name}")
        out_path = convert_pdf(pdf, ocr)
        print(f"  -> {out_path}")


if __name__ == "__main__":
    main()
