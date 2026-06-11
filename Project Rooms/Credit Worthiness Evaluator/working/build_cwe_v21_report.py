from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(r"C:\Codex\Wiki Files")
OUT_DIR = PROJECT_ROOT / r"Project Rooms\Credit Worthiness Evaluator\outputs\320 Rose Pl - Ever Cardoza"
SOURCE_DIR = PROJECT_ROOT / r"Project Rooms\Credit Worthiness Evaluator\sources\320 Rose Pl - Ever Cardoza"
BUYER_FOLDER = Path(r"C:\Users\wesbr\Buy Your Home\Buy Your Home - Property\28-SYH-320 Rose Pl\Selling\Ever Cordoza")
SPREADSHEET_PATH = Path(r"C:\Users\wesbr\Buy Your Home\Buy Your Home - Property\28_Project Management - 320 Rose Pl.xlsm")

CONTRACT_PACKAGE = BUYER_FOLDER / "Contract Package"
REPORT_DIR = CONTRACT_PACKAGE
ARCHIVE_DIR = CONTRACT_PACKAGE / "Credit Worthiness Archive"
OLD_CW_DIR = CONTRACT_PACKAGE / "Credit Worthiness"
WRONG_REPORT_DIR = BUYER_FOLDER / "Contract Package Report"


def money(value):
    return "${:,.2f}".format(value)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_table(doc, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, value in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[index], value, True)
    for row in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], title, True)
    set_cell_text(table.rows[1].cells[0], body)


def add_page_numbers(doc):
    paragraph = doc.sections[0].footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Page ")
    for instruction in ("PAGE", "NUMPAGES"):
        if instruction == "NUMPAGES":
            paragraph.add_run(" of ")
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        text = OxmlElement("w:instrText")
        text.set(qn("xml:space"), "preserve")
        text.text = instruction
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        paragraph._p.append(begin)
        paragraph._p.append(text)
        paragraph._p.append(end)


def archive_existing_teams_reports():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    CONTRACT_PACKAGE.mkdir(parents=True, exist_ok=True)
    archive_action = None

    if OLD_CW_DIR.exists():
        if ARCHIVE_DIR.exists():
            for child in OLD_CW_DIR.iterdir():
                destination = ARCHIVE_DIR / child.name
                if destination.exists():
                    stem = child.stem if child.is_file() else child.name
                    suffix = child.suffix if child.is_file() else ""
                    counter = 1
                    while True:
                        candidate = ARCHIVE_DIR / f"{stem} {counter}{suffix}"
                        if not candidate.exists():
                            destination = candidate
                            break
                        counter += 1
                shutil.move(str(child), str(destination))
            try:
                OLD_CW_DIR.rmdir()
            except OSError:
                pass
            archive_action = f"Moved contents from {OLD_CW_DIR} into existing {ARCHIVE_DIR} without overwriting."
        else:
            shutil.move(str(OLD_CW_DIR), str(ARCHIVE_DIR))
            archive_action = f"Renamed {OLD_CW_DIR} to {ARCHIVE_DIR}."

    current_name = "320 Rose Ever Amarildo Cardoza Bolanos - Creditworthiness Evaluation Report - Document Preparation Ready.docx"
    current_report = REPORT_DIR / current_name
    archived_current = None

    if WRONG_REPORT_DIR.exists():
        for wrong_file in WRONG_REPORT_DIR.glob("*Creditworthiness Evaluation Report*.docx"):
            ARCHIVE_DIR.mkdir(parents=True, exist_ok=True)
            counter = 1
            while True:
                candidate = ARCHIVE_DIR / f"v{counter} {wrong_file.name}"
                if not candidate.exists():
                    break
                counter += 1
            shutil.move(str(wrong_file), str(candidate))
            archived_current = candidate
            archive_action = (archive_action + " " if archive_action else "") + f"Moved prior current report from {WRONG_REPORT_DIR} to {candidate}."
        try:
            WRONG_REPORT_DIR.rmdir()
        except OSError:
            pass

    if current_report.exists():
        ARCHIVE_DIR.mkdir(parents=True, exist_ok=True)
        counter = 1
        while True:
            candidate = ARCHIVE_DIR / f"v{counter} {current_name}"
            if not candidate.exists():
                break
            counter += 1
        shutil.move(str(current_report), str(candidate))
        archived_current = candidate

    return current_report, archive_action, archived_current


def next_project_report_path():
    base = OUT_DIR / "26-06-11 320 Rose Ever Amarildo Cardoza Bolanos - Creditworthiness Evaluation Report - Document Preparation Ready v22.docx"
    if not base.exists():
        return base
    counter = 1
    while True:
        candidate = OUT_DIR / f"26-06-11 320 Rose Ever Amarildo Cardoza Bolanos - Creditworthiness Evaluation Report - Document Preparation Ready v22 {counter}.docx"
        if not candidate.exists():
            return candidate
        counter += 1


def build_report(path):
    purchase_price = 346500
    loan_amount = 325710
    payment = 2504.43
    down_payment = 20790
    earnest = 6000
    stated_income = 9500
    credit_debt = 2162
    housing_ratio = payment / stated_income
    back_end = (payment + credit_debt) / stated_income
    verified_rent = 1850
    rent_delta = payment - verified_rent
    business_avg_deposits = 27773.92
    business_avg_withdrawals = 27622.03
    business_total_deposits = 472156.56
    business_total_withdrawals = 469574.46
    business_latest_balance = 5058.11
    personal_avg_deposits = 2040.14

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(9.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("UPDATED CREDITWORTHINESS EVALUATION REPORT")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph(
        "Document-preparation readiness rerun for Ever Amarildo Cardoza Bolanos "
        "and Maria Geraldina Sarmiento, 320 Rose Pl, Wendell, North Carolina"
    )

    add_table(
        doc,
        [
            ["Prepared by", "Investment Services, LLC"],
            ["Report date", "June 11, 2026"],
            ["Subject buyer", "Ever Amarildo Cardoza Bolanos"],
            ["Additional signing buyer from spreadsheet", "Maria Geraldina Sarmiento"],
            ["Property", "320 Rose Pl, Wendell, NC 27591"],
            ["Purpose", "Current evidence rerun for Sell Your Home, LLC document-preparation decision and Contract for Deed closing-package support."],
            ["Overall result", "Ready for document preparation; closing execution remains subject to required affidavits, funds proof, final legal-name/spelling confirmation, and attorney/compliance review."],
        ],
        [2.2, 5.8],
    )

    doc.add_paragraph()
    add_callout(
        doc,
        "Executive Recommendation",
        "Proceed with Contract for Deed document preparation. The current file is supportable for document preparation based on verified rent history with related company Buy Your Home, LLC, satisfactory screening/credit indicators, observed business cash flow, and affidavit-supported business judgment. The file should not be marked ready for closing execution until closing deliverables are completed, funds-to-close are proved or accepted by the closing process, buyer name spelling is confirmed, and counsel/compliance review confirms the seller-financing and Contract for Deed package.",
    )

    def heading(text):
        doc.add_heading(text, level=1)

    def paragraph(text):
        doc.add_paragraph(text)

    heading("1. What Changed In This Rerun")
    paragraph("This rerun uses the refreshed buyer source mirror copied from the Teams buyer folder on June 11, 2026. Four newer files were copied into the project-room source mirror, primarily current Contract Package email/report materials. No new tax return, profit-and-loss statement, payroll record, owner-draw analysis, or final funds-to-close proof was identified in the refreshed source set.")
    paragraph("The live project spreadsheet was rechecked. It identifies Ever Amarildo Cardoza Bolanos as Selling-Buyer1 and Maria Geraldina Sarmiento as Selling-Buyer2. It shows Sell Your Home, LLC as seller, a purchase price of $346,500, loan amount of $325,710, monthly payment of $2,504.43, and manager field of Jose Fabre Jr. Older materials that used different manager or name-spelling language should be treated as requiring final authority/name review before execution.")

    heading("2. Transaction Snapshot")
    add_table(
        doc,
        [
            ["Item", "Current source value"],
            ["Buyer 1", "Ever Amarildo Cardoza Bolanos"],
            ["Buyer 2 / signing buyer", "Maria Geraldina Sarmiento"],
            ["Buyer address", "4121 Tensity Dr, Raleigh, NC 27604"],
            ["Seller", "Sell Your Home, LLC"],
            ["Property", "320 Rose Pl, Wendell, NC 27591"],
            ["County", "Wake County, North Carolina"],
            ["Purchase price", money(purchase_price)],
            ["Earnest money", money(earnest)],
            ["Down payment", money(down_payment)],
            ["Loan amount", money(loan_amount)],
            ["Monthly payment", money(payment)],
            ["Loan start / end", "July 1, 2026 / June 1, 2056"],
            ["Trustee", "Investment Services LLC"],
            ["Manager field in spreadsheet", "Jose Fabre Jr."],
        ],
        [2.5, 5.5],
    )

    heading("3. Screening Strengths")
    paragraph("The Zillow screening materials identify Ever Amarildo Cardoza as identity verified, self-employed, and residing at 4121 Tensity Dr since June 1, 2025. The stated monthly self-employment income is $9,500. Screening shows no pets and a stated move reason of buying a house.")
    paragraph("The credit report dated June 5, 2026 shows a 706 VantageScore 4.0, 100% on-time payments, no collections found, and no bankruptcies found. Reported monthly debt payments are $2,162, with total debt of $86,790. The positive credit history is a meaningful compensating factor, but the auto-loan debt load remains material.")

    heading("4. Ability-To-Repay Findings")
    paragraph("On strict underwriting evidence, borrower-level net self-employment income remains incompletely documented because the file still lacks tax returns, a profit-and-loss statement, payroll records, or an owner-draw analysis. The evaluator should not convert gross business receipts into verified personal net income without those records.")
    paragraph("For document-preparation readiness, the file remains supportable because the proposed housing payment is only $654.43 above the verified related-company rent of $1,850, the buyer has a direct rent-payment track record with Buy Your Home, LLC, and the available business records show sustained gross activity sufficient to support a business-judgment review. This is not the same as final closing execution approval.")

    heading("5. Income And Business Cash Flow")
    add_table(
        doc,
        [
            ["Metric", "Current evidence"],
            ["Applicant-stated self-employed income", money(stated_income) + " per month"],
            ["Business bank statements reviewed", "17 statement periods summarized"],
            ["Average business deposits", money(business_avg_deposits) + " per month"],
            ["Average business withdrawals", money(business_avg_withdrawals) + " per month"],
            ["Total business deposits summarized", money(business_total_deposits)],
            ["Total business withdrawals summarized", money(business_total_withdrawals)],
            ["Latest business ending balance", money(business_latest_balance)],
            ["Personal bank statements summarized", "12 statement periods; average deposits " + money(personal_avg_deposits) + " per month"],
            ["Receipts/receivables", "Receipt packages support gross receipts only; they do not establish borrower-level net income by themselves."],
        ],
        [2.8, 5.2],
    )
    paragraph("The business account activity is consistent with an active operating business, but the same records also show business withdrawals nearly equal to deposits. That pattern supports the need for business-judgment documentation, affidavit support, and attorney/compliance review rather than a clean income-verification approval.")

    heading("6. Debt Burden")
    add_table(
        doc,
        [
            ["Debt item", "Amount / status"],
            ["Credit-report estimated monthly payments", money(credit_debt)],
            ["Total reported debt", "$86,790"],
            ["Collections", "None found"],
            ["Bankruptcies", "None found"],
            ["Current rent verified by related company", money(verified_rent) + " per month"],
            ["Increase from verified rent to proposed payment", money(rent_delta) + " per month"],
        ],
        [3.2, 4.8],
    )
    paragraph("The debt load is high enough that a strict debt-to-income view depends heavily on the stated income figure and the business-judgment treatment of nonstandard income evidence.")

    heading("7. Payment Ratios")
    add_table(
        doc,
        [
            ["Ratio test", "Calculation", "Result"],
            ["Housing ratio using stated income", f"{money(payment)} / {money(stated_income)}", f"{housing_ratio:.1%}"],
            ["Back-end ratio using stated income", f"({money(payment)} + {money(credit_debt)}) / {money(stated_income)}", f"{back_end:.1%}"],
            ["Increase over verified rent", f"{money(payment)} - {money(verified_rent)}", money(rent_delta)],
        ],
        [2.8, 3.4, 1.8],
    )
    paragraph("Using stated income, the back-end ratio is approximately 49.1%. That ratio is elevated, but it is not by itself a document-preparation stop when the seller/business decides to proceed based on the buyer's verified rent history, credit performance, and required closing affidavits. It remains a compliance and business-judgment issue for final review.")

    heading("8. Cash To Close And Reserves")
    paragraph("The spreadsheet requires $6,000 earnest money and total down payment of $20,790. Current support materials include an affidavit package addressing observed cash, receivables, and foreign-fund access, but the file still needs final funds-to-close proof or closing-accepted handling before execution.")
    add_table(
        doc,
        [
            ["Item", "Status"],
            ["Earnest money", "Spreadsheet shows $6,000; final proof/receipt handling still required."],
            ["Remaining down payment", "Total down payment shown as $20,790; final closing proof still required."],
            ["Cash/reserve observation affidavit", "Useful support if signed/notarized and accepted by counsel/closing process."],
            ["Receivables", "Do not count as liquid reserves unless collection is proved or accepted by closing reviewer."],
        ],
        [2.6, 5.4],
    )

    heading("9. Decision")
    add_callout(doc, "Decision", "Credit/evaluation status: supportable for document preparation under a business-judgment and closing-deliverables approach. Closing-package readiness status: Ready for document preparation. Not ready for closing execution until listed closing deliverables and legal/compliance review are complete.")

    heading("10. Conditions Or Documents Required Before Approval/Closing")
    for item in [
        "Signed/notarized support affidavits or counsel-approved substitutes.",
        "Final proof or closing-accepted handling for earnest money, down payment, closing costs, and any expected post-closing reserves.",
        "Final legal-name spelling confirmation for Ever Amarildo Cardoza Bolanos and Maria Geraldina Sarmiento.",
        "Resolution of any signature-authority conflict created by spreadsheet manager fields versus affidavit or prior workflow materials.",
        "Attorney/compliance review of seller-financing, ability-to-repay, disclosure, state Contract for Deed, title, recording, notice, and adverse-action issues.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    heading("11. CFD Closing-Package Document Requests")
    add_table(
        doc,
        [
            ["Requested item", "Purpose", "Status"],
            ["Affidavit of Related-Company Rent Payment History", "Documents rent history with Buy Your Home, LLC and related-company limitation.", "Closing deliverable if relied on."],
            ["Affidavit of Cash Reserves and Receivables Observation", "Supports observed cash/reserve/receivable facts without overstating collection or transfer.", "Closing deliverable or substitute proof needed."],
            ["Affidavit of Receipt Package Review and Acceptance", "Documents business acceptance of paid receipts as gross-receipt support, not net-income proof.", "Closing deliverable if receipts are relied on."],
            ["Affidavit of Business Judgment Approval Direction", "Documents management decision to proceed despite nonstandard income proof and elevated back-end ratio.", "Closing deliverable if management proceeds on this basis; update payment/authority values before final use."],
        ],
        [2.4, 3.8, 1.8],
    )
    paragraph("The file is ready for CFD document preparation. The affidavits and final proof items should be treated as closing deliverables or final-execution conditions, not blockers to preparing the closing package.")

    heading("12. Legal/Compliance Review Items")
    for item in [
        "Confirm whether the transaction is subject to TILA/Regulation Z, Dodd-Frank ability-to-repay requirements, and/or any seller-financer exclusion.",
        "Confirm North Carolina Contract for Deed requirements, recording/notice duties, title/adverse-condition handling, and disclosure package sufficiency.",
        "Confirm seller, trust/trustee, manager, and exact signature authority for Sell Your Home, LLC / Investment Services LLC / project spreadsheet fields.",
        "Confirm whether adverse-action notice or similar documentation is needed if any term or approval condition is considered less favorable than requested.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    heading("13. Recommended Next Step")
    paragraph("Proceed with preparation of the Contract for Deed closing package, including the listed affidavit/support deliverables. Before execution, obtain final funds proof or closing-accepted substitute documentation, confirm buyer legal names and signature roles, resolve manager/signature authority, and complete attorney/compliance review.")

    add_page_numbers(doc)
    doc.save(path)

    return {
        "housing_ratio": housing_ratio,
        "back_end": back_end,
        "rent_delta": rent_delta,
        "business_avg_deposits": business_avg_deposits,
        "business_avg_withdrawals": business_avg_withdrawals,
    }


def write_summary(path, project_docx, teams_current, archive_action, metrics):
    summary = f"""# 320 Rose Pl - Ever Cardoza Current Evidence Evaluation Summary v22

## Run Summary

- Run type: current CWE rerun after filing-rule update.
- Buyer source folder refreshed from: `{BUYER_FOLDER}`
- Project-room source copy: `{SOURCE_DIR}`
- Source refresh date/time: June 11, 2026 at 5:19 PM Eastern.
- Current evaluation mode: evidence mode, not approval-assumption mode.
- Full DOCX report: `{project_docx}`
- Teams current report: `{teams_current}`
- Teams archive folder: `{ARCHIVE_DIR}`
- Archive action: {archive_action or 'No old Credit Worthiness folder needed to be renamed.'}

## Current Evidence Result

Credit/evaluation status: supportable for document preparation under a business-judgment/closing-deliverables approach; not final closing-execution approval.

Closing-package readiness status: Ready for document preparation.

## Spreadsheet Transaction Snapshot

Source: live project spreadsheet `{SPREADSHEET_PATH}`, `Docs` worksheet.

| Field | Current value |
| --- | --- |
| Buyer 1 | Ever Amarildo Cardoza Bolanos |
| Buyer 2 | Maria Geraldina Sarmiento |
| Buyer address | 4121 Tensity Dr; Raleigh, NC 27604 |
| Seller | Sell Your Home, LLC |
| Purchase price | $346,500 |
| Earnest money | $6,000 |
| Down payment | $20,790 |
| Loan amount | $325,710 |
| Monthly payment | $2,504.43 |
| Loan start | 2026-07-01 |
| Loan end | 2056-06-01 |
| County | Wake |
| Trustee | Investment Services LLC |
| Manager field | Jose Fabre Jr. |

## Key Underwriting Numbers

- Stated self-employed income: $9,500/month.
- Credit-report monthly debt payments: $2,162/month.
- Housing ratio using stated income: {metrics['housing_ratio']:.1%}.
- Back-end ratio using stated income: {metrics['back_end']:.1%}.
- Verified related-company rent: $1,850/month.
- Proposed payment increase over verified rent: ${metrics['rent_delta']:,.2f}/month.
- Average business deposits summarized: ${metrics['business_avg_deposits']:,.2f}/month.
- Average business withdrawals summarized: ${metrics['business_avg_withdrawals']:,.2f}/month.

## Missing Before Closing Execution

- Signed/notarized support affidavits or counsel-approved substitutes.
- Funds-to-close proof or closing-accepted handling.
- Buyer legal-name spelling confirmation at closing.
- Seller/trustee/manager/signature authority confirmation.
- Attorney/compliance review.

## CFD Readiness

Ready for document preparation; closing execution subject to required affidavits, funds proof, final legal-name spelling confirmation, signature-authority confirmation, and attorney/compliance review.
"""
    path.write_text(summary, encoding="utf-8")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    teams_current, archive_action, archived_current = archive_existing_teams_reports()
    project_docx = next_project_report_path()
    metrics = build_report(project_docx)
    shutil.copy2(project_docx, teams_current)
    summary_path = OUT_DIR / "26-06-11 Current Evidence Evaluation Summary v22.md"
    write_summary(summary_path, project_docx, teams_current, archive_action, metrics)

    print(f"PROJECT_DOCX={project_docx}")
    print(f"TEAMS_CURRENT={teams_current}")
    print(f"ARCHIVE_DIR={ARCHIVE_DIR}")
    if archived_current:
        print(f"ARCHIVED_CURRENT={archived_current}")
    print(f"ARCHIVE_ACTION={archive_action}")
    print(f"SUMMARY={summary_path}")


if __name__ == "__main__":
    main()
