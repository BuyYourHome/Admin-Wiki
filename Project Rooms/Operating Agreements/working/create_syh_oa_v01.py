from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


BASE = Path(__file__).resolve().parent
SRC = BASE.parent / "sources" / "Jeff Watson" / "Simplified OA Subfiles" / "Reassembled OA Check - 00 - Reassembled Verification - Simplified OA with Jeff Watson tracked edits.docx"
OUT_DIR = BASE / "SYH"
OUT = OUT_DIR / "SYH V01 - Sell Your Home OA Draft - from Reassembled OA Check.docx"
NOTES = OUT_DIR / "SYH V01 - Sell Your Home OA Draft Notes.md"


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}r"):
            p.remove(child)


def add_page_number_run(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, text, fld_end])


def apply_version_footer(doc):
    for section in list(doc.sections)[:-1]:
        section.footer.is_linked_to_previous = False
        paragraph = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        clear_paragraph(paragraph)
        paragraph.alignment = 1
        paragraph.add_run("Page ")
        add_page_number_run(paragraph)
        paragraph.add_run(" | SYH V01 - Sell Your Home OA Draft")

    if len(doc.sections) > 1:
        legend_footer = doc.sections[-1].footer
        legend_footer.is_linked_to_previous = False
        for paragraph in legend_footer.paragraphs:
            clear_paragraph(paragraph)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(SRC)
    apply_version_footer(doc)
    doc.save(OUT)
    NOTES.write_text(
        "\n".join(
            [
                "# SYH V01 - Sell Your Home OA Draft Notes",
                "",
                "Mode: `SYH`.",
                "",
                "Source used:",
                "",
                "`sources/Jeff Watson/Simplified OA Subfiles/Reassembled OA Check - 00 - Reassembled Verification - Simplified OA with Jeff Watson tracked edits.docx`",
                "",
                "Output DOCX:",
                "",
                "`SYH V01 - Sell Your Home OA Draft - from Reassembled OA Check.docx`",
                "",
                "Build treatment:",
                "",
                "- Built from the current verified Reassembled OA Check.",
                "- No entity-conversion changes were applied.",
                "- Preserves current SYH source-stack text, colors, strikeouts, certification layout, Exhibit A, and Drafting Legend.",
                "- Adds the required page/version footer to agreement sections.",
                "- Leaves the final Drafting Legend section without a page/version footer.",
                "",
                "Open facts before finalization:",
                "",
                "- Confirm current controlling SYH member, manager, effective-date, certification, and signature facts.",
                "- Confirm attorney/CPA review before treating legal, tax, ownership, authority, or retirement-account language as final.",
                "",
            ]
        ),
        encoding="utf-8",
    )
    print(OUT)
    print(NOTES)


if __name__ == "__main__":
    build()
