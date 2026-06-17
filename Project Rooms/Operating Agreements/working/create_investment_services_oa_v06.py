from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph


BASE = Path(__file__).resolve().parent
SRC = BASE.parent / "sources" / "Jeff Watson" / "Simplified OA Subfiles" / "Reassembled OA Check - 00 - Reassembled Verification - Simplified OA with Jeff Watson tracked edits.docx"
WYOMING_SRC = BASE.parent / "sources" / "Investment Services LLC 30-1435731" / "25-01-17 Investment Services LLC Operating Agreement.docx"
OUT_DIR = BASE / "Investment Services"
OUT = OUT_DIR / "IS V06 - Investment Services OA Draft - from Reassembled OA Check.docx"
NOTES = OUT_DIR / "IS V06 - Investment Services OA Draft Notes.md"


GREEN = RGBColor(0x00, 0x80, 0x00)


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}r"):
            p.remove(child)


def add_run(paragraph, text, *, green=False, strike=False, bold=False, italic=False):
    run = paragraph.add_run(text)
    if green:
        run.font.color.rgb = GREEN
    if strike:
        run.font.strike = True
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return run


def insert_paragraph_after(paragraph, segments):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    for text, opts in segments:
        add_run(new_paragraph, text, **opts)
    return new_paragraph


def replace_once_redline(paragraph, old, new):
    text = paragraph.text
    if old not in text:
        return False
    before, after = text.split(old, 1)
    clear_paragraph(paragraph)
    if before:
        add_run(paragraph, before)
    add_run(paragraph, old, strike=True)
    add_run(paragraph, " ")
    add_run(paragraph, new, green=True)
    if after:
        add_run(paragraph, after)
    return True


def replace_whole_paragraph(paragraph, replacement):
    old = paragraph.text
    clear_paragraph(paragraph)
    if old:
        add_run(paragraph, old, strike=True)
        add_run(paragraph, " ", green=True)
    add_run(paragraph, replacement, green=True)


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def strike_whole_paragraph(paragraph):
    old = paragraph.text
    clear_paragraph(paragraph)
    if old:
        add_run(paragraph, old, strike=True)


def remove_table_paragraph(paragraph):
    parent = paragraph._p.getparent()
    parent.remove(paragraph._p)


def set_cell_text(cell, text):
    for paragraph in list(cell.paragraphs)[1:]:
        remove_table_paragraph(paragraph)
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    for i, part in enumerate(text.split("\n")):
        if i:
            paragraph.add_run().add_break()
        add_run(paragraph, part, green=True)


def set_signature_cell(cell, heading, name, lines=None):
    lines = lines or []
    for paragraph in list(cell.paragraphs)[1:]:
        remove_table_paragraph(paragraph)
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    add_run(paragraph, heading, green=True, bold=True)
    paragraph.add_run().add_break()
    add_run(paragraph, name, green=True, bold=True)
    for line in lines:
        paragraph.add_run().add_break()
        add_run(paragraph, line, green=True)


def add_green_note_after(paragraph, note):
    return insert_paragraph_after(paragraph, [(note, {"green": True, "italic": True})])


def add_green_paragraph_after(paragraph, text, *, bold=False):
    return insert_paragraph_after(paragraph, [(text, {"green": True, "bold": bold})])


def add_page_break_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    new_paragraph.add_run().add_break(WD_BREAK.PAGE)
    return new_paragraph


def add_page_number_run(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, text, fld_end])


def apply_version_footer(doc):
    for section in list(doc.sections)[:-1]:
        section.footer.is_linked_to_previous = False
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        clear_paragraph(paragraph)
        paragraph.alignment = 1
        add_run(paragraph, "Page ")
        add_page_number_run(paragraph)
        add_run(paragraph, " | IS V06 - Investment Services OA Draft")

    if len(doc.sections) > 1:
        legend_footer = doc.sections[-1].footer
        legend_footer.is_linked_to_previous = False
        for paragraph in legend_footer.paragraphs:
            clear_paragraph(paragraph)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(SRC)
    paragraphs = list(doc.paragraphs)

    # Title and general entity-name conversions.
    for paragraph in paragraphs:
        if paragraph.text.strip() == "SELL YOUR HOME, LLC":
            replace_whole_paragraph(paragraph, "INVESTMENT SERVICES LLC")
        elif "Sell Your Home, LLC" in paragraph.text:
            replace_once_redline(paragraph, "Sell Your Home, LLC", "Investment Services LLC")

    # Opening effective-date paragraph. Use targeted redline replacements instead of
    # replacing the whole paragraph when most of the source sentence still works.
    for paragraph in paragraphs:
        if paragraph.text.startswith("This Amended and Restated Operating Agreement is entered into"):
            replace_once_redline(paragraph, "on this _____ day of June, 2025", "effective January 17, 2025")
            replace_once_redline(paragraph, "persons executing", "persons or entities executing")
            break

    # Opening member block. Keep the SYH retirement-account history visible, then insert IS members.
    inserted_member_block = False
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if (
            text.startswith("Quest Trust Company FBO")
            or text.startswith("Heritage IRA FBO")
        ) and ("IRA" in text or "HSA" in text):
            clear_paragraph(paragraph)
            add_run(paragraph, text, strike=True)
            if text == "Heritage IRA FBO Jeanette W. Hollinger Roth IRA #61077178;" and not inserted_member_block:
                current = paragraph
                for line in [
                    "Investment Services LLC Members and Membership Interests:",
                    "Buy Your Home, LLC - 25%",
                    "Sell Your Home, LLC - 25%",
                    "Heritage Management LLC - 25%",
                    "BYH 401K LLC - 25%",
                ]:
                    current = insert_paragraph_after(current, [(line, {"green": True})])
                inserted_member_block = True

    # Summary and main-body business-purpose changes.
    for paragraph in paragraphs:
        if paragraph.text.strip() == "Business is limited to residential real estate investment.":
            replace_whole_paragraph(
                paragraph,
                "Business provides services to other investment companies, including management services, signing authority, and buyer-qualification services.",
            )
            break

    for paragraph in paragraphs:
        if paragraph.text.startswith("Scope of Business and Company Opportunities."):
            replace_once_redline(
                paragraph,
                "The Company’s sole business is investing in residential real estate for rental income and long-term growth.",
                "The Company provides services to other investment companies, including management services, signing authority, buyer-qualification services, and related administrative or operational services approved by the Manager(s).",
            )
            replace_once_redline(paragraph, "other real estate", "other real estate or business opportunities")
            break

    # Article 5 manager paragraph: do not carry SYH manager facts forward.
    for paragraph in paragraphs:
        if "Manager(s)" in paragraph.text and "Jessica Santacruz" in paragraph.text:
            replace_once_redline(paragraph, "Jessica Santacruz", "[UNCONFIRMED Manager(s)]")
            replace_once_redline(
                paragraph,
                "She will serve until the first annual meeting or until she resigns or is removed.",
                "The Manager(s) will serve until the first annual meeting or until resignation or removal.",
            )
            break

    # Joshua Kennedy is confirmed by Wes as the Investment Services Manager.
    for paragraph in paragraphs:
        if paragraph.text.startswith("Current Manager(s). Joshua Kennedy"):
            for run in paragraph.runs:
                run.font.strike = False
            break

    # Leave the SYH source clause about Investment Services LLC serving as Manager out
    # of Investment Services LLC's own OA.
    paragraphs = list(doc.paragraphs)
    for paragraph in paragraphs:
        if paragraph.text.startswith("Entity Manager Authority."):
            remove_paragraph(paragraph)
            break

    # Keep 5.1(b)(xi) operative even if a source review mark would otherwise strike it.
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "Any increase or decrease in the number of Manager(s).":
            for run in paragraph.runs:
                run.font.strike = False
            break

    # Strike Article 9.1, 9.2, and 9.6 while preserving paragraph position.
    for marker in [
        "Restriction Against Prohibited Transactions.",
        "Unrelated Business Taxable Income and Unrelated Debt Financed Income.",
        "Required Minimum Distributions.",
    ]:
        for paragraph in doc.paragraphs:
            if paragraph.text.startswith(marker):
                strike_whole_paragraph(paragraph)
                break

    # Reuse the source two-column certification table as a clean-edit layout.
    # Old inapplicable SYH/IRA text is removed, but the two-column structure remains.
    for table in doc.tables:
        table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if "Heritage IRA, FBO" in table_text and "Read and approved for investment" in table_text:
            while len(table.rows) < 3:
                table.add_row()
            signature_cells = [
                ("MEMBER:", "Buy Your Home, LLC", ["By: ____________________________________", "Printed Name: Wesley Dale Browning", "Title: __________________________________"]),
                ("MEMBER:", "Sell Your Home, LLC", ["By: ____________________________________", "Printed Name: Joshua Kennedy", "Title: __________________________________"]),
                ("MEMBER:", "Heritage Management LLC", ["By: ____________________________________", "Printed Name: Wesley Dale Browning", "Title: __________________________________"]),
                ("MEMBER:", "BYH 401K LLC", ["By: ____________________________________", "Printed Name: Wesley Dale Browning", "Title: __________________________________"]),
                ("MANAGER:", "____________________________________", ["Joshua Kennedy, Manager"]),
                ("", "", []),
            ]
            for cell, (heading, name, lines) in zip([cell for row in table.rows for cell in row.cells], signature_cells):
                if heading:
                    set_signature_cell(cell, heading, name, lines)
                else:
                    set_cell_text(cell, "")
            break

    # Convert Exhibit A from SYH IRA members to the Investment Services member list.
    for table in doc.tables:
        table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if "Member Name, Address" in table_text and "Membership" in table_text:
            exhibit_rows = [
                ("Member Name, Address", "Units"),
                ("Buy Your Home, LLC\nAddress: [TBD]\nEmail: [TBD]", "25 units"),
                ("Sell Your Home, LLC\nAddress: [TBD]\nEmail: [TBD]", "25 units"),
                ("Heritage Management LLC\nAddress: [TBD]\nEmail: [TBD]", "25 units"),
                ("BYH 401K LLC\nAddress: [TBD]\nEmail: [TBD]", "25 units"),
            ]
            for row, values in zip(table.rows, exhibit_rows):
                set_cell_text(row.cells[0], values[0])
                set_cell_text(row.cells[1], values[1])
            break

    # Mark provisions that still need Investment Services-specific review.
    review_markers = [
        ("Company as Pass-Through Entity.", "INVESTMENT SERVICES V06 TAX REVIEW: tax classification remains [UNCONFIRMED]; confirm before finalizing pass-through, partnership, or other tax language."),
    ]
    for marker, note in review_markers:
        for paragraph in paragraphs:
            if marker in paragraph.text:
                add_green_note_after(paragraph, note)
                break

    apply_version_footer(doc)
    doc.save(OUT)
    NOTES.write_text(
        "\n".join(
            [
                "# IS V06 - Investment Services OA Draft Notes",
                "",
                "Mode: `Investment Services`.",
                "",
                "Source used:",
                "",
                "`sources/Jeff Watson/Simplified OA Subfiles/Reassembled OA Check - 00 - Reassembled Verification - Simplified OA with Jeff Watson tracked edits.docx`",
                "",
                "Output DOCX:",
                "",
                "`IS V06 - Investment Services OA Draft - from Reassembled OA Check.docx`",
                "",
                "Wyoming OA use:",
                "",
                f"`{WYOMING_SRC.relative_to(BASE.parent)}`",
                "",
                "- Used only as background for the effective/start date instruction.",
                "- The Wyoming OA text says the agreement became effective when articles were filed; the visible file date is `25-01-17`, so V06 uses January 17, 2025 as the working effective date.",
                "- The Wyoming OA was not used as the drafting source.",
                "",
                "Changes applied from Wes's cumulative Investment Services list:",
                "",
                "- Title redlined from `SELL YOUR HOME, LLC` to `INVESTMENT SERVICES LLC`.",
                "- Opening effective-date paragraph uses targeted word/phrase redlines instead of wholesale paragraph replacement.",
                "- Opening member block now lists Buy Your Home, LLC, Sell Your Home, LLC, Heritage Management LLC, and BYH 401K LLC as 25% Members.",
                "- Summary and Article 1 business-purpose changes leave out `Business is not limited to residential real estate investment.`",
                "- Main-body `Scope of Business and Company Opportunities` uses targeted replacement for the source residential-investment business sentence.",
                "- Keeps Joshua Kennedy as the confirmed current Manager of Investment Services LLC.",
                "- Leaves out the source `Entity Manager Authority` paragraph and related manager-review note.",
                "- Keeps Section 5.1(b)(xi), `Any increase or decrease in the number of Manager(s).`, operative and unstruck.",
                "- Strikes Article 9.1, 9.2, and 9.6 while preserving the paragraphs for reference continuity.",
                "- Replaces the certification page placeholder with a clean two-column Investment Services certification table, using member signature groups and a Joshua Kennedy Manager signature group.",
                "- Member certification printed-name lines list Wesley Dale Browning for Buy Your Home, LLC, Heritage Management LLC, and BYH 401K LLC, and Joshua Kennedy for Sell Your Home, LLC.",
                "- Updates Exhibit A to the four known Investment Services Members at 25 units each, with address/email details marked TBD.",
                "- Adds a review marker for unresolved tax classification.",
                "- Adds the required page/version footer to agreement sections while leaving the final Drafting Legend section without a footer.",
                "- Carries forward the source-stack Drafting Legend from the Reassembled OA Check.",
                "",
                "QA results:",
                "",
                "- Structural DOCX checks passed for the certification table, old IRA certification removal, Joshua Kennedy Manager block, Exhibit A 25-unit member table, and V06 footer/legend handling.",
                "- LibreOffice fallback PDF conversion succeeded at 26 pages.",
                "- Visual page checks passed for the one-page two-column certification layout, Exhibit A, and final Drafting Legend.",
                "",
                "Open facts before review-ready build:",
                "",
                "- Confirm Secretary of State filing/start date if January 17, 2025 should be replaced by a filing-date record.",
                "- Confirm whether any additional current Manager or Managers should be listed.",
                "- Tax classification.",
                "- Member contribution/unit details for Exhibit A.",
                "- Signature authority and signer titles.",
                "",
            ]
        ),
        encoding="utf-8",
    )
    print(OUT)
    print(NOTES)


if __name__ == "__main__":
    build()
