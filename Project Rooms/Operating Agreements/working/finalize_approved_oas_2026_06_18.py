from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"C:\Codex\Wiki Files\Project Rooms\Operating Agreements")
OUTPUTS = ROOT / "outputs"

FILES = [
    {
        "source": ROOT / "working" / "SYH" / "SYH Final Candidate 01 - Sell Your Home OA.docx",
        "output_dir": OUTPUTS / "SYH",
        "output_name": "26-06-18 Sell Your Home LLC Operating Agreement - Approved Final.docx",
        "footer": "Approved Final - Sell Your Home LLC Operating Agreement",
    },
    {
        "source": ROOT / "working" / "Investment Services" / "IS Final Candidate 01 - Investment Services OA.docx",
        "output_dir": OUTPUTS / "Investment Services",
        "output_name": "26-06-18 Investment Services LLC Operating Agreement - Approved Final.docx",
        "footer": "Approved Final - Investment Services LLC Operating Agreement",
    },
]


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr)
    run._r.append(fld_char_end)


def set_footer(section, footer_text):
    section.different_first_page_header_footer = False
    section.footer.is_linked_to_previous = False
    footer = section.footer
    for paragraph in footer.paragraphs:
        paragraph.clear()
    paragraph = footer.paragraphs[0]
    paragraph.alignment = 1
    paragraph.add_run("Page ")
    add_page_field(paragraph)
    paragraph.add_run(f" | {footer_text}")


def finalize_file(spec):
    doc = Document(spec["source"])
    for section in doc.sections:
        set_footer(section, spec["footer"])
    spec["output_dir"].mkdir(parents=True, exist_ok=True)
    output_path = spec["output_dir"] / spec["output_name"]
    doc.save(output_path)
    return output_path


def main():
    for spec in FILES:
        print(finalize_file(spec))


if __name__ == "__main__":
    main()
