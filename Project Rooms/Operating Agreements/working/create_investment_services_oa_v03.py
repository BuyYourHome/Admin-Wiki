from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph


BASE = Path(__file__).resolve().parent
SRC = BASE.parent / "sources" / "Jeff Watson" / "Simplified OA Subfiles" / "Reassembled OA Check - 00 - Reassembled Verification - Simplified OA with Jeff Watson tracked edits.docx"
WYOMING_SRC = BASE.parent / "sources" / "Investment Services LLC 30-1435731" / "25-01-17 Investment Services LLC Operating Agreement.docx"
OUT_DIR = BASE / "Investment Services"
OUT = OUT_DIR / "IS V03 - Investment Services OA Draft - from Reassembled OA Check.docx"
NOTES = OUT_DIR / "IS V03 - Investment Services OA Draft Notes.md"


GREEN = RGBColor(0x00, 0x80, 0x00)


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}r"):
            p.remove(child)


def add_run(paragraph, text, *, green=False, strike=False, bold=False, italic=False):
    run = paragraph.add_run(text)
    if green:
        run.font.color.rgb = GREEN
    if strike:
        run.font.strike = True
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return run


def insert_paragraph_after(paragraph, segments):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    for text, opts in segments:
        add_run(new_paragraph, text, **opts)
    return new_paragraph


def replace_once_redline(paragraph, old, new):
    text = paragraph.text
    if old not in text:
        return False
    before, after = text.split(old, 1)
    clear_paragraph(paragraph)
    if before:
        add_run(paragraph, before)
    add_run(paragraph, old, strike=True)
    add_run(paragraph, " ")
    add_run(paragraph, new, green=True)
    if after:
        add_run(paragraph, after)
    return True


def replace_whole_paragraph(paragraph, replacement):
    old = paragraph.text
    clear_paragraph(paragraph)
    if old:
        add_run(paragraph, old, strike=True)
        add_run(paragraph, " ", green=True)
    add_run(paragraph, replacement, green=True)


def add_green_note_after(paragraph, note):
    return insert_paragraph_after(paragraph, [(note, {"green": True, "italic": True})])


def add_page_number_run(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, text, fld_end])


def apply_version_footer(doc):
    for section in list(doc.sections)[:-1]:
        section.footer.is_linked_to_previous = False
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        clear_paragraph(paragraph)
        paragraph.alignment = 1
        add_run(paragraph, "Page ")
        add_page_number_run(paragraph)
        add_run(paragraph, " | IS V03 - Investment Services OA Draft")

    if len(doc.sections) > 1:
        legend_footer = doc.sections[-1].footer
        legend_footer.is_linked_to_previous = False
        for paragraph in legend_footer.paragraphs:
            clear_paragraph(paragraph)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(SRC)
    paragraphs = list(doc.paragraphs)

    # Title and general entity-name conversions.
    for paragraph in paragraphs:
        if paragraph.text.strip() == "SELL YOUR HOME, LLC":
            replace_whole_paragraph(paragraph, "INVESTMENT SERVICES LLC")
        elif "Sell Your Home, LLC" in paragraph.text:
            replace_once_redline(paragraph, "Sell Your Home, LLC", "Investment Services LLC")

    # Opening effective-date paragraph. The start date is taken from the Wyoming OA file date,
    # while the Wyoming OA text is used only as background for the effective-date concept.
    for paragraph in paragraphs:
        if paragraph.text.startswith("This Amended and Restated Operating Agreement is entered into"):
            replace_whole_paragraph(
                paragraph,
                "This Amended and Restated Operating Agreement is entered into effective January 17, 2025, by and among the persons or entities executing this Agreement as Members, or their representatives, on the following terms and conditions.",
            )
            break

    # Opening member block. Keep the SYH retirement-account history visible, then insert IS members.
    inserted_member_block = False
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if (
            text.startswith("Quest Trust Company FBO")
            or text.startswith("Heritage IRA FBO")
        ) and ("IRA" in text or "HSA" in text):
            clear_paragraph(paragraph)
            add_run(paragraph, text, strike=True)
            if text == "Heritage IRA FBO Jeanette W. Hollinger Roth IRA #61077178;" and not inserted_member_block:
                current = paragraph
                for line in [
                    "Investment Services LLC Members and Membership Interests:",
                    "Buy Your Home, LLC - 25%",
                    "Sell Your Home, LLC - 25%",
                    "Heritage Management LLC - 25%",
                    "BYH 401K LLC - 25%",
                ]:
                    current = insert_paragraph_after(current, [(line, {"green": True})])
                inserted_member_block = True

    # Summary and main-body business-purpose changes.
    for paragraph in paragraphs:
        if paragraph.text.strip() == "Business is limited to residential real estate investment.":
            replace_whole_paragraph(
                paragraph,
                "Business is not limited to residential real estate investment. The Company provides services to other investment companies, including management services, signing authority, and buyer-qualification services.",
            )
            break

    for paragraph in paragraphs:
        if paragraph.text.startswith("Scope of Business and Company Opportunities."):
            replace_whole_paragraph(
                paragraph,
                "Scope of Business and Company Opportunities. The Company is not an investment company and its business is not limited to residential real estate investment. The Company provides services to other investment companies, including management services, signing authority, buyer-qualification services, and related administrative or operational services approved by the Manager(s). Members can invest in other real estate or business opportunities on their own, as long as it does not significantly harm the Company.",
            )
            break

    # Article 5 manager paragraph: do not carry Jessica or BYH manager facts forward.
    for paragraph in paragraphs:
        if "Manager(s)" in paragraph.text and "Jessica Santacruz" in paragraph.text:
            replace_once_redline(paragraph, "Jessica Santacruz", "[UNCONFIRMED Manager(s)]")
            replace_once_redline(
                paragraph,
                "She will serve until the first annual meeting or until she resigns or is removed.",
                "The Manager(s) will serve until the first annual meeting or until resignation or removal.",
            )
            break

    # Mark provisions that still need Investment Services-specific review.
    review_markers = [
        ("Company as Pass-Through Entity.", "INVESTMENT SERVICES V03 TAX REVIEW: tax classification remains [UNCONFIRMED]; confirm before finalizing pass-through, partnership, or other tax language."),
        ("Restriction Against Prohibited Transactions.", "INVESTMENT SERVICES V03 RETIREMENT REVIEW: applicability of IRA/ESA/HSA prohibited-transaction provisions remains [UNCONFIRMED] for Investment Services."),
        ("Unrelated Business Taxable Income", "INVESTMENT SERVICES V03 RETIREMENT/TAX REVIEW: UBTI/UDFI language requires Investment Services-specific ownership and tax review."),
        ("Required Minimum Distributions.", "INVESTMENT SERVICES V03 RETIREMENT REVIEW: RMD language requires confirmation of whether any retirement-account member exists."),
        ("Entity Manager Authority.", "INVESTMENT SERVICES V03 MANAGER REVIEW: this source clause was added for entities where Investment Services LLC may act as manager. Confirm whether and how it should apply to Investment Services LLC itself."),
    ]
    for marker, note in review_markers:
        for paragraph in paragraphs:
            if marker in paragraph.text:
                add_green_note_after(paragraph, note)
                break

    for paragraph in paragraphs:
        if paragraph.text.strip() == "CERTIFICATION":
            add_green_note_after(
                paragraph,
                "INVESTMENT SERVICES V03 CERTIFICATION NOTE: replace certification/signature blocks with Investment Services-specific Member(s), Manager(s), signer titles, and effective date after confirmation.",
            )
            break

    for paragraph in paragraphs:
        if paragraph.text.strip() == "EXHIBIT A":
            add_green_note_after(
                paragraph,
                "INVESTMENT SERVICES V03 EXHIBIT A NOTE: update Exhibit A to show Buy Your Home, LLC, Sell Your Home, LLC, Heritage Management LLC, and BYH 401K LLC as equal 25% Members, plus any confirmed contribution/unit details.",
            )
            break

    apply_version_footer(doc)
    doc.save(OUT)
    NOTES.write_text(
        "\n".join(
            [
                "# IS V03 - Investment Services OA Draft Notes",
                "",
                "Mode: `Investment Services`.",
                "",
                "Source used:",
                "",
                "`sources/Jeff Watson/Simplified OA Subfiles/Reassembled OA Check - 00 - Reassembled Verification - Simplified OA with Jeff Watson tracked edits.docx`",
                "",
                "Output DOCX:",
                "",
                "`IS V03 - Investment Services OA Draft - from Reassembled OA Check.docx`",
                "",
                "Wyoming OA use:",
                "",
                f"`{WYOMING_SRC.relative_to(BASE.parent)}`",
                "",
                "- Used only as background for the effective/start date instruction.",
                "- The Wyoming OA text says the agreement became effective when articles were filed; the visible file date is `25-01-17`, so V03 uses January 17, 2025 as the working effective date.",
                "- The Wyoming OA was not used as the drafting source.",
                "",
                "Changes applied from Wes's cumulative Investment Services list:",
                "",
                "- Title redlined from `SELL YOUR HOME, LLC` to `INVESTMENT SERVICES LLC`.",
                "- Opening effective-date paragraph redlined to January 17, 2025.",
                "- Opening member block now lists Buy Your Home, LLC, Sell Your Home, LLC, Heritage Management LLC, and BYH 401K LLC as 25% Members.",
                "- Summary business-purpose line now states Investment Services is not limited to residential real estate investment and provides services to other investment companies.",
                "- Main-body `Scope of Business and Company Opportunities` paragraph now states Investment Services is not an investment company and provides management, signing-authority, buyer-qualification, and related services.",
                "- Adds review markers for unresolved tax, retirement-account, manager-authority, certification/signature, and Exhibit A issues.",
                "- Adds the required page/version footer to agreement sections while leaving the final Drafting Legend section without a footer.",
                "- Carries forward the source-stack Drafting Legend from the Reassembled OA Check.",
                "",
                "Open facts before review-ready build:",
                "",
                "- Confirm Secretary of State filing/start date if January 17, 2025 should be replaced by a filing-date record.",
                "- Current Manager or Managers.",
                "- Tax classification.",
                "- Member contribution/unit details for Exhibit A.",
                "- Signature authority and signer titles.",
                "- Whether the SYH source clause allowing Investment Services LLC to serve as Manager should apply, be revised, or be omitted in the Investment Services LLC OA itself.",
                "",
            ]
        ),
        encoding="utf-8",
    )
    print(OUT)
    print(NOTES)


if __name__ == "__main__":
    build()
