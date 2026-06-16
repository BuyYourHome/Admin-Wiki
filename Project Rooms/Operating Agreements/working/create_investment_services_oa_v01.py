from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph


BASE = Path(__file__).resolve().parent
SRC = BASE.parent / "sources" / "Jeff Watson" / "Simplified OA Subfiles" / "Reassembled OA Check - 00 - Reassembled Verification - Simplified OA with Jeff Watson tracked edits.docx"
OUT_DIR = BASE / "Investment Services"
OUT = OUT_DIR / "IS V01 - Investment Services OA Draft - from Reassembled OA Check.docx"
NOTES = OUT_DIR / "IS V01 - Investment Services OA Draft Notes.md"


GREEN = RGBColor(0x00, 0x80, 0x00)


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}r"):
            p.remove(child)


def add_run(paragraph, text, *, green=False, strike=False, bold=False, italic=False):
    run = paragraph.add_run(text)
    if green:
        run.font.color.rgb = GREEN
    if strike:
        run.font.strike = True
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return run


def replace_once_redline(paragraph, old, new):
    text = paragraph.text
    if old not in text:
        return False
    before, after = text.split(old, 1)
    clear_paragraph(paragraph)
    if before:
        add_run(paragraph, before)
    add_run(paragraph, old, strike=True)
    add_run(paragraph, new, green=True)
    if after:
        add_run(paragraph, after)
    return True


def add_green_note_after(paragraph, note):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    add_run(new_paragraph, note, green=True, italic=True)
    return new_paragraph


def insert_paragraph_after(paragraph, segments):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    for text, opts in segments:
        add_run(new_paragraph, text, **opts)
    return new_paragraph


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(SRC)

    paragraphs = list(doc.paragraphs)

    # Entity-name conversions. Keep source text visible and add target entity in green.
    # Do not loop here: after redlining, paragraph.text still includes the struck source text.
    for paragraph in paragraphs:
        if "Sell Your Home, LLC" in paragraph.text:
            replace_once_redline(paragraph, "Sell Your Home, LLC", "Investment Services LLC")

    # Opening retirement-account member lines: strike source lines and insert an Investment Services placeholder.
    inserted_member_placeholder = False
    for idx, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if (
            text.startswith("Heritage IRA FBO")
            or text.startswith("Quest Trust Company FBO")
        ) and ("IRA" in text or "HSA" in text):
            clear_paragraph(paragraph)
            add_run(paragraph, text, strike=True)
            if not inserted_member_placeholder and idx >= 10:
                pass

    for paragraph in paragraphs:
        if paragraph.text.strip() == "Heritage IRA FBO Jeanette W. Hollinger Roth IRA #61077178;":
            insert_paragraph_after(
                paragraph,
                [
                    ("Investment Services LLC member(s): ", {"green": True, "bold": True}),
                    ("[UNCONFIRMED - confirm current Member(s), ownership percentages or units, contributions, and whether any retirement-account ownership applies.]", {"green": True}),
                ],
            )
            inserted_member_placeholder = True
            break

    # Article 5 manager paragraph: do not carry Jessica or BYH manager facts forward.
    for paragraph in paragraphs:
        if "Manager(s) – Number and Tenure." in paragraph.text and "Jessica" in paragraph.text:
            replace_once_redline(paragraph, "Jessica Santacruz", "[UNCONFIRMED Manager(s)]")
            replace_once_redline(
                paragraph,
                "She will serve until the first annual meeting or until she resigns or is removed.",
                "The Manager(s) will serve until the first annual meeting or until resignation or removal.",
            )
            break

    # Mark retirement-account and tax provisions that need Investment Services review.
    review_markers = [
        ("Company as Pass-Through Entity.", "INVESTMENT SERVICES V01 TAX REVIEW: tax classification is [UNCONFIRMED]; do not assume BYH S-corp treatment or SYH retirement-account pass-through language without CPA/attorney review."),
        ("Restriction Against Prohibited Transactions.", "INVESTMENT SERVICES V01 RETIREMENT REVIEW: applicability of IRA/ESA/HSA prohibited-transaction provisions is [UNCONFIRMED] for Investment Services."),
        ("Unrelated Business Taxable Income", "INVESTMENT SERVICES V01 RETIREMENT/TAX REVIEW: UBTI/UDFI language requires Investment Services-specific ownership and tax review."),
        ("Required Minimum Distributions.", "INVESTMENT SERVICES V01 RETIREMENT REVIEW: RMD language requires confirmation of whether retirement-account members exist."),
    ]
    for marker, note in review_markers:
        for paragraph in paragraphs:
            if marker in paragraph.text:
                add_green_note_after(paragraph, note)
                break

    # Certification and Exhibit A need clean replacement once facts are known.
    for paragraph in paragraphs:
        if paragraph.text.strip() == "CERTIFICATION":
            add_green_note_after(
                paragraph,
                "INVESTMENT SERVICES V01 CERTIFICATION NOTE: replace certification/signature blocks with Investment Services-specific Member(s), Manager(s), signer titles, and effective date after confirmation.",
            )
            break

    for paragraph in paragraphs:
        if paragraph.text.strip() == "EXHIBIT A":
            add_green_note_after(
                paragraph,
                "INVESTMENT SERVICES V01 EXHIBIT A NOTE: replace Exhibit A with clean Investment Services member, ownership, unit/percentage, and contribution information after confirmation.",
            )
            break

    doc.save(OUT)

    NOTES.write_text(
        "\n".join(
            [
                "# IS V01 - Investment Services OA Draft Notes",
                "",
                "Mode: `Investment Services`.",
                "",
                "Source used:",
                "",
                "`sources/Jeff Watson/Simplified OA Subfiles/Reassembled OA Check - 00 - Reassembled Verification - Simplified OA with Jeff Watson tracked edits.docx`",
                "",
                "Output DOCX:",
                "",
                "`IS V01 - Investment Services OA Draft - from Reassembled OA Check.docx`",
                "",
                "Build method:",
                "",
                "- Generated from the Reassembled OA Check / SYH framework.",
                "- Follows the project-room rule that new OA drafts using the SYH/Simplified OA framework begin from the Reassembled OA Check unless Wes explicitly names another source.",
                "- Did not use the Wyoming Investment Services OA as the drafting source.",
                "- Used BYH-like process discipline but did not copy BYH entity facts.",
                "- Source refresh from Teams/SharePoint was not available in this local tool session; the local Reassembled OA Check source was used.",
                "",
                "First-pass changes:",
                "",
                "- Redlined `Sell Your Home, LLC` to `Investment Services LLC` where present.",
                "- Struck opening retirement-account member lines and inserted an Investment Services member placeholder.",
                "- Redlined the initial Manager name to `[UNCONFIRMED Manager(s)]` and changed the immediate pronoun sentence to Manager(s)-neutral language.",
                "- Added green Investment Services review notes for tax classification, retirement-account provisions, certification/signature blocks, and Exhibit A.",
                "",
                "Open facts before next build:",
                "",
                "- Current Member or Members.",
                "- Membership percentages or units.",
                "- Manager or Managers.",
                "- Tax classification.",
                "- Effective date.",
                "- Signature authority and signer titles.",
                "- Exhibit A content.",
                "",
            ]
        ),
        encoding="utf-8",
    )

    print(OUT)
    print(NOTES)


if __name__ == "__main__":
    main()
