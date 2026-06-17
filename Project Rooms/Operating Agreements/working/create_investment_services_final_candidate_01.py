from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


ROOT = Path(r"C:\Codex\Wiki Files\Project Rooms\Operating Agreements")
MODE_DIR = ROOT / "working" / "Investment Services"
SOURCE = MODE_DIR / "IS V06 - Investment Services OA Draft - from Reassembled OA Check.docx"
OUTPUT = MODE_DIR / "IS Final Candidate 01 - Investment Services OA.docx"
NOTES = MODE_DIR / "IS Final Candidate 01 - Investment Services OA Notes.md"

FINAL_FOOTER = "IS Final Candidate 01 - Investment Services OA"
PRESERVED_STRUCK_HEADINGS = (
    "Restriction Against Prohibited Transactions.",
    "Unrelated Business Taxable Income and Unrelated Debt Financed Income.",
    "Required Minimum Distributions.",
)
REMOVE_MARKERS = (
    "INVESTMENT SERVICES V06 TAX REVIEW:",
    "[UNCONFIRMED]",
    "Drafting Legend",
    "This drafting legend is not part of the Operating Agreement.",
    "Black text reflects Jeff Watson source text",
    "Blue text reflects Wes Browning simplified replacement language",
    "Red text reflects Jeff Watson review edits",
    "Green text reflects Codex/Wes derivative changes",
)


def remove_element(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_paragraph(paragraph):
    remove_element(paragraph._p)


def paragraph_has_section_break(paragraph):
    return bool(paragraph._p.xpath(".//w:sectPr"))


def is_preserved_struck_clause(text):
    return any(text.startswith(heading) for heading in PRESERVED_STRUCK_HEADINGS)


def is_numbered_paragraph(paragraph):
    return bool(paragraph._p.xpath("./w:pPr/w:numPr"))


def remove_run(run):
    remove_element(run._r)


def remove_run_color(run):
    rpr = run._r.get_or_add_rPr()
    for child in list(rpr):
        if child.tag == qn("w:color"):
            rpr.remove(child)


def force_black(run):
    remove_run_color(run)
    run.font.color.rgb = RGBColor(0, 0, 0)


def clean_runs(paragraph):
    text = paragraph.text.strip()
    preserve_struck = is_preserved_struck_clause(text) or is_numbered_paragraph(paragraph)

    for run in list(paragraph.runs):
        if preserve_struck:
            run.font.hidden = False
            run.font.strike = bool(run.font.strike)
            force_black(run)
            continue

        if run.font.strike:
            # Physically deleting numbered paragraphs can make Word/LibreOffice
            # recalculate downstream Article numbering. Hide rejected source text
            # so the visible/signing copy is clean while numbering remains stable.
            run.font.hidden = True
            run.font.strike = False
            force_black(run)
            continue

        run.font.hidden = False
        run.font.strike = False
        force_black(run)


def clean_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in list(cell.paragraphs):
                if any(marker in paragraph.text for marker in REMOVE_MARKERS):
                    remove_paragraph(paragraph)
                    continue
                clean_runs(paragraph)
            for nested in cell.tables:
                clean_table(nested)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr)
    run._r.append(fld_char_end)


def set_footer(section):
    section.different_first_page_header_footer = False
    section.footer.is_linked_to_previous = False
    footer = section.footer
    for paragraph in footer.paragraphs:
        paragraph.clear()
    paragraph = footer.paragraphs[0]
    paragraph.alignment = 1
    paragraph.add_run("Page ")
    add_page_field(paragraph)
    paragraph.add_run(f" | {FINAL_FOOTER}")
    for run in paragraph.runs:
        force_black(run)


def remove_legend_section(doc):
    legend_index = None
    paragraphs = list(doc.paragraphs)
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == "Drafting Legend":
            legend_index = index
            break
    if legend_index is None:
        return

    start = legend_index
    if legend_index > 0 and paragraph_has_section_break(paragraphs[legend_index - 1]):
        start = legend_index - 1

    for paragraph in paragraphs[start:]:
        remove_paragraph(paragraph)


def main():
    doc = Document(SOURCE)
    remove_legend_section(doc)

    for paragraph in list(doc.paragraphs):
        if any(marker in paragraph.text for marker in REMOVE_MARKERS):
            remove_paragraph(paragraph)
            continue
        clean_runs(paragraph)

    for table in doc.tables:
        clean_table(table)

    for section in doc.sections:
        set_footer(section)

    doc.save(OUTPUT)

    NOTES.write_text(
        "\n".join(
            [
                "# IS Final Candidate 01 - Investment Services OA Notes",
                "",
                f"Source draft: `{SOURCE.name}`.",
                "",
                "Final-candidate cleanup applied:",
                "- Removed the separate Drafting Legend section from the signing candidate.",
                "- Removed Investment Services draft review markers.",
                "- Accepted partial word/phrase redlines into clean text only outside numbered paragraphs and approved clean-edit areas.",
                "- Converted surviving candidate text to black/default final text.",
                "- Preserved numbered paragraphs and numbered paragraph positions; unwanted numbered text remains visibly struck rather than deleted or hidden.",
                "- Kept the two-column certification format and the current member/manager signature names from IS V06.",
                "",
                "Approval status: not approved final. Do not copy to Teams until Wes approves this candidate.",
                "",
            ]
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
