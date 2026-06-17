from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"C:\Codex\Wiki Files\Project Rooms\Operating Agreements")
MODE_DIR = ROOT / "working" / "SYH"
SOURCE = MODE_DIR / "SYH V01 - Sell Your Home OA Draft - from Reassembled OA Check.docx"
OUTPUT = MODE_DIR / "SYH Final Candidate 01 - Sell Your Home OA.docx"
NOTES = MODE_DIR / "SYH Final Candidate 01 - Sell Your Home OA Notes.md"

FINAL_FOOTER = "SYH Final Candidate 01 - Sell Your Home OA"


def remove_element(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_paragraph(paragraph):
    remove_element(paragraph._p)


def paragraph_has_section_break(paragraph):
    return bool(paragraph._p.xpath(".//w:sectPr"))


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr)
    run._r.append(fld_char_end)


def set_footer(section):
    section.different_first_page_header_footer = False
    section.footer.is_linked_to_previous = False
    footer = section.footer
    for paragraph in footer.paragraphs:
        paragraph.clear()
    paragraph = footer.paragraphs[0]
    paragraph.alignment = 1
    paragraph.add_run("Page ")
    add_page_field(paragraph)
    paragraph.add_run(f" | {FINAL_FOOTER}")


def remove_legend_section(doc):
    paragraphs = list(doc.paragraphs)
    legend_index = None
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == "Drafting Legend":
            legend_index = index
            break
    if legend_index is None:
        return

    start = legend_index
    if legend_index > 0 and paragraph_has_section_break(paragraphs[legend_index - 1]):
        start = legend_index - 1

    for paragraph in paragraphs[start:]:
        remove_paragraph(paragraph)


def main():
    doc = Document(SOURCE)
    remove_legend_section(doc)
    for section in doc.sections:
        set_footer(section)
    doc.save(OUTPUT)

    NOTES.write_text(
        "\n".join(
            [
                "# SYH Final Candidate 01 - Sell Your Home OA Notes",
                "",
                f"Source draft: `{SOURCE.name}`.",
                "",
                "Final-candidate treatment:",
                "- Built from the current SYH V01 draft, which was built from the Reassembled OA Check.",
                "- Removed the separate non-operative Drafting Legend section.",
                "- Preserved agreement text, source colors, formatting, tables, numbering, and certification layout.",
                "- Applied the final-candidate footer to the remaining agreement section.",
                "- No numbered paragraphs were deleted or hidden.",
                "",
                "Approval status: not approved final. Do not copy to Teams until Wes approves this candidate.",
                "",
            ]
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
