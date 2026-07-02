from __future__ import annotations

import re
import shutil
import subprocess
import zipfile
import os
import argparse
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook


WORKBOOK = Path(
    os.environ.get(
        "TENSITY_WORKBOOK",
        r"C:\Users\wesbr\Buy Your Home\Buy Your Home - Property\24_Project Management - 4121 Tensity Dr 2.xlsm",
    )
)
OUTPUT_DIR = Path(r"C:\Codex\Wiki Files\Project Rooms\Property Trade Evaluation\outputs")
TEAMS_PACKAGE = Path(r"C:\Users\wesbr\Buy Your Home\Buy Your Home - Property\24-HM-4121 Tensity Dr\Selling\Contract Package")
ARCHIVE_DIR = TEAMS_PACKAGE / "Archive" / "Contract"
SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")

CFD_DOC_NAME = "4121 Tensity Dr - Contract for Deed Term Sheet - DRAFT.docx"
PS_DOC_NAME = "4121 Tensity Dr - Purchase and Sale Term Sheet - DRAFT.docx"
PS_TRADE_EQUITY_DOC_NAME = "4121 Tensity Dr - Purchase and Sale Term Sheet with Trade Property Credit - DRAFT.docx"
RUN_DATE = datetime.now()
VERSION_DATE = RUN_DATE.strftime("%y-%m-%d")
PREPARED_DATE = RUN_DATE.strftime("%B %#d, %Y")


def clean(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    return str(value).strip()


def money(value) -> str:
    try:
        return f"${float(value):,.2f}"
    except (TypeError, ValueError):
        return "[not listed]"


def numeric(value) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return 0.0


def positive_money(value) -> bool:
    return numeric(value) > 0.005


def first_present(*values):
    for value in values:
        if value is not None and clean(value) != "":
            return value
    return None


def percent(value) -> str:
    try:
        value = float(value)
    except (TypeError, ValueError):
        return clean(value) or "[not listed]"
    return f"{value:.3%}" if abs(value) < 1 else f"{value:.3f}%"


def date_text(value) -> str:
    if hasattr(value, "strftime"):
        return value.strftime("%B %#d, %Y")
    return clean(value) or "[not listed]"


def normalize_trust_date(date_value, year_value) -> str:
    date_part = clean(date_value)
    year_part = clean(year_value)
    date_part = re.sub(r"^([A-Za-z]+),\s*", r"\1 ", date_part)
    date_part = re.sub(r"\s+", " ", date_part).strip()
    if not date_part:
        return ""
    if year_part and not re.search(r"\b(?:19|20)\d{2}\b", date_part):
        return f"{date_part}, {year_part}"
    date_part = re.sub(r"\s+((?:19|20)\d{2})\b", r", \1", date_part)
    return date_part


def clean_trust_name(trust_value: str) -> str:
    trust_text = clean(trust_value)
    return re.split(r"\s+dated\s+", trust_text, maxsplit=1, flags=re.IGNORECASE)[0].strip()


def docx_text(path: Path) -> str:
    if not path.exists() or path.suffix.lower() != ".docx":
        return ""
    try:
        with zipfile.ZipFile(path) as zf:
            parts = []
            for name in zf.namelist():
                if name.startswith("word/") and name.endswith(".xml"):
                    parts.append(zf.read(name).decode("utf-8", errors="ignore"))
            return "\n".join(parts)
    except zipfile.BadZipFile:
        return ""


def output_names(mode: str, trade_credit=None) -> tuple[str, str]:
    if mode == "ps-trade-equity":
        doc_name = PS_TRADE_EQUITY_DOC_NAME if positive_money(trade_credit) else PS_DOC_NAME
    else:
        doc_name = CFD_DOC_NAME
    return doc_name, doc_name.replace(".docx", ".pdf")


def resolve_mode(requested_mode: str, seller_financed, trade_credit, buyer_outside_financing=None) -> str:
    if requested_mode != "auto":
        return requested_mode
    if positive_money(seller_financed):
        return "cfd"
    if positive_money(trade_credit) or positive_money(buyer_outside_financing):
        return "ps-trade-equity"
    raise ValueError("Auto mode could not choose a term sheet type: no seller financing or trade credit was found.")


def next_package_version(doc_name: str) -> str:
    pattern = re.compile(rf"{re.escape(VERSION_DATE)} V(\d+)")
    candidates = [
        OUTPUT_DIR / doc_name,
        TEAMS_PACKAGE / doc_name,
        *ARCHIVE_DIR.glob(f"v?? - {doc_name}"),
    ]
    highest = 0
    for path in candidates:
        for match in pattern.finditer(docx_text(path)):
            highest = max(highest, int(match.group(1)))
    return f"{VERSION_DATE} V{highest + 1}"


def looks_like_address(value) -> bool:
    text = clean(value).lower()
    return bool(re.search(r"\d", text)) and any(
        marker in text
        for marker in (" st", " street", " dr", " drive", " ave", " avenue", " rd", " road", " ln", " lane", " unit")
    )


def docs_values():
    wb = load_workbook(WORKBOOK, data_only=True, read_only=True, keep_vba=True)
    ws = wb["Docs"]
    values = {}
    for row in ws.iter_rows(values_only=True):
        row = list(row)
        for idx, key in enumerate(row[:-1]):
            if key in (None, ""):
                continue
            label = str(key).strip()
            val = row[idx + 1]
            if label in values:
                suffix = 2
                while f"{label}__{suffix}" in values:
                    suffix += 1
                values[f"{label}__{suffix}"] = val
            else:
                values[label] = val
    return values


def trade_properties():
    wb = load_workbook(WORKBOOK, data_only=True, read_only=True, keep_vba=True)
    ws = wb["Trade Properties"]
    headers = [clean(ws.cell(13, c).value) for c in range(1, 16)]
    rows = []
    for r in range(14, 17):
        record = {headers[c - 1]: ws.cell(r, c).value for c in range(1, 16)}
        if record.get("Include in Contract?") is True:
            rows.append(record)
    return rows


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_text, fld_end])


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_nowrap(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    no_wrap = OxmlElement("w:noWrap")
    tc_pr.append(no_wrap)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_footer(section, version: str):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.text = ""
    p.add_run("Page ")
    add_field(p, "PAGE")
    p.add_run(" of ")
    add_field(p, "NUMPAGES")
    p.add_run(f" | {version}")
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(90, 90, 90)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.style = "Heading 2"
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p


def join_parts(parts):
    parts = [part for part in parts if part]
    if len(parts) <= 1:
        return "".join(parts)
    if len(parts) == 2:
        return " and ".join(parts)
    return ", ".join(parts[:-1]) + ", and " + parts[-1]


def add_terms_table(doc, rows, widths=(2.25, 4.9), font_size=9.5):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(rows[0]):
        cell = hdr.cells[i]
        set_cell_shading(cell, "D9EAF7")
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_width(cell, Inches(widths[i]))
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_width(cells[i], Inches(widths[i]))
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(clean(text))
            run.font.size = Pt(font_size)
            if i == 0:
                run.bold = True
    return table


def build_docx(version: str, mode: str, doc_name: str):
    v = docs_values()
    trades = trade_properties()

    buyer2 = clean(v.get("Selling-Buyer2"))
    buyer_address_parts = []
    if looks_like_address(buyer2):
        buyer_address_parts.append(buyer2)
    buyer_address_parts.extend(
        part
        for part in [clean(v.get("Selling-Buyer Add1")), clean(v.get("Selling-Buyer Add1__2")), clean(v.get("Selling-Buyer Add2"))]
        if part
    )

    seller = clean(v.get("Selling -Seller")) or clean(v.get("SellingSeller"))
    trust = clean(v.get("Trust"))
    trust_name = clean_trust_name(trust)
    trust_date = normalize_trust_date(v.get("Trust Date"), v.get("Year"))
    trust_display = f"{trust_name} dated {trust_date}" if trust_name and trust_date else trust
    trustee = clean(v.get("Trustee"))
    manager = clean(v.get("Manger") or v.get("Manager"))
    trustee_address = ", ".join(
        part
        for part in [
            clean(v.get("Trustee Address1:")),
            clean(v.get("Trustee Address2:") or v.get("Trustee Address2")),
        ]
        if part
    )
    property_line = ", ".join(part for part in [clean(v.get("Address")), clean(v.get("City-State"))] if part)
    property_short = clean(v.get("Address")) or "4121 Tensity Dr"

    purchase_price = v.get("Selling Purchase Price:")
    earnest = v.get("Selling Earnest Money:")
    cash_due = v.get("Selling Cash Due at Closing:")
    trade_credit = v.get("Selling Trade Credit Total:")
    total_credits = v.get("Selling Total Down Payment/Credits:")
    refinance_first_mortgage = first_present(
        v.get("Selling Refinance 1st Mortgage:"),
        v.get("Selling Refinance 1st Mortgage"),
    )
    seller_financed = first_present(v.get("Selling Seller-Financed Principal:"), v.get("Selling Note Principal Sum:"))
    note_payment = v.get("Selling Note Normal Payment:")
    is_ps_trade_equity = mode == "ps-trade-equity"
    purchase_financing = seller_financed
    has_purchase_financing = positive_money(purchase_financing)
    has_seller_financing = positive_money(seller_financed)
    has_trade_credit = positive_money(trade_credit) or bool(trades)
    has_refinance_first_mortgage = positive_money(refinance_first_mortgage)
    active_buyer_funding_total = refinance_first_mortgage if has_refinance_first_mortgage else total_credits
    buyer_closing_funds = (
        max(0.0, numeric(refinance_first_mortgage) - numeric(earnest) - numeric(trade_credit))
        if has_refinance_first_mortgage
        else cash_due
    )
    buyer_closing_funds_label = (
        "refinance first-mortgage proceeds brought to closing"
        if has_refinance_first_mortgage
        else "cash due at closing"
    )
    buyer_closing_funds_row_label = (
        "Refinance first-mortgage proceeds brought to closing"
        if has_refinance_first_mortgage
        else "Cash due at closing"
    )
    active_buyer_funding_row_label = (
        "Total refinance first-mortgage proceeds / credits"
        if has_refinance_first_mortgage
        else "Total cash and credits"
    )
    financing_label = (
        "seller-financed balance"
        if is_ps_trade_equity
        else "seller financing"
    )
    purchase_financing_row_label = "Seller-financed balance" if is_ps_trade_equity else "Amount financed by seller"
    primary_agreement = "Purchase and Sale Agreement" if is_ps_trade_equity else "Contract for Deed"
    term_sheet_title = (
        ("Purchase and Sale Term Sheet with Trade Property Credit" if has_trade_credit else "Purchase and Sale Term Sheet")
        if is_ps_trade_equity
        else "Contract for Deed Term Sheet"
    )
    term_sheet_subtitle = (
        f"Proposed Purchase and Sale Agreement - {property_line}"
        if is_ps_trade_equity
        else f"Proposed Contract for Deed Sale - {property_line}"
    )
    trade_property_phrase = "included trade properties" if len(trades) != 1 else "included trade property"
    purchase_components = [
        f"{money(earnest)} earnest money" if positive_money(earnest) else "",
        f"{money(buyer_closing_funds)} {buyer_closing_funds_label}" if positive_money(buyer_closing_funds) else "",
        f"{money(trade_credit)} credit for {trade_property_phrase}" if positive_money(trade_credit) else "",
        f"{money(purchase_financing)} {financing_label}" if has_purchase_financing else "",
    ]
    purchase_component_text = join_parts(purchase_components)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    add_footer(section, version)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(term_sheet_title)
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(term_sheet_subtitle).bold = True

    prepared = doc.add_paragraph()
    prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prepared.add_run(f"Prepared: {PREPARED_DATE} | Draft version: {version}")

    status = doc.add_table(rows=1, cols=1)
    status.style = "Table Grid"
    set_cell_shading(status.cell(0, 0), "FFF2CC")
    if is_ps_trade_equity:
        supporting_docs = (
            "trade-property addendum, ownership transfer paperwork, "
            if has_trade_credit
            else "buyer financing or lender paperwork, "
        )
        purpose_text = (
            "Purpose: This term sheet is a plain-English summary of the proposed purchase-and-sale deal for "
            "review and discussion. It is not the final contract and does not transfer title. The final Purchase "
            f"and Sale Agreement, deed, {supporting_docs}buyer qualification review, and attorney review still "
            "control. If this summary conflicts with the final signed documents, the final signed documents control."
        )
    else:
        purpose_text = (
            "Purpose: This term sheet is a plain-English summary of the proposed deal for review and discussion. "
            "It is not the final contract and does not transfer title. The final Contract for Deed, any needed note, "
            "memorandum, trade-property addendum, title or ownership paperwork, buyer qualification review, and "
            "attorney review still control. If this summary conflicts with the final signed documents, the final signed "
            "documents control."
        )
    status.cell(0, 0).paragraphs[0].add_run(purpose_text).italic = True

    section_no = 1

    def add_numbered_heading(text):
        nonlocal section_no
        heading = add_heading(doc, f"{section_no}. {text}")
        section_no += 1
        return heading

    add_numbered_heading("Plain-Language Deal Summary")
    add_body(
        doc,
        f"{clean(v.get('Selling-Buyer1'))} would purchase {property_short} under a proposed {primary_agreement} "
        f"for {money(purchase_price)}. The proposed deal is made up of {purchase_component_text}.",
    )

    add_numbered_heading("Parties and Property")
    add_terms_table(
        doc,
        [
            ["Item", "Details"],
            ["Seller", seller],
            ["Seller signature structure for final documents", f"{trust_display}, by and through {trustee}, Trustee; Manager: {manager}"],
            ["Trustee mailing address", trustee_address],
            ["Buyer", clean(v.get("Selling Note Maker (Buyer):")) or clean(v.get("Selling-Buyer1"))],
            ["Buyer mailing address", ", ".join(buyer_address_parts) or "[not listed]"],
            ["Property", property_line],
            ["County", clean(v.get("County")) or "[not listed]"],
            ["Parcel ID", clean(v.get("Property Parcel ID")) or "[not listed]"],
            ["Legal description summary", clean(v.get("Brief Legal Description")) or "[not listed]"],
        ],
    )

    add_numbered_heading("Proposed Purchase Structure")
    structure_parts = ["the earnest money", buyer_closing_funds_label]
    if has_trade_credit:
        structure_parts.append(f"credit for {trade_property_phrase}")
    if has_purchase_financing:
        structure_parts.append(
            "the seller-financed balance"
            if is_ps_trade_equity
            else "the amount financed by the seller"
        )
    add_body(
        doc,
        f"The proposed {primary_agreement} separates {join_parts(structure_parts)} so each part can be shown "
        "clearly in the final documents.",
    )
    purchase_rows = [
        ["Item", "Details"],
        ["Proposed purchase price", money(purchase_price)],
        ["Earnest money at signing", money(earnest)],
        [buyer_closing_funds_row_label, money(buyer_closing_funds)],
        [active_buyer_funding_row_label, money(active_buyer_funding_total)],
    ]
    if has_trade_credit:
        purchase_rows.insert(4, [f"Credit for {trade_property_phrase}", money(trade_credit)])
    if has_purchase_financing:
        purchase_rows.append(
            [
                purchase_financing_row_label,
                money(purchase_financing),
            ]
        )
    allocation = " + ".join(
        part
        for part in [
            f"{money(earnest)} earnest money" if positive_money(earnest) else "",
            f"{money(buyer_closing_funds)} {buyer_closing_funds_label}" if positive_money(buyer_closing_funds) else "",
            f"{money(trade_credit)} trade-property credit" if positive_money(trade_credit) else "",
            f"{money(purchase_financing)} {financing_label}" if has_purchase_financing else "",
        ]
        if part
    )
    purchase_rows.append(["How the purchase price is made up", f"{allocation} = {money(purchase_price)} proposed purchase price"])
    add_terms_table(
        doc,
        purchase_rows,
    )

    if has_trade_credit:
        add_numbered_heading("Trade Properties Included as Equity Credit" if is_ps_trade_equity else "Trade Properties Included as Credit")
        add_body(
            doc,
            "The buyer is expected to transfer title, ownership rights, or other required ownership documents for "
            f"the {trade_property_phrase} below as part of the deal. The credit is not final until the required "
            "paperwork, approvals, possession items, and other closing conditions are completed and accepted.",
        )
        trade_rows = [["Property", "Type", "Credit", "What still has to happen"]]
        for item in trades:
            trade_rows.append(
                [
                    clean(item.get("Property Address")),
                    clean(item.get("Home Type")),
                    money(item.get("Net Trade Credit")),
                    "Required before final credit is counted",
                ]
            )
        trade_rows.append(["Total credit", "", money(trade_credit), "Subject to trade-property addendum"])
        table = add_terms_table(doc, trade_rows, widths=(2.1, 1.0, 1.05, 3.0), font_size=8.5)
        for row in table.rows:
            for cell in row.cells:
                set_cell_nowrap(cell)
        for cell in table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for cell in table.rows[-1].cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True

        add_numbered_heading("What Must Happen for the Trade Credit to Count")
        for text in [
            "The trade credit is conditional. Before it counts as final, the buyer must deliver acceptable ownership documents, lien releases or payoff proof, possession/keys, tenant and rent information where applicable, and required third-party approvals.",
            "The final trade-property addendum should spell out ownership transfer, possession, approvals, liens, tenant rents, verification, condition, and remedies.",
            "If a required item is not delivered and accepted, the final documents should say whether the missing credit becomes additional cash due, a purchase-price adjustment, a closing-condition issue, or another agreed remedy. The proposed values are working numbers until accepted.",
        ]:
            add_body(doc, text)

    if has_seller_financing:
        add_numbered_heading("Seller-Financed Balance")
        seller_financing_body = (
            "The seller-financed balance is expected to be shown in a note signed by the buyer as part of the "
            "purchase-and-sale closing package, unless the attorney recommends a different final structure."
            if is_ps_trade_equity
            else "The seller-financed balance is expected to be shown in a note signed by the buyer and carried under the Contract for Deed, unless the attorney recommends a different final structure."
        )
        add_body(doc, seller_financing_body)
        add_terms_table(
            doc,
            [
                ["Item", "Details"],
                ["Buyer signing the note", clean(v.get("Selling Note Maker (Buyer):"))],
                ["Seller receiving payments", clean(v.get("Selling Note Payee (Seller):"))],
                ["Purpose of the note", clean(v.get("Selling Note Purpose/Direction:"))],
                ["Principal sum", money(seller_financed)],
                ["Interest rate", percent(v.get("Selling Note Interest Rate:"))],
                ["Principal and interest payment", money(note_payment)],
                ["Installment months", clean(v.get("Selling Note Installment Months:"))],
                ["First payment date", date_text(v.get("Selling Note First Payment Date:"))],
                ["Final scheduled payment date", date_text(v.get("Selling Note Final Scheduled Payment Date:"))],
                ["Balloon payment", money(v.get("Selling Note Balloon Payment:"))],
                ["Balloon wording", clean(v.get("Selling Note Balloon Term:")) or "[not listed]"],
            ],
        )

        add_numbered_heading("Estimated Monthly Payment")
        tax = v.get("Tax Escrow:")
        insurance = v.get("Property Insurance:")
        hoa = v.get("HOA:")
        total_payment = numeric(note_payment) + numeric(tax) + numeric(insurance) + numeric(hoa)
        add_terms_table(
            doc,
            [
                ["Item", "Details"],
                ["Principal and interest payment", money(note_payment)],
                ["Property tax escrow", money(tax)],
                ["Insurance escrow", money(insurance)],
                ["HOA", money(hoa)],
                ["Estimated total monthly payment if escrows are collected", money(total_payment)],
            ],
        )

    add_numbered_heading("Documents Expected Later")
    if is_ps_trade_equity:
        expected_docs = [
            f"Purchase and Sale Agreement for {property_short}.",
            "Deed from seller to buyer to be delivered at closing, subject to attorney-approved final documents.",
            "Buyer refinance / first-mortgage documents required by the buyer's lender, if applicable.",
            "Buyer and seller closing statement showing earnest money, refinance proceeds brought to closing, seller financing, prorations, and other closing adjustments.",
            "Any attorney-required disclosures, acknowledgments, or closing affidavits.",
        ]
        if has_seller_financing:
            expected_docs.insert(
                3,
                f"Promissory Note from buyer to seller showing {money(seller_financed)} principal, {percent(v.get('Selling Note Interest Rate:'))} interest, and {money(note_payment)} monthly principal-and-interest payment, unless counsel directs a different final structure.",
            )
            expected_docs.insert(
                4,
                "Seller-financing security document, deed of trust, or other attorney-approved collateral document if required for the seller-financed balance.",
            )
        if has_trade_credit:
            expected_docs.insert(2, f"Trade-property addendum or exhibit for {trade_property_phrase}.")
            expected_docs.insert(
                3,
                "Ownership/title transfer documents for the trade properties, including any required title, bill of sale, deed, assignment, park approval, or similar transfer document based on property type.",
            )
            expected_docs.insert(
                4,
                "Lien releases, payoff confirmations, or other proof that trade-property liens and adverse interests are resolved as required by the final agreement.",
            )
    else:
        expected_docs = [
            f"Contract for Deed Agreement for {property_short}.",
            "Memorandum of Contract for Deed for recording, if attorney-approved.",
            f"Trade-property addendum for {trade_property_phrase}.",
            "Buyer acknowledgment addendum and any required disclosures.",
        ]
        if has_seller_financing:
            expected_docs.insert(
                1,
                f"Promissory Note from buyer to seller showing {money(seller_financed)} principal, {percent(v.get('Selling Note Interest Rate:'))} interest, and {money(note_payment)} monthly principal-and-interest payment, unless counsel directs a different final structure.",
            )
    for text in expected_docs:
        add_body(doc, text)

    add_numbered_heading("Confirmation Status")
    status_rows = [["Status", "Item", "Notes"]]
    if not clean(v.get("Property Parcel ID")):
        status_rows.append(["Needs attention", "Parcel ID", "Parcel ID is not filled in yet."])
    if not clean(v.get("Brief Legal Description")) and not clean(v.get("Legal Description")):
        status_rows.append(["Needs attention", "Legal description", "Legal description is not filled in yet."])
    buyer_address = ", ".join(buyer_address_parts)
    if buyer_address:
        status_rows.append(["Resolved", "Buyer mailing address", f"Listed as {buyer_address}."])
    else:
        status_rows.append(["Needs attention", "Buyer mailing address", "Buyer mailing address is not filled in yet."])
    trust_text = f"{trust_display}, by and through {trustee}, Trustee; Manager: {manager}"
    trust_date_resolved = bool(
        trust_display
        and trustee
        and manager
        and re.search(r"\bdated\b", trust_display, flags=re.IGNORECASE)
        and re.search(r"\b[A-Za-z]+ \d{1,2}, (?:19|20)\d{2}\b", trust_display)
    )
    if trust_date_resolved:
        status_rows.append(["Resolved", "Trust date wording", f"Resolves as: {trust_text}."])
    else:
        status_rows.append(
            [
                "Needs attention",
                "Trust date wording",
                "Trust name, date, trustee, or manager wording is incomplete or not cleanly resolved.",
            ]
        )
    if has_trade_credit:
        status_rows.append(
            [
                "Needs attention",
                "Trade-property conditions",
                "Ownership documents, lien status, required third-party approvals, tenant/rent information where applicable, possession/keys, and condition acceptance are required before trade credit should be treated as final.",
            ]
        )
    status_rows.append(
        [
            "Needs attention",
            "Attorney review",
            (
                (
                    "Attorney review should confirm whether the purchase-and-sale structure, deed, buyer refinance / first-mortgage funding, seller-financing terms, and trade-property addendum are the preferred final structure."
                    if has_trade_credit
                    else "Attorney review should confirm whether the purchase-and-sale structure, deed, buyer refinance / first-mortgage funding, and seller-financing terms are the preferred final structure."
                )
                if is_ps_trade_equity
                else "Attorney review should confirm whether the buyer note, contract for deed, memorandum, and trade-property addendum are the preferred final structure."
                if has_seller_financing
                else "Attorney review should confirm whether the no-seller-financing structure, contract for deed, memorandum, and trade-property addendum are the preferred final structure."
            ),
        ]
    )
    add_terms_table(doc, status_rows, widths=(1.25, 1.95, 3.95), font_size=8.8)

    add_numbered_heading("Review Acknowledgment")
    add_body(doc, "The parties may use this term sheet to discuss the proposed deal before final contract documents are prepared. A signature below means only that the person received it for review.")
    for text in [
        "Buyer 1: _____________________________________________________  Date: ____________________",
        "Buyer 2: _____________________________________________________  Date: ____________________",
        "Seller / Trustee: ____________________________________________  Date: ____________________",
    ]:
        add_body(doc, text)

    out = OUTPUT_DIR / doc_name
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def convert_pdf(docx_path: Path):
    subprocess.run(
        [
            str(SOFFICE),
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(docx_path.parent),
            str(docx_path),
        ],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        timeout=120,
    )
    return docx_path.with_suffix(".pdf")


def archive_current_team_files(doc_name: str, pdf_name: str):
    ARCHIVE_DIR.mkdir(parents=True, exist_ok=True)
    existing = []
    for path in ARCHIVE_DIR.glob(f"*{doc_name}"):
        match = re.match(r"v(\d{2}) - ", path.name)
        if match:
            existing.append(int(match.group(1)))
    next_archive = max(existing, default=0) + 1
    for source in [TEAMS_PACKAGE / doc_name, TEAMS_PACKAGE / pdf_name]:
        if source.exists():
            dest = ARCHIVE_DIR / f"v{next_archive:02d} - {source.name}"
            shutil.copy2(source, dest)


def main():
    parser = argparse.ArgumentParser(description="Build Tensity term sheets from the project spreadsheet.")
    parser.add_argument(
        "--mode",
        choices=["auto", "cfd", "ps-trade-equity"],
        default="auto",
        help="auto chooses CFD when seller financing exists, otherwise P&S trade-equity when trade credit exists.",
    )
    args = parser.parse_args()

    values = docs_values()
    seller_financed = first_present(values.get("Selling Seller-Financed Principal:"), values.get("Selling Note Principal Sum:"))
    buyer_outside_financing = first_present(
        values.get("Selling Buyer Outside Financing Amount:"),
        values.get("Selling Residual Financing Amount:"),
    )
    trade_credit = values.get("Selling Trade Credit Total:")
    mode = resolve_mode(args.mode, seller_financed, trade_credit, buyer_outside_financing)
    doc_name, pdf_name = output_names(mode, trade_credit)
    version = next_package_version(doc_name)
    docx = build_docx(version, mode, doc_name)
    pdf = convert_pdf(docx)
    TEAMS_PACKAGE.mkdir(parents=True, exist_ok=True)
    archive_current_team_files(doc_name, pdf_name)
    shutil.copy2(docx, TEAMS_PACKAGE / doc_name)
    shutil.copy2(pdf, TEAMS_PACKAGE / pdf_name)
    print(f"mode={mode}")
    print(docx)
    print(pdf)
    print(TEAMS_PACKAGE / doc_name)
    print(TEAMS_PACKAGE / pdf_name)


if __name__ == "__main__":
    main()
