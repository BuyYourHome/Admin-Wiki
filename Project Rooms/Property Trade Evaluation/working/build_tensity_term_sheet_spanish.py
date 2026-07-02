from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from build_tensity_term_sheet import (
    ARCHIVE_DIR,
    OUTPUT_DIR,
    TEAMS_PACKAGE,
    VERSION_DATE,
    add_field,
    clean,
    clean_trust_name,
    convert_pdf,
    date_text,
    docs_values,
    first_present,
    looks_like_address,
    money,
    next_package_version,
    normalize_trust_date,
    numeric,
    percent,
    positive_money,
    set_cell_nowrap,
    set_cell_shading,
    set_cell_width,
    set_repeat_table_header,
    trade_properties,
)


DOC_NAME = "4121 Tensity Dr - Purchase and Sale Term Sheet - SPANISH DRAFT.docx"
PDF_NAME = DOC_NAME.replace(".docx", ".pdf")


def spanish_prepared_date() -> str:
    from build_tensity_term_sheet import RUN_DATE

    months = {
        1: "enero",
        2: "febrero",
        3: "marzo",
        4: "abril",
        5: "mayo",
        6: "junio",
        7: "julio",
        8: "agosto",
        9: "septiembre",
        10: "octubre",
        11: "noviembre",
        12: "diciembre",
    }
    return f"{RUN_DATE.day} de {months[RUN_DATE.month]} de {RUN_DATE.year}"


def date_text_es(value) -> str:
    months = {
        1: "enero",
        2: "febrero",
        3: "marzo",
        4: "abril",
        5: "mayo",
        6: "junio",
        7: "julio",
        8: "agosto",
        9: "septiembre",
        10: "octubre",
        11: "noviembre",
        12: "diciembre",
    }
    if hasattr(value, "year"):
        return f"{value.day} de {months[value.month]} de {value.year}"
    return clean(value) or "[no indicado]"


def text_date_es(value: str) -> str:
    text = clean(value)
    month_map = {
        "January": "enero",
        "February": "febrero",
        "March": "marzo",
        "April": "abril",
        "May": "mayo",
        "June": "junio",
        "July": "julio",
        "August": "agosto",
        "September": "septiembre",
        "October": "octubre",
        "November": "noviembre",
        "December": "diciembre",
    }
    match = re.match(r"^([A-Za-z]+)\s+(\d{1,2}),\s*((?:19|20)\d{2})$", text)
    if not match:
        return text
    month, day, year = match.groups()
    return f"{int(day)} de {month_map.get(month, month)} de {year}"


def note_purpose_es(value) -> str:
    text = clean(value)
    normalized = text.lower().replace("-", " ")
    if "buyer note to seller" in normalized and "seller financed principal" in normalized:
        return "Pagaré del comprador al vendedor por el principal financiado por el vendedor"
    return text


def join_parts_es(parts):
    parts = [part for part in parts if part]
    if len(parts) <= 1:
        return "".join(parts)
    if len(parts) == 2:
        return " y ".join(parts)
    return ", ".join(parts[:-1]) + " y " + parts[-1]


def add_footer(section, version: str):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.text = ""
    p.add_run("Página ")
    add_field(p, "PAGE")
    p.add_run(" de ")
    add_field(p, "NUMPAGES")
    p.add_run(f" | {version}")
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(90, 90, 90)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.style = "Heading 2"
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_terms_table(doc, rows, widths=(2.25, 4.9), font_size=9.5):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(rows[0]):
        cell = hdr.cells[i]
        set_cell_shading(cell, "D9EAF7")
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_width(cell, Inches(widths[i]))
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_width(cells[i], Inches(widths[i]))
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(clean(text))
            run.font.size = Pt(font_size)
            if i == 0:
                run.bold = True
    return table


def build_docx(version: str) -> Path:
    v = docs_values()
    trades = trade_properties()

    buyer2 = clean(v.get("Selling-Buyer2"))
    buyer_address_parts = []
    if looks_like_address(buyer2):
        buyer_address_parts.append(buyer2)
    buyer_address_parts.extend(
        part
        for part in [
            clean(v.get("Selling-Buyer Add1")),
            clean(v.get("Selling-Buyer Add1__2")),
            clean(v.get("Selling-Buyer Add2")),
        ]
        if part
    )

    seller = clean(v.get("Selling -Seller")) or clean(v.get("SellingSeller"))
    trust = clean(v.get("Trust"))
    trust_name = clean_trust_name(trust)
    trust_date = normalize_trust_date(v.get("Trust Date"), v.get("Year"))
    trust_display = f"{trust_name} dated {trust_date}" if trust_name and trust_date else trust
    trust_display_es = f"{trust_name} de fecha {text_date_es(trust_date)}" if trust_name and trust_date else trust
    trustee = clean(v.get("Trustee"))
    manager = clean(v.get("Manger") or v.get("Manager"))
    trustee_address = ", ".join(
        part
        for part in [
            clean(v.get("Trustee Address1:")),
            clean(v.get("Trustee Address2:") or v.get("Trustee Address2")),
        ]
        if part
    )
    property_line = ", ".join(part for part in [clean(v.get("Address")), clean(v.get("City-State"))] if part)
    property_short = clean(v.get("Address")) or "4121 Tensity Dr"

    purchase_price = v.get("Selling Purchase Price:")
    earnest = v.get("Selling Earnest Money:")
    cash_due = v.get("Selling Cash Due at Closing:")
    trade_credit = v.get("Selling Trade Credit Total:")
    total_credits = v.get("Selling Total Down Payment/Credits:")
    refinance_first_mortgage = first_present(
        v.get("Selling Refinance 1st Mortgage:"),
        v.get("Selling Refinance 1st Mortgage"),
    )
    seller_financed = first_present(v.get("Selling Seller-Financed Principal:"), v.get("Selling Note Principal Sum:"))
    note_payment = v.get("Selling Note Normal Payment:")

    has_trade_credit = positive_money(trade_credit) or bool(trades)
    has_refinance_first_mortgage = positive_money(refinance_first_mortgage)
    has_seller_financing = positive_money(seller_financed)
    active_buyer_funding_total = refinance_first_mortgage if has_refinance_first_mortgage else total_credits
    buyer_closing_funds = (
        max(0.0, numeric(refinance_first_mortgage) - numeric(earnest) - numeric(trade_credit))
        if has_refinance_first_mortgage
        else cash_due
    )
    buyer_closing_funds_label = (
        "fondos del refinanciamiento / primera hipoteca que se entregarían al cierre"
        if has_refinance_first_mortgage
        else "efectivo debido al cierre"
    )
    buyer_closing_funds_row_label = (
        "Fondos del refinanciamiento / primera hipoteca al cierre"
        if has_refinance_first_mortgage
        else "Efectivo debido al cierre"
    )
    active_buyer_funding_row_label = (
        "Total de fondos / créditos del refinanciamiento"
        if has_refinance_first_mortgage
        else "Total de efectivo y créditos"
    )
    trade_property_phrase = "propiedades entregadas en intercambio" if len(trades) != 1 else "propiedad entregada en intercambio"
    purchase_components = [
        f"{money(earnest)} de dinero de seriedad" if positive_money(earnest) else "",
        f"{money(buyer_closing_funds)} de {buyer_closing_funds_label}" if positive_money(buyer_closing_funds) else "",
        f"{money(trade_credit)} de crédito por {trade_property_phrase}" if positive_money(trade_credit) else "",
        f"{money(seller_financed)} de saldo financiado por el vendedor" if has_seller_financing else "",
    ]
    purchase_component_text = join_parts_es(purchase_components)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    add_footer(section, version)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Hoja de Términos de Compraventa")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Contrato de Compraventa Propuesto - {property_line}").bold = True

    prepared = doc.add_paragraph()
    prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prepared.add_run(f"Preparado: {spanish_prepared_date()} | Versión del borrador: {version}")

    status = doc.add_table(rows=1, cols=1)
    status.style = "Table Grid"
    set_cell_shading(status.cell(0, 0), "FFF2CC")
    purpose_text = (
        "Propósito: Esta hoja de términos es un resumen en español, en lenguaje sencillo, de la compraventa propuesta "
        "para revisión y conversación. No es el contrato final y no transfiere título. El Contrato de Compraventa final, "
        "la escritura, los documentos de financiamiento del comprador, cualquier pagaré del comprador, la revisión de "
        "calificación del comprador y la revisión del abogado todavía controlan. Si este resumen entra en conflicto con "
        "los documentos finales firmados, controlan los documentos finales firmados."
    )
    status.cell(0, 0).paragraphs[0].add_run(purpose_text).italic = True

    section_no = 1

    def add_numbered_heading(text):
        nonlocal section_no
        heading = add_heading(doc, f"{section_no}. {text}")
        section_no += 1
        return heading

    add_numbered_heading("Resumen en lenguaje sencillo")
    add_body(
        doc,
        f"{clean(v.get('Selling-Buyer1'))} compraría {property_short} bajo un Contrato de Compraventa propuesto "
        f"por {money(purchase_price)}. La estructura propuesta se compone de {purchase_component_text}.",
    )

    add_numbered_heading("Partes y propiedad")
    add_terms_table(
        doc,
        [
            ["Término", "Detalle"],
            ["Vendedor", seller],
            ["Estructura de firma del vendedor para documentos finales", f"{trust_display_es}, por medio de {trustee}, Fiduciario; Gerente: {manager}"],
            ["Dirección postal del fiduciario", trustee_address],
            ["Comprador", clean(v.get("Selling Note Maker (Buyer):")) or clean(v.get("Selling-Buyer1"))],
            ["Dirección postal del comprador", ", ".join(buyer_address_parts) or "[no indicada]"],
            ["Propiedad", property_line],
            ["Condado", clean(v.get("County")) or "[no indicado]"],
            ["Número de parcela", clean(v.get("Property Parcel ID")) or "[no indicado]"],
            ["Resumen de descripción legal", clean(v.get("Brief Legal Description")) or "[no indicado]"],
        ],
    )

    add_numbered_heading("Estructura propuesta de compra")
    structure_parts = ["el dinero de seriedad", buyer_closing_funds_label]
    if has_trade_credit:
        structure_parts.append(f"el crédito por {trade_property_phrase}")
    if has_seller_financing:
        structure_parts.append("el saldo financiado por el vendedor")
    add_body(
        doc,
        "El Contrato de Compraventa propuesto separa "
        f"{join_parts_es(structure_parts)} para que cada parte pueda mostrarse claramente en los documentos finales.",
    )
    purchase_rows = [
        ["Término", "Detalle"],
        ["Precio de compra propuesto", money(purchase_price)],
        ["Dinero de seriedad al firmar", money(earnest)],
        [buyer_closing_funds_row_label, money(buyer_closing_funds)],
        [active_buyer_funding_row_label, money(active_buyer_funding_total)],
    ]
    if has_trade_credit:
        purchase_rows.insert(4, [f"Crédito por {trade_property_phrase}", money(trade_credit)])
    if has_seller_financing:
        purchase_rows.append(["Saldo financiado por el vendedor", money(seller_financed)])
    allocation = " + ".join(
        part
        for part in [
            f"{money(earnest)} de dinero de seriedad" if positive_money(earnest) else "",
            f"{money(buyer_closing_funds)} de {buyer_closing_funds_label}" if positive_money(buyer_closing_funds) else "",
            f"{money(trade_credit)} de crédito por intercambio" if positive_money(trade_credit) else "",
            f"{money(seller_financed)} de saldo financiado por el vendedor" if has_seller_financing else "",
        ]
        if part
    )
    purchase_rows.append(["Cómo se compone el precio de compra", f"{allocation} = {money(purchase_price)} de precio de compra propuesto"])
    add_terms_table(doc, purchase_rows)

    if has_trade_credit:
        add_numbered_heading("Propiedades de intercambio incluidas como crédito")
        add_body(
            doc,
            "Se espera que el comprador transfiera título, derechos de propiedad u otros documentos requeridos de propiedad "
            f"para las {trade_property_phrase} como parte del trato. El crédito no es final hasta que la documentación requerida, "
            "aprobaciones, posesión y demás condiciones de cierre sean completadas y aceptadas.",
        )
        trade_rows = [["Propiedad", "Tipo", "Crédito", "Lo que falta"]]
        type_map = {"Single wide": "Sencilla", "Double wide": "Doble"}
        for item in trades:
            trade_rows.append(
                [
                    clean(item.get("Property Address")),
                    type_map.get(clean(item.get("Home Type")), clean(item.get("Home Type"))),
                    money(item.get("Net Trade Credit")),
                    "Requerido antes de contar el crédito final",
                ]
            )
        trade_rows.append(["Crédito total", "", money(trade_credit), "Sujeto al anexo de intercambio"])
        table = add_terms_table(doc, trade_rows, widths=(2.1, 1.0, 1.05, 3.0), font_size=8.5)
        for row in table.rows:
            for cell in row.cells:
                set_cell_nowrap(cell)
        for cell in table.rows[-1].cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True

        add_numbered_heading("Condiciones para que el crédito de intercambio cuente")
        for text in [
            "El crédito de intercambio es condicional. Antes de que cuente como final, el comprador debe entregar documentos de propiedad aceptables, liberaciones de gravámenes o prueba de pago, posesión/llaves, información de inquilinos y rentas cuando corresponda, y aprobaciones requeridas de terceros.",
            "El anexo final de propiedades de intercambio debe detallar transferencia de propiedad, posesión, aprobaciones, gravámenes, rentas de inquilinos, verificación, condición y remedios.",
            "Si un requisito no se entrega y acepta, los documentos finales deben indicar si el crédito faltante se convierte en efectivo adicional debido, ajuste al precio de compra, asunto de condición de cierre u otro remedio acordado. Los valores propuestos son cifras de trabajo hasta ser aceptados.",
        ]:
            add_body(doc, text)

    if has_seller_financing:
        add_numbered_heading("Saldo financiado por el vendedor")
        add_body(
            doc,
            "Se espera que el saldo financiado por el vendedor se muestre en un pagaré firmado por el comprador como parte "
            "del paquete de cierre de compraventa, salvo que el abogado recomiende una estructura final diferente.",
        )
        add_terms_table(
            doc,
            [
                ["Término", "Detalle"],
                ["Comprador que firma el pagaré", clean(v.get("Selling Note Maker (Buyer):"))],
                ["Vendedor que recibe pagos", clean(v.get("Selling Note Payee (Seller):"))],
                ["Propósito del pagaré", note_purpose_es(v.get("Selling Note Purpose/Direction:"))],
                ["Suma principal", money(seller_financed)],
                ["Tasa de interés", percent(v.get("Selling Note Interest Rate:"))],
                ["Pago de principal e interés", money(note_payment)],
                ["Meses de cuotas", clean(v.get("Selling Note Installment Months:"))],
                ["Fecha del primer pago", date_text_es(v.get("Selling Note First Payment Date:"))],
                ["Fecha final programada de pago", date_text_es(v.get("Selling Note Final Scheduled Payment Date:"))],
                ["Pago global", money(v.get("Selling Note Balloon Payment:"))],
                ["Texto de pago global", clean(v.get("Selling Note Balloon Term:")) or "[no indicado]"],
            ],
        )

        add_numbered_heading("Pago mensual estimado")
        tax = v.get("Tax Escrow:")
        insurance = v.get("Property Insurance:")
        hoa = v.get("HOA:")
        total_payment = numeric(note_payment) + numeric(tax) + numeric(insurance) + numeric(hoa)
        add_terms_table(
            doc,
            [
                ["Término", "Detalle"],
                ["Pago de principal e interés", money(note_payment)],
                ["Reserva para impuestos", money(tax)],
                ["Reserva para seguro", money(insurance)],
                ["HOA", money(hoa)],
                ["Pago mensual total estimado si se cobran reservas", money(total_payment)],
            ],
        )

    add_numbered_heading("Documentos esperados después")
    expected_docs = [
        f"Contrato de Compraventa para {property_short}.",
        "Escritura del vendedor al comprador que se entregará al cierre, sujeta a documentos finales aprobados por abogado.",
        "Documentos de refinanciamiento / primera hipoteca del comprador requeridos por el prestamista del comprador, si corresponde.",
        "Estado de cierre del comprador y vendedor que muestre dinero de seriedad, fondos de refinanciamiento entregados al cierre, financiamiento del vendedor, prorrateos y otros ajustes de cierre.",
        "Cualquier divulgación, reconocimiento o declaración jurada de cierre requerida por abogado.",
    ]
    if has_seller_financing:
        expected_docs.insert(
            3,
            f"Pagaré del comprador al vendedor que muestre {money(seller_financed)} de principal, {percent(v.get('Selling Note Interest Rate:'))} de interés y {money(note_payment)} de pago mensual de principal e interés, salvo que el abogado indique una estructura final diferente.",
        )
        expected_docs.insert(
            4,
            "Documento de garantía para el financiamiento del vendedor, escritura de fideicomiso u otro documento de garantía aprobado por abogado si se requiere para el saldo financiado por el vendedor.",
        )
    if has_trade_credit:
        expected_docs.insert(2, f"Anexo o exhibición de propiedades de intercambio para las {trade_property_phrase}.")
        expected_docs.insert(
            3,
            "Documentos de transferencia de propiedad/título para las propiedades de intercambio, incluyendo título, bill of sale, escritura, cesión, aprobación del parque o documento similar según el tipo de propiedad.",
        )
        expected_docs.insert(
            4,
            "Liberaciones de gravámenes, confirmaciones de pago u otra prueba de que gravámenes e intereses adversos de las propiedades de intercambio estén resueltos según lo requiera el acuerdo final.",
        )
    for text in expected_docs:
        add_body(doc, text)

    add_numbered_heading("Estado de confirmación")
    status_rows = [["Estado", "Asunto", "Notas"]]
    if not clean(v.get("Property Parcel ID")):
        status_rows.append(["Necesita atención", "Número de parcela", "El número de parcela todavía no está completado."])
    if not clean(v.get("Brief Legal Description")) and not clean(v.get("Legal Description")):
        status_rows.append(["Necesita atención", "Descripción legal", "La descripción legal todavía no está completada."])
    buyer_address = ", ".join(buyer_address_parts)
    if buyer_address:
        status_rows.append(["Resuelto", "Dirección postal del comprador", f"Indicada como {buyer_address}."])
    else:
        status_rows.append(["Necesita atención", "Dirección postal del comprador", "La dirección postal del comprador todavía no está completada."])
    trust_text = f"{trust_display_es}, por medio de {trustee}, Fiduciario; Gerente: {manager}"
    trust_date_resolved = bool(
        trust_display
        and trustee
        and manager
        and re.search(r"\bdated\b", trust_display, flags=re.IGNORECASE)
        and re.search(r"\b[A-Za-z]+ \d{1,2}, (?:19|20)\d{2}\b", trust_display)
    )
    if trust_date_resolved:
        status_rows.append(["Resuelto", "Redacción de fecha del fideicomiso", f"Se resuelve como: {trust_text}."])
    else:
        status_rows.append(
            [
                "Necesita atención",
                "Redacción de fecha del fideicomiso",
                "El nombre, fecha, fiduciario o gerente del fideicomiso está incompleto o no se resuelve limpiamente.",
            ]
        )
    if has_trade_credit:
        status_rows.append(
            [
                "Necesita atención",
                "Condiciones de propiedades de intercambio",
                "Se requieren documentos de propiedad, estado de gravámenes, aprobaciones de terceros, información de inquilinos/rentas cuando corresponda, posesión/llaves y aceptación de condición antes de tratar el crédito como final.",
            ]
        )
    status_rows.append(
        [
            "Necesita atención",
            "Revisión de abogado",
            "El abogado debe confirmar si la estructura de compraventa, escritura, financiamiento / primera hipoteca del comprador y términos de financiamiento del vendedor son la estructura final preferida.",
        ]
    )
    add_terms_table(doc, status_rows, widths=(1.25, 1.95, 3.95), font_size=8.8)

    add_numbered_heading("Reconocimiento de revisión")
    add_body(
        doc,
        "Las partes pueden usar esta hoja de términos para conversar sobre el trato propuesto antes de preparar los documentos finales. "
        "Una firma abajo significa solamente que la persona la recibió para revisión.",
    )
    for text in [
        "Comprador 1: _____________________________________________________  Fecha: ____________________",
        "Comprador 2: _____________________________________________________  Fecha: ____________________",
        "Vendedor / Fiduciario: ____________________________________________  Fecha: ____________________",
    ]:
        add_body(doc, text)

    out = OUTPUT_DIR / DOC_NAME
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def archive_current_team_files():
    ARCHIVE_DIR.mkdir(parents=True, exist_ok=True)
    existing = []
    for path in ARCHIVE_DIR.glob(f"*{DOC_NAME}"):
        match = re.match(r"v(\d{2}) - ", path.name)
        if match:
            existing.append(int(match.group(1)))
    next_archive = max(existing, default=0) + 1
    for source in [TEAMS_PACKAGE / DOC_NAME, TEAMS_PACKAGE / PDF_NAME]:
        if source.exists():
            shutil.copy2(source, ARCHIVE_DIR / f"v{next_archive:02d} - {source.name}")


def main():
    version = next_package_version(DOC_NAME)
    docx = build_docx(version)
    pdf = convert_pdf(docx)
    TEAMS_PACKAGE.mkdir(parents=True, exist_ok=True)
    archive_current_team_files()
    shutil.copy2(docx, TEAMS_PACKAGE / DOC_NAME)
    shutil.copy2(pdf, TEAMS_PACKAGE / PDF_NAME)
    print(docx)
    print(pdf)
    print(TEAMS_PACKAGE / DOC_NAME)
    print(TEAMS_PACKAGE / PDF_NAME)


if __name__ == "__main__":
    main()
