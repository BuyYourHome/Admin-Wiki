from __future__ import annotations

import re
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Codex\Wiki Files\Project Rooms\Jennys Drawings")
SOURCE = ROOT / "outputs" / "Drawn by Grace - Review Manuscript.md"
IMAGE_DIR = ROOT / "sources" / "images" / "bible-album"
OUTPUT = ROOT / "outputs" / "Drawn by Grace - Visual Companion - REVIEW.docx"


NAVY = "203748"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GOLD = "B08D3B"
MUTED = "666666"
PARCHMENT = "FFFCF2"
WHITE = "FFFFFF"


PLATES = {
    "JD-001": ("JD-001-you-are-my-shield.jpg", "You are my Shield", "Genesis 15:1 | Drawing dated September 13, 2017"),
    "JD-002": ("JD-002-you-are-the-god-who-sees.jpg", "You are the God who sees", "Genesis 16:13 | Drawing dated October 13, 2017"),
    "JD-003": ("JD-003-your-faith-needs-a-fight.jpg", "Your faith needs a fight", "Genesis 17:15-21 | Drawing dated February 13, 2018"),
    "JD-004": ("JD-004-grant-me-success-today.jpg", "O Lord, grant me success today", "Genesis 24:12 | Drawing date to be confirmed"),
    "JD-005": ("JD-005-to-break-out-look-in.jpg", "To break out, you've got to look in", "Genesis 27:30-40 | Drawing dated March 15, 2018"),
    "JD-006": ("JD-006-admit-you-are-broken.jpg", "You'll never be fixed until you admit you're broken", "Genesis 32:24-32 | Drawing dated June 14, 2018"),
    "JD-007": ("JD-007-maker-has-naming-rights.jpg", "Only the Maker has naming rights; rename your sorrow", "Genesis 35:17-18 | Drawing dated March 9, 2018"),
    "JD-008": ("JD-008-prison-to-palace.jpg", "God's plan: prison to palace", "Genesis 41 | Drawing dated September 21, 2017"),
    "JD-009": ("JD-009-discover-the-source-in-the-stretch.jpg", "In the stretch, discover the source", "Exodus 4 | Drawing dated January 15, 2017"),
    "JD-010": ("JD-010-past-is-a-guidepost.jpg", "The past is a guidepost to tomorrow", "Exodus 1-2 | Drawing dated September 19, 2017"),
    "JD-011": ("JD-011-rescued.jpg", "Rescued", "Exodus 3:8 | Drawing dated October 5, 2017"),
    "JD-012": ("JD-012-what-is-in-your-hands.jpg", "What has God put in your hands?", "Exodus 4:17 | Drawing dated November 1, 2017"),
    "JD-013": ("JD-013-take-the-first-step.jpg", "God positions you; take the first step", "Exodus 23:20 | Drawing dated January 4, 2018"),
    "JD-014": ("JD-014-what-has-god-prepared.jpg", "Are you missing what God has prepared?", "Exodus 23:20-21 | Drawing dated January 3, 2018"),
    "JD-015": ("JD-015-talking-with-god.jpg", "Talking with God...or to Him?", "Exodus 25:22 | Drawing dated October 26, 2017"),
    "JD-016": ("JD-016-what-limits-you.jpg", "What is limiting you?", "Numbers 13:33 | Drawing dated March 21, 2018"),
    "JD-017": ("JD-017-guide-me.jpg", "Guide me", "Deuteronomy 1:21-30 | Drawing dated September 8, 2017"),
    "JD-018": ("JD-018-ingredients-not-cake.jpg", "Ingredients, not cake", "Deuteronomy 30:15 | Drawing dated November 30, 2017"),
}


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in edges:
            continue
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edges[edge].items():
            element.set(qn("w:" + key), str(value))


def set_cell_margins(cell, top=160, start=160, bottom=160, end=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    col = OxmlElement("w:gridCol")
    col.set(qn("w:w"), str(width_dxa))
    grid.append(col)
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa))
            tc_w.set(qn("w:type"), "dxa")


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def parse_markdown(path):
    text = path.read_text(encoding="utf-8")
    sections = []
    current = None
    for line in text.splitlines():
        if line.startswith("## "):
            if current:
                sections.append(current)
            current = {"heading": line[3:].strip(), "paragraphs": [], "plate": None}
        elif current is not None:
            plate = re.fullmatch(r"\[\[PLATE:(JD-\d{3})\]\]", line.strip())
            if plate:
                current["plate"] = plate.group(1)
            elif line.strip():
                current["paragraphs"].append(line.strip())
    if current:
        sections.append(current)
    return sections


def add_body_paragraph(doc, markdown_text):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    parts = re.split(r"(\*[^*]+\*|\"[^\"]+\")", markdown_text)
    for part in parts:
        if not part:
            continue
        if part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)
    return p


def add_plate(doc, plate_id):
    filename, title, meta = PLATES[plate_id]
    image_path = IMAGE_DIR / filename
    if not image_path.exists():
        raise FileNotFoundError(image_path)

    doc.add_page_break()
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, 8496)  # named art_plate_frame override: 5.9 inches
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=180, start=180, bottom=160, end=180)
    set_cell_border(
        cell,
        top={"val": "double", "sz": 14, "color": GOLD, "space": 0},
        bottom={"val": "double", "sz": 14, "color": GOLD, "space": 0},
        left={"val": "double", "sz": 14, "color": GOLD, "space": 0},
        right={"val": "double", "sz": 14, "color": GOLD, "space": 0},
    )
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), PARCHMENT)
    cell._tc.get_or_add_tcPr().append(shading)

    image_p = cell.paragraphs[0]
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.paragraph_format.space_before = Pt(0)
    image_p.paragraph_format.space_after = Pt(6)
    with Image.open(image_path) as im:
        width_px, height_px = im.size
    max_width = 5.45
    # Leave room for Word's mandatory paragraph after a table so a framed
    # full-page plate does not force an otherwise blank following page.
    max_height = 7.45
    ratio = width_px / height_px
    width = min(max_width, max_height * ratio)
    height = width / ratio
    if height > max_height:
        height = max_height
        width = height * ratio
    run = image_p.add_run()
    shape = run.add_picture(str(image_path), width=Inches(width), height=Inches(height))
    shape._inline.docPr.set("descr", f"Jenny Browning drawing: {title}")

    caption = cell.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(0)
    title_run = caption.add_run(title)
    set_run_font(title_run, size=10.5, color=NAVY, bold=True)
    meta_p = cell.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_before = Pt(1)
    meta_p.paragraph_format.space_after = Pt(0)
    meta_run = meta_p.add_run(meta)
    set_run_font(meta_run, size=8.5, color=MUTED, italic=True)


def build():
    sections = parse_markdown(SOURCE)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    sec.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run("DRAWN BY GRACE  |  A VISUAL COMPANION")
    set_run_font(hr, size=8.5, color=MUTED, bold=True)
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("DRAWN BY GRACE  |  ")
    set_run_font(fr, size=8.5, color=MUTED)
    add_page_field(fp)

    # Editorial-cover header pattern with narrative_proposal typography.
    for _ in range(5):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = kicker.add_run("A VISUAL COMPANION")
    set_run_font(kr, size=10.5, color=GOLD, bold=True)
    kicker.paragraph_format.space_after = Pt(16)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Drawn by Grace")
    set_run_font(tr, size=30, color=NAVY, bold=True)
    title.paragraph_format.space_after = Pt(9)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("A Visual Companion to The Gracious Millionaire")
    set_run_font(sr, size=15, color=DARK_BLUE, italic=True)
    subtitle.paragraph_format.space_after = Pt(34)
    artist = doc.add_paragraph()
    artist.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = artist.add_run("Featuring drawings by Jenny Browning")
    set_run_font(ar, size=11.5, color=NAVY, bold=True)
    artist.paragraph_format.space_after = Pt(6)
    review = doc.add_paragraph()
    review.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = review.add_run("REVIEW MANUSCRIPT  |  JULY 2026")
    set_run_font(rr, size=9.5, color=MUTED)

    for index, section in enumerate(sections):
        doc.add_page_break()
        heading = doc.add_paragraph(section["heading"], style="Heading 1")
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for para in section["paragraphs"]:
            add_body_paragraph(doc, para)
        if section["plate"]:
            add_plate(doc, section["plate"])

    doc.core_properties.title = "Drawn by Grace: A Visual Companion to The Gracious Millionaire"
    doc.core_properties.subject = "Review manuscript featuring drawings by Jenny Browning"
    doc.core_properties.author = "Wes and Jenny Browning"
    doc.core_properties.comments = "Review manuscript prepared with AI assistance; see provenance note."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
