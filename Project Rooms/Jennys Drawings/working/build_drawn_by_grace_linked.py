from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor

import build_drawn_by_grace as base


ROOT = Path(r"C:\Codex\Wiki Files\Project Rooms\Jennys Drawings")
OUTPUT = ROOT / "outputs" / "Drawn by Grace - Linked Review Manuscript.docx"
TEAMS_FOLDER_URL = (
    "https://lifeisanadventure.sharepoint.com/sites/SellYourHome/Shared%20Documents/"
    "Marketing/Jennys%20Drawings%20-%20Gracious%20Millionaire%20Companion"
)


def add_external_link(paragraph, text, url, *, bold=False, size=11):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), base.BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(int(size * 2)))
    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), "Calibri")
    font.set(qn("w:hAnsi"), "Calibri")
    run_properties.extend([font, color, underline, size_node])
    if bold:
        run_properties.append(OxmlElement("w:b"))
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_internal_link(paragraph, text, anchor, *, size=10.5):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), base.BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(int(size * 2)))
    run_properties.extend([color, underline, size_node])
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bookmark(paragraph, name, bookmark_id):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def slug(index):
    return f"section_{index:02d}"


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def set_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (16, base.BLUE, 18, 10),
        "Heading 2": (13, base.BLUE, 12, 6),
        "Heading 3": (12, base.DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("DRAWN BY GRACE  |  LINKED REVIEW MANUSCRIPT")
    base.set_run_font(run, size=8.5, color=base.MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("DRAWN BY GRACE  |  ")
    base.set_run_font(run, size=8.5, color=base.MUTED)
    add_page_number(footer)


def add_cover(doc):
    for _ in range(5):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_run_font(kicker.add_run("A LIGHTWEIGHT VISUAL COMPANION"), size=10.5, color=base.GOLD, bold=True)
    kicker.paragraph_format.space_after = Pt(16)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_run_font(title.add_run("Drawn by Grace"), size=30, color=base.NAVY, bold=True)
    title.paragraph_format.space_after = Pt(9)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_run_font(
        subtitle.add_run("A Visual Companion to The Gracious Millionaire"),
        size=15,
        color=base.DARK_BLUE,
        italic=True,
    )
    subtitle.paragraph_format.space_after = Pt(30)
    artist = doc.add_paragraph()
    artist.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_run_font(artist.add_run("Featuring drawings by Jenny Browning"), size=11.5, color=base.NAVY, bold=True)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(16)
    add_external_link(note, "Open the complete drawing folder in Teams", TEAMS_FOLDER_URL, bold=True, size=11)
    review = doc.add_paragraph()
    review.alignment = WD_ALIGN_PARAGRAPH.CENTER
    review.paragraph_format.space_before = Pt(18)
    base.set_run_font(review.add_run("LINKED REVIEW MANUSCRIPT  |  JULY 2026"), size=9.5, color=base.MUTED)


def add_contents(doc, sections):
    doc.add_page_break()
    heading = doc.add_paragraph("Clickable Outline", style="Heading 1")
    add_bookmark(heading, "outline", 1)
    intro = doc.add_paragraph(
        "Use this page to jump within the manuscript. Numbered companion sections also include a secure link to the full-resolution drawing stored in the Marketing channel."
    )
    intro.paragraph_format.space_after = Pt(14)
    for index, section in enumerate(sections, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_after = Pt(5)
        add_internal_link(p, section["heading"], slug(index), size=10.5)


def add_art_link(doc, plate_id):
    filename, title, meta = base.PLATES[plate_id]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    base.set_run_font(p.add_run("DRAWING  |  "), size=10, color=base.GOLD, bold=True)
    add_external_link(p, title, f"{TEAMS_FOLDER_URL}/{filename}", bold=True, size=11)
    meta_p = doc.add_paragraph(meta)
    meta_p.paragraph_format.space_before = Pt(0)
    meta_p.paragraph_format.space_after = Pt(4)
    for run in meta_p.runs:
        base.set_run_font(run, size=9, color=base.MUTED, italic=True)
    open_p = doc.add_paragraph()
    open_p.paragraph_format.space_before = Pt(0)
    open_p.paragraph_format.space_after = Pt(0)
    add_external_link(open_p, "Open the full-resolution drawing in Teams", f"{TEAMS_FOLDER_URL}/{filename}", size=10.5)


def build():
    parsed = base.parse_markdown(base.SOURCE)
    sections = parsed[1:]  # The first source section is represented by the designed cover.
    doc = Document()
    set_styles(doc)
    add_cover(doc)
    add_contents(doc, sections)

    bookmark_id = 10
    for index, section in enumerate(sections, start=1):
        doc.add_page_break()
        heading = doc.add_paragraph(section["heading"], style="Heading 1")
        add_bookmark(heading, slug(index), bookmark_id)
        bookmark_id += 1
        for paragraph in section["paragraphs"]:
            base.add_body_paragraph(doc, paragraph)
        if section["plate"]:
            add_art_link(doc, section["plate"])
        back = doc.add_paragraph()
        back.paragraph_format.space_before = Pt(14)
        add_internal_link(back, "Back to clickable outline", "outline", size=9.5)

    doc.core_properties.title = "Drawn by Grace: Linked Review Manuscript"
    doc.core_properties.subject = "Lightweight review manuscript with Teams-hosted drawing links"
    doc.core_properties.author = "Wes and Jenny Browning"
    doc.core_properties.comments = "Review manuscript prepared with AI assistance; linked drawings are stored in the Buy Your Home Marketing channel."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
